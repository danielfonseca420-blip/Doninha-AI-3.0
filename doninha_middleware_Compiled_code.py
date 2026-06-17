# Doninha IA Standalone
from __future__ import annotations

from __future__ import annotations
from langchain_community.vectorstores import Chroma
from typing import Any, Dict
from typing import Any, Dict, List, Optional
from typing import Any, Dict, List, Optional, Tuple
from typing import Any, Dict, Optional
from typing import Dict, List, Optional
from typing import Dict, List, Optional, Any
from typing import Dict, List, Optional, Any, Tuple
from typing import Dict, List, Optional, Tuple
from typing import Dict, List, Optional, Tuple, Any
from typing import Dict, List, Tuple, Optional
from typing import List
from typing import List, Optional, Tuple
from typing import List, Tuple, Iterator
from typing import List, Tuple, Optional
from typing import Optional
from typing import Optional, Dict, Any
from typing import Optional, List
import json
import logging
import math
import os
import re
import sys
import time

from agente_busca_web import run_search_for_context
from collections import Counter
from dataclasses import dataclass
from dataclasses import dataclass, field
from docx import Document
from enum import Enum
from langchain_community.embeddings import HuggingFaceEmbeddings
from numpy import dot
from numpy.linalg import norm
from pathlib import Path
from pydantic import BaseModel
from sentence_transformers import SentenceTransformer
from torch.utils.data import Dataset
from transformers import AutoModel, AutoTokenizer
from uuid import uuid4
import argparse
import ollama
import sentencepiece as spm
import torch
import torch.nn as nn
import torch.nn.functional as F
import uvicorn
import yaml


# ========== layer_titles.py ==========
"""
Nomes das camadas L1L7 usados nas instrues de gerao de texto.
"""

LAYER_TITLES = {
    "l1": "Demarcao de Conceitos Fundamentais",
    "l2": "Premissas e proposies centrais",
    "l3": "Anlise da Estrutura Lgico-filosfica",
    "l4": "Comparao da equivalncia entre Estrutura formal e Mundo Emprico",
    "l5": "Sntese Intermediria derivada das etapas anteriores",
    "l6": "Concluso do raciocnio",
    "l7": "Sntese Final e Redao",
}
# ========== config_loader.py ==========
"""
Carregamento de configurao centralizada.
==========================================
L config.yaml (ou variveis de ambiente) e expe um nico dicionrio
para pipeline, API e agentes.
"""

import os
from pathlib import Path
from typing import Any, Dict

# Diretrio raiz do projeto
PROJECT_ROOT = Path(__file__).resolve().parent
CONFIG_PATH = PROJECT_ROOT / "config.yaml"


def load_config(config_path: str | Path | None = None) -> Dict[str, Any]:
    """
    Carrega a configurao a partir de config.yaml.
    Se o arquivo no existir ou estiver incompleto, usa defaults e env.
    """
    path = Path(config_path) if config_path else CONFIG_PATH
    config: Dict[str, Any] = {
        "knowledge_base": {
            "path": "",
            "chroma_path": os.getenv("VECTOR_DB_PATH", "meu_vector_db"),
            "default_kb": "",
            "domain_specific_kbs": {},
        },
        "l3": {
            "model_path": "truth_scoring_model.pt",
            "backbone": "bert-base-multilingual-cased",
        },
        "l4": {
            "russell_concepts_path": "l4_russell_concepts.json",
        },
        "l4_chain_verification": {
            "provider": os.getenv("L4_COVE_PROVIDER", "template"),
            "ollama_model": os.getenv("L4_COVE_OLLAMA_MODEL", "doninha8:latest"),
            "ollama_host": os.getenv("L4_COVE_OLLAMA_HOST", "http://localhost:11434"),
            "custom_lm_path": "",
        },
        "generation": {
            "provider": os.getenv("GENERATION_PROVIDER", "ollama"),
            "custom_lm_path": "",
            "ollama_model": os.getenv("OLLAMA_MODEL", "doninha8:latest"),
            "ollama_host": os.getenv("OLLAMA_HOST", "http://localhost:11434"),
        },
        "finalization": {
            "provider": os.getenv("FINALIZATION_PROVIDER", "ollama"),
            "custom_lm_path": "",
            "ollama_model": os.getenv("FINALIZATION_OLLAMA_MODEL", "doninha8:latest"),
            "ollama_host": os.getenv("FINALIZATION_OLLAMA_HOST", "http://localhost:11434"),
        },
        "l7": {
            "provider": os.getenv("L7_PROVIDER", "ollama"),
            "custom_lm_path": "",
            "model": os.getenv("L7_MODEL", "doninha8:latest"),
            "ollama_host": os.getenv("L7_OLLAMA_HOST", "http://localhost:11434"),
        },
        "agent": {
            "use_agent": os.getenv("USE_AGENT", "false").lower() == "true",
            "vector_db_path": os.getenv("VECTOR_DB_PATH", "meu_vector_db"),
            "embedding_model": "sentence-transformers/all-MiniLM-L6-v2",
        },
        "api": {
            "host": os.getenv("API_HOST", "0.0.0.0"),
            "port": int(os.getenv("API_PORT", "8000")),
        },
        "chat": {
            "max_turns_in_context": 10,
        },
    }

    if path.exists():
        try:
            import yaml
            with open(path, "r", encoding="utf-8") as f:
                loaded = yaml.safe_load(f)
            if isinstance(loaded, dict):
                _deep_merge(config, loaded)
        except Exception:
            pass

    # Resolve paths relativos ao projeto
    for key in ["model_path", "russell_concepts_path", "custom_lm_path", "path", "default_kb"]:
        for section in ["l3", "l4", "l4_chain_verification", "generation", "finalization", "l7", "knowledge_base"]:
            if section in config and key in config[section]:
                val = config[section][key]
                if val and not Path(val).is_absolute():
                    config[section][key] = str(PROJECT_ROOT / val)

    if "knowledge_base" in config and isinstance(config["knowledge_base"].get("domain_specific_kbs"), dict):
        for name, value in config["knowledge_base"]["domain_specific_kbs"].items():
            if value and not Path(value).is_absolute():
                config["knowledge_base"]["domain_specific_kbs"][name] = str(PROJECT_ROOT / value)

    return config


def _deep_merge(base: Dict[str, Any], override: Dict[str, Any]) -> None:
    """Mescla override em base recursivamente."""
    for k, v in override.items():
        if k in base and isinstance(base[k], dict) and isinstance(v, dict):
            _deep_merge(base[k], v)
        else:
            base[k] = v
# ========== chat_session.py ==========
"""
Sesso de chat com histrico.
=============================
Mantm as ltimas N trocas (usurio/assistente) e monta o contexto
para o pipeline ou para o gerador.
"""

from dataclasses import dataclass, field
from typing import List, Optional, Tuple

@dataclass
class Turn:
    role: str  # "user" | "assistant"
    content: str


class ChatSession:
    """
    Histrico de mensagens para dilogo multi-turno.
    """

    def __init__(self, max_turns: int = 10):
        self.max_turns = max(1, max_turns)
        self.turns: List[Turn] = []

    def add_user(self, content: str) -> None:
        self.turns.append(Turn(role="user", content=content.strip()))

    def add_assistant(self, content: str) -> None:
        self.turns.append(Turn(role="assistant", content=content.strip()))

    def get_context_for_prompt(self, current_prompt: str, max_turns_in_context: Optional[int] = None) -> str:
        """
        Retorna um nico texto com as ltimas N trocas + pergunta atual,
        para ser usado como contexto (ex.: prefixo da pergunta ou resumo).
        """
        n = max_turns_in_context if max_turns_in_context is not None else self.max_turns
        n = max(0, n)
        recent = self.turns[-n * 2 :] if n else []  # pares user/assistant
        parts = []
        for t in recent:
            prefix = "Usurio" if t.role == "user" else "Assistente"
            parts.append(f"{prefix}: {t.content}")
        if parts:
            parts.append(f"Usurio: {current_prompt}")
            return "\n".join(parts)
        return current_prompt

    def get_last_user_prompt(self) -> str:
        """Retorna a ltima mensagem do usurio (para pipeline que no usa contexto)."""
        for t in reversed(self.turns):
            if t.role == "user":
                return t.content
        return ""

    def clear(self) -> None:
        self.turns.clear()

    def turn_count(self) -> int:
        return len(self.turns)
# ========== paraconsistent_rules.py ==========
"""
Conjunto de regras para sistema paraconsistente (LPA)
======================================================
Extrado de data/Fuzzy.txt  Lgica Paraconsistente Anotada (da Costa et al.).

Convenes do documento:
  -  (mu) = grau de crena   [0,1], eixo x no QUPC
  -  (lambda) = grau de descrena  [0,1], eixo y no QUPC
  - Gc  = Grau de Certeza     =       [1, 1]
  - Gct = Grau de Contradio =  +   1  [1, 1]

Valores de controle (Figura 3):
  Vscc  = Valor superior de controle de certeza     = 1/2
  Vicc  = Valor inferior de controle de certeza    = -1/2
  Vscct = Valor superior de controle de contradio = 1/2
  Vicct = Valor inferior de controle de contradio = -1/2

Doze estados lgicos (reticulado discretizado):
  T, V, F, , QV, QF e regies de transio (QFV, F, V, QV, TV, QVT, etc.).
"""

from dataclasses import dataclass
from typing import List, Tuple, Iterator
import re
import os

#  Constantes extradas do Fuzzy.txt 
VSCC = 0.5   # Valor superior de controle de certeza
VICC = -0.5  # Valor inferior de controle de certeza
VSCC_T = 0.5   # Valor superior de controle de contradio
VICC_T = -0.5  # Valor inferior de controle de contradio

# Doze estados lgicos do reticulado (para-analisador)
STATE_T = "Inconsistente"           # T  = (1,1)
STATE_V = "Verdadeiro"              # V  = (1,0)
STATE_F = "Falso"                   # F  = (0,1)
STATE_BOT = "Indeterminado"         #   = (0,0)
STATE_QV = "Quase_Verdadeiro"       # QV
STATE_QF = "Quase_Falso"            # QF
STATE_QF_TO_V = "QF_to_V"
STATE_BOT_TO_F = "Indeterminado_to_F"
STATE_BOT_TO_V = "Indeterminado_to_V"
STATE_QV_TO_BOT = "QV_to_Indeterminado"
STATE_T_TO_V = "Inconsistente_to_V"
STATE_QV_TO_T = "QV_to_Inconsistente"

ALL_STATES_12: List[str] = [
    STATE_T, STATE_V, STATE_F, STATE_BOT,
    STATE_QV, STATE_QF,
    STATE_QF_TO_V, STATE_BOT_TO_F, STATE_BOT_TO_V,
    STATE_QV_TO_BOT, STATE_T_TO_V, STATE_QV_TO_T,
]


@dataclass
class ParaconsistentRules:
    """Parmetros do sistema paraconsistente (ajustveis)."""
    vscc: float = VSCC
    vicc: float = VICC
    vscct: float = VSCC_T
    vicct: float = VICC_T

    @staticmethod
    def gc(mu: float, lam: float) -> float:
        """Grau de Certeza: Gc =     [1, 1]."""
        return mu - lam

    @staticmethod
    def gct(mu: float, lam: float) -> float:
        """Grau de Contradio: Gct =  +   1  [1, 1]."""
        return mu + lam - 1.0

    def state_12(self, mu: float, lam: float) -> str:
        """
        Para-analisador: discretiza (, ) em um dos 12 estados lgicos.
        Regras conforme Fuzzy.txt  regies no QUPC delimitadas por Vscc, Vicc, Vscct, Vicct.
        """
        gc = self.gc(mu, lam)
        gct = self.gct(mu, lam)

        # Alto grau de contradio positiva  Inconsistente (T)
        if gct >= self.vscct:
            if gc >= self.vscc:
                return STATE_T_TO_V      # transio TV
            if gc <= self.vicc:
                return STATE_QV_TO_T     # transio QVT (ou TF)
            return STATE_T
        # Alto grau de contradio negativa  Indeterminado ()
        if gct <= self.vicct:
            if gc >= self.vscc:
                return STATE_BOT_TO_V
            if gc <= self.vicc:
                return STATE_BOT_TO_F
            return STATE_BOT
        # Contradio baixa (zona central em Gct)
        if gc >= self.vscc:
            return STATE_V if gct <= 0 else STATE_QV
        if gc <= self.vicc:
            return STATE_F if gct <= 0 else STATE_QF
        # Certeza em zona intermediria
        if gct > 0:
            return STATE_QV_TO_BOT
        return STATE_QF_TO_V


# Instncia global com valores padro do documento
DEFAULT_RULES = ParaconsistentRules()


def state_from_rules(mu: float, lam: float, rules: ParaconsistentRules | None = None) -> str:
    """Retorna o estado lgico de 12 valores para (, ) segundo as regras do Fuzzy.txt."""
    r = rules or DEFAULT_RULES
    return r.state_12(mu, lam)


def state_12_to_simple(state_12: str) -> str:
    """
    Mapeia os 12 estados do reticulado para os 4 estados usados pelo
    TruthScoringModel / ParaconsistentValue: Verdadeiro | Falso | Intermedirio | Indeterminado.
    """
    if state_12 in (STATE_V, STATE_QV, STATE_T_TO_V, STATE_BOT_TO_V):
        return "Verdadeiro"
    if state_12 in (STATE_F, STATE_QF, STATE_BOT_TO_F, STATE_QF_TO_V):
        return "Falso"
    if state_12 in (STATE_T, STATE_QV_TO_T, STATE_QV_TO_BOT):
        return "Intermedirio"
    return "Indeterminado"


def truth_value_from_annotations(mu: float, lam: float) -> float:
    """Valor-verdade escalar em [0,1] a partir de (, ), compatvel com L3."""
    return round((mu + (1.0 - lam)) / 2.0, 4)


def parse_rules_from_fuzzy_text(content: str) -> ParaconsistentRules:
    """
    Extrai valores de controle do texto do arquivo Fuzzy.txt quando possvel.
    Se no encontrar, retorna DEFAULT_RULES.
    """
    rules = ParaconsistentRules()
    # Procura padres como "Vscc=1/2", "Vscct= 1/2", "1/2" prximo a Vscc, etc.
    vscc_m = re.search(r"Vscc\s*=\s*Vscct\s*=\s*1/2", content, re.I)
    if vscc_m:
        rules.vscc = 0.5
        rules.vscct = 0.5
    vicc_m = re.search(r"Vicc\s*(?:e|e\s*Vicct)?\s*[=:]?\s*-?\s*1/2", content, re.I)
    if vicc_m or "Vicc" in content and "-1/2" in content:
        rules.vicc = -0.5
        rules.vicct = -0.5
    return rules


def load_rules_from_fuzzy_file(path: str | None = None) -> ParaconsistentRules:
    """
    Carrega o contedo de data/Fuzzy.txt e retorna ParaconsistentRules.
    Se o arquivo no existir ou no for legvel, retorna DEFAULT_RULES.
    """
    if path is None:
        base = os.path.dirname(os.path.abspath(__file__))
        path = os.path.join(base, "data", "Fuzzy.txt")
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return parse_rules_from_fuzzy_text(content)
    except Exception:
        return DEFAULT_RULES


def generate_training_pairs(
    rules: ParaconsistentRules | None = None,
    grid_step: float = 0.1,
) -> Iterator[Tuple[float, float, str, float]]:
    """
    Gera pares (, , estado_12, valor_verdade) para treinar a camada L3
    a partir do conjunto de regras (para-analisador).
    til para criar dataset sinttico que segue exatamente o Fuzzy.txt.
    """
    r = rules or DEFAULT_RULES
    mu = 0.0
    while mu <= 1.0:
        lam = 0.0
        while lam <= 1.0:
            state = r.state_12(mu, lam)
            truth = truth_value_from_annotations(mu, lam)
            yield (mu, lam, state, truth)
            lam = round(lam + grid_step, 2)
        mu = round(mu + grid_step, 2)


def get_rules_training_examples(
    rules: ParaconsistentRules | None = None,
    grid_step: float = 0.1,
) -> List[Tuple[float, float, str, float]]:
    """Lista de (, , estado_12, valor_verdade) para uso no treinamento."""
    return list(generate_training_pairs(rules=rules, grid_step=grid_step))
# ========== custom_tokenizer.py ==========
"""
Tokenizador SentencePiece (BPE/Unigram) customizado
===================================================

Treina um modelo SentencePiece a partir de um corpus de texto (por exemplo,
o texto do artigo/README) e expe uma interface simples de encode/decode
para ser usada pelo modelo de linguagem customizado.
"""

from dataclasses import dataclass
from typing import List
import os

import sentencepiece as spm


SPECIAL_TOKENS = ["<pad>", "<bos>", "<eos>"]


@dataclass
class SPConfig:
    model_prefix: str = "sp_epistemologia"
    vocab_size: int = 2000
    model_type: str = "bpe"  # ou "unigram"


def train_sentencepiece(
    input_files: List[str],
    config: SPConfig = SPConfig(),
) -> None:
    """
    Treina um modelo SentencePiece a partir de uma lista de arquivos de texto.
    Gera `config.model_prefix.model` e `.vocab` na pasta atual.
    """
    input_str = ",".join(input_files)
    user_defined_symbols = ",".join(SPECIAL_TOKENS)

    spm.SentencePieceTrainer.Train(
        input=input_str,
        model_prefix=config.model_prefix,
        vocab_size=config.vocab_size,
        model_type=config.model_type,
        character_coverage=0.9995,
        bos_id=-1,
        eos_id=-1,
        pad_id=-1,
        user_defined_symbols=user_defined_symbols,
    )


class CustomSPTokenizer:
    """
    Wrapper simples em torno de SentencePieceProcessor.

    Convenes:
      - `<bos>`  adicionado no incio da sequncia.
      - `<eos>`  adicionado no final (opcional).
    """

    def __init__(self, model_prefix: str = "sp_epistemologia") -> None:
        model_file = f"{model_prefix}.model"
        if not os.path.exists(model_file):
            raise FileNotFoundError(
                f"Modelo SentencePiece '{model_file}' no encontrado. "
                f"Treine primeiro com train_sentencepiece()."
            )
        self.sp = spm.SentencePieceProcessor(model_file=model_file)
        # Mapeia ids das user_defined_symbols
        self.pad_id = self.sp.piece_to_id("<pad>")
        self.bos_id = self.sp.piece_to_id("<bos>")
        self.eos_id = self.sp.piece_to_id("<eos>")

    def encode(self, text: str, add_bos: bool = True, add_eos: bool = True) -> List[int]:
        pieces = self.sp.encode(text, out_type=int)
        ids: List[int] = []
        if add_bos and self.bos_id >= 0:
            ids.append(self.bos_id)
        ids.extend(pieces)
        if add_eos and self.eos_id >= 0:
            ids.append(self.eos_id)
        return ids

    def decode(self, ids: List[int]) -> str:
        # remove tokens especiais se presentes
        filtered = [
            i
            for i in ids
            if i not in {self.bos_id, self.eos_id, self.pad_id} and i >= 0
        ]
        return self.sp.decode(filtered)

    @property
    def vocab_size(self) -> int:
        return self.sp.vocab_size()
# ========== custom_lm_model.py ==========
"""
Modelo de Linguagem Customizado (TransformerEncoder + BPE)
==========================================================

Pequeno modelo de linguagem causal baseado em TransformerEncoder, usando
tokens produzidos pelo `CustomSPTokenizer` (SentencePiece).

Funes principais:
  - `EpistemicLanguageModel`: arquitetura PyTorch
  - `generate_text`: funo de gerao autoregressiva
  - helpers para salvar/carregar pesos
"""

from dataclasses import dataclass
from typing import Optional, List

import torch
import torch.nn as nn
import torch.nn.functional as F


@dataclass
class LMConfig:
    vocab_size: int
    d_model: int = 256
    n_heads: int = 4
    num_layers: int = 4
    dim_feedforward: int = 512
    max_seq_len: int = 256
    dropout: float = 0.1


class EpistemicLanguageModel(nn.Module):
    """
    Modelo de linguagem simples (causal) com TransformerEncoder.
    """

    def __init__(self, config: LMConfig) -> None:
        super().__init__()
        self.config = config

        self.token_emb = nn.Embedding(config.vocab_size, config.d_model)
        self.pos_emb = nn.Embedding(config.max_seq_len, config.d_model)

        encoder_layer = nn.TransformerEncoderLayer(
            d_model=config.d_model,
            nhead=config.n_heads,
            dim_feedforward=config.dim_feedforward,
            dropout=config.dropout,
            batch_first=True,
        )
        self.encoder = nn.TransformerEncoder(
            encoder_layer,
            num_layers=config.num_layers,
        )
        self.lm_head = nn.Linear(config.d_model, config.vocab_size, bias=False)

    def forward(
        self,
        input_ids: torch.Tensor,  # (batch, seq_len)
        attention_mask: Optional[torch.Tensor] = None,
    ) -> torch.Tensor:
        bsz, seq_len = input_ids.shape
        device = input_ids.device

        pos_ids = torch.arange(seq_len, device=device).unsqueeze(0).expand(bsz, -1)
        x = self.token_emb(input_ids) + self.pos_emb(pos_ids)

        # Mscara causal: cada posio s v tokens anteriores
        causal_mask = torch.triu(
            torch.ones(seq_len, seq_len, device=device, dtype=torch.bool),
            diagonal=1,
        )

        if attention_mask is not None:
            # attention_mask: (batch, seq_len) com 1 para tokens vlidos, 0 para pad
            # A API de TransformerEncoder usa src_key_padding_mask com True = pad
            key_padding_mask = attention_mask == 0
        else:
            key_padding_mask = None

        hidden = self.encoder(
            x,
            mask=causal_mask,
            src_key_padding_mask=key_padding_mask,
        )
        logits = self.lm_head(hidden)
        return logits


def generate_text(
    model: EpistemicLanguageModel,
    tokenizer,
    prompt: str,
    max_new_tokens: int = 50,
    temperature: float = 1.0,
    top_k: int = 50,
    device: Optional[torch.device] = None,
) -> str:
    """
    Gerao autoregressiva simples.
    """
    if device is None:
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    model.eval()
    model.to(device)

    ids = tokenizer.encode(prompt, add_bos=True, add_eos=False)
    input_ids = torch.tensor([ids], dtype=torch.long, device=device)

    for _ in range(max_new_tokens):
        if input_ids.size(1) >= model.config.max_seq_len:
            break

        with torch.no_grad():
            logits = model(input_ids)  # (1, seq_len, vocab)
            next_token_logits = logits[0, -1, :] / max(temperature, 1e-4)

            if top_k > 0:
                values, indices = torch.topk(next_token_logits, k=min(top_k, next_token_logits.size(-1)))
                probs = F.softmax(values, dim=-1)
                next_token = indices[torch.multinomial(probs, num_samples=1)]
            else:
                probs = F.softmax(next_token_logits, dim=-1)
                next_token = torch.multinomial(probs, num_samples=1)

        input_ids = torch.cat([input_ids, next_token.view(1, 1)], dim=1)

    generated_ids: List[int] = input_ids[0].tolist()
    return tokenizer.decode(generated_ids)


def save_lm(model: EpistemicLanguageModel, path: str) -> None:
    torch.save({"config": model.config.__dict__, "state_dict": model.state_dict()}, path)


def load_lm(path: str, vocab_size: int) -> EpistemicLanguageModel:
    data = torch.load(path, map_location="cpu")
    cfg_dict = data.get("config", {})
    cfg_dict["vocab_size"] = vocab_size  # garante compatibilidade
    config = LMConfig(**cfg_dict)
    model = EpistemicLanguageModel(config)
    model.load_state_dict(data["state_dict"])
    return model
# ========== neural_truth_model.py ==========
"""
MDULO NEURAL  TruthScoringModel
=================================

Modelo PyTorch baseado em Transformer (via `transformers`) que recebe
proposies textuais (sada de L2) e produz:

  - logits de classe para o estado paraconsistente:
        Verdadeiro | Falso | Intermedirio | Indeterminado
  - um escalar v  [0,1] representando o valor-verdade aproximado
    (compatvel com `ParaconsistentValue.truth_value`).

Este mdulo NO  acoplado diretamente ao pipeline; ele pode ser
instanciado e passado opcionalmente para a `ParaconsistentEngine`,
que o usar para calcular (, ) neurais em vez das heursticas puras.
"""

from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional

import torch
import torch.nn as nn
from torch.utils.data import Dataset
from transformers import AutoModel, AutoTokenizer


# 
# Rtulos paraconsistentes
# 

LABEL2ID: Dict[str, int] = {
    "Verdadeiro": 0,
    "Falso": 1,
    "Intermedirio": 2,
    "Indeterminado": 3,
}

ID2LABEL: Dict[int, str] = {v: k for k, v in LABEL2ID.items()}


@dataclass
class PropositionExample:
    """Exemplo supervisionado para treinamento do modelo neural."""

    text: str
    label_state: str          # uma das chaves de LABEL2ID
    truth_value: float        # valor-verdade escalar em [0,1]


class PropositionDataset(Dataset):
    """
    Dataset simples de proposies rotuladas com estado paraconsistente
    e valor-verdade escalar.
    """

    def __init__(self, examples: List[PropositionExample], tokenizer, max_length: int = 64):
        self.examples = examples
        self.tokenizer = tokenizer
        self.max_length = max_length

    def __len__(self) -> int:
        return len(self.examples)

    def __getitem__(self, idx: int) -> Dict[str, torch.Tensor]:
        ex = self.examples[idx]
        enc = self.tokenizer(
            ex.text,
            truncation=True,
            padding="max_length",
            max_length=self.max_length,
            return_tensors="pt",
        )
        item = {k: v.squeeze(0) for k, v in enc.items()}
        item["labels_state"] = torch.tensor(LABEL2ID[ex.label_state], dtype=torch.long)
        item["labels_truth"] = torch.tensor(float(ex.truth_value), dtype=torch.float)
        return item


class TruthScoringModel(nn.Module):
    """
    Modelo hbrido:
      - backbone TransformerEncoder (BERT-like)
      - cabea de classificao para o estado lgico
      - cabea de regresso para valor-verdade escalar
    """

    def __init__(
        self,
        backbone_name: str = "bert-base-multilingual-cased",
        num_labels: int = 4,
    ) -> None:
        super().__init__()
        self.backbone = AutoModel.from_pretrained(backbone_name)
        hidden_size = self.backbone.config.hidden_size

        self.classifier = nn.Linear(hidden_size, num_labels)
        self.truth_head = nn.Sequential(
            nn.Linear(hidden_size, hidden_size),
            nn.ReLU(),
            nn.Linear(hidden_size, 1),
            nn.Sigmoid(),  # restringe para [0,1]
        )

    def forward(
        self,
        input_ids: torch.Tensor,
        attention_mask: torch.Tensor,
        labels_state: Optional[torch.Tensor] = None,
        labels_truth: Optional[torch.Tensor] = None,
    ) -> Dict[str, torch.Tensor]:
        outputs = self.backbone(input_ids=input_ids, attention_mask=attention_mask)
        cls_emb = outputs.last_hidden_state[:, 0, :]

        logits_state = self.classifier(cls_emb)
        truth_score = self.truth_head(cls_emb).squeeze(-1)

        loss: Optional[torch.Tensor] = None
        if labels_state is not None and labels_truth is not None:
            ce_loss = nn.CrossEntropyLoss()(logits_state, labels_state)
            mse_loss = nn.MSELoss()(truth_score, labels_truth)
            loss = ce_loss + 0.5 * mse_loss

        return {
            "logits_state": logits_state,
            "truth_score": truth_score,
            "loss": loss,
        }


# 
# Helpers de inferncia
# 

def load_tokenizer(backbone_name: str = "bert-base-multilingual-cased"):
    """Cria um tokenizer compatvel com o backbone."""
    return AutoTokenizer.from_pretrained(backbone_name)


def score_proposition(
    model: TruthScoringModel,
    tokenizer,
    text: str,
    device: Optional[torch.device] = None,
) -> Tuple[str, float]:
    """
    Executa inferncia para uma nica proposio textual.
    Retorna (label_state, truth_score).
    """
    if device is None:
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.eval()
    enc = tokenizer(
        text,
        truncation=True,
        padding="max_length",
        max_length=64,
        return_tensors="pt",
    )
    enc = {k: v.to(device) for k, v in enc.items()}
    model = model.to(device)
    with torch.no_grad():
        out = model(**enc)
    logits = out["logits_state"]
    truth_score = float(out["truth_score"].cpu().item())
    pred_id = int(logits.argmax(dim=-1).cpu().item())
    pred_label = ID2LABEL[pred_id]
    return pred_label, truth_score


def neural_annotations(
    model: TruthScoringModel,
    tokenizer,
    text: str,
) -> Tuple[float, float, str, float]:
    """
    Mapeia a sada do modelo neural para (, ) compatveis com L3.
    Retorna (mu, lam, state_label, truth_score).
    """
    state, v = score_proposition(model, tokenizer, text)
    if state == "Verdadeiro":
        mu = v
        lam = 1.0 - v
    elif state == "Falso":
        mu = 1.0 - v
        lam = v
    elif state == "Intermedirio":
        mu = 0.4 + 0.2 * v
        lam = 0.4 + 0.2 * (1.0 - v)
    else:  # Indeterminado
        mu = 0.3
        lam = 0.3
    return float(mu), float(lam), state, float(v)
# ========== knowledge_base.py ==========
"""
Base de conhecimento escalvel.
===============================
Carrega KB a partir de arquivo (JSON) e opcionalmente enriquece com
retrieval em ChromaDB (RAG). Mantm interface termo -> grau [0,1] para L3/L4.
"""

import json
import os
import re
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional

# KB padro (fallback quando no h arquivo)
SEED_KNOWLEDGE_BASE: Dict[str, float] = {
    "quente": 0.85, "frio": 0.85, "morno": 0.70, "aquecido": 0.80, "gelado": 0.80,
    "temperatura": 0.90, "graus": 0.88, "escaldante": 0.75, "tpido": 0.65,
    "verdadeiro": 0.95, "falso": 0.95, "contradio": 0.80, "proposio": 0.85,
    "silogismo": 0.75, "conhecimento": 0.90, "inteligncia": 0.85, "conscincia": 0.70,
    "razo": 0.88, "verdade": 0.92, "gua": 0.95, "lquido": 0.90, "h2o": 0.90,
}


def _extract_texts_from_doc(doc: Any) -> List[str]:
    if isinstance(doc, str):
        return [doc]
    if isinstance(doc, dict):
        for key in ("text", "content", "body", "summary", "description"):
            value = doc.get(key)
            if isinstance(value, str) and value.strip():
                return [value]
        text_parts: List[str] = []
        for value in doc.values():
            if isinstance(value, str) and value.strip():
                text_parts.append(value)
        if text_parts:
            return [" ".join(text_parts)]
    if isinstance(doc, list):
        texts = []
        for item in doc:
            texts.extend(_extract_texts_from_doc(item))
        return texts
    return []


def _term_weights_from_texts(texts: List[str], max_terms: int = 2000) -> Dict[str, float]:
    counts: Counter[str] = Counter()
    for text in texts:
        words = re.findall(r"[a-z]+", text.lower())
        for word in words:
            if len(word) > 3:
                counts[word] += 1
    if not counts:
        return {}
    most_common = counts.most_common(max_terms)
    max_value = most_common[0][1]
    return {term: min(1.0, count / max_value) for term, count in most_common}


def _normalize_counter(counts: Counter[str]) -> Dict[str, float]:
    if not counts:
        return {}
    most_common = counts.most_common(2000)
    max_value = most_common[0][1]
    return {term: min(1.0, count / max_value) for term, count in most_common}


def _count_terms_in_text(text: str, counts: Counter[str]) -> None:
    for word in re.findall(r"[a-z]+", text.lower()):
        if len(word) > 3:
            counts[word] += 1


def _load_kb_from_jsonl(path: Path, max_docs: int = 20000) -> Dict[str, float]:
    counts: Counter[str] = Counter()
    with open(path, "r", encoding="utf-8") as f:
        for index, line in enumerate(f):
            if index >= max_docs:
                break
            line = line.strip()
            if not line:
                continue
            try:
                record = json.loads(line)
            except json.JSONDecodeError:
                continue
            for text in _extract_texts_from_doc(record):
                _count_terms_in_text(text, counts)
    return _normalize_counter(counts)


def load_kb_from_file(path: str | Path) -> Dict[str, float]:
    """
    Carrega dicionrio termo -> grau [0,1] de um arquivo JSON.
    Suporta JSON de termo->peso, lista de documentos e NDJSON.
    """
    path = Path(path)
    if not path.exists():
        return {}
    try:
        if path.suffix.lower() in {".jsonl", ".ndjson"}:
            return _load_kb_from_jsonl(path)

        if path.stat().st_size > 1_000_000_000:
            return _load_kb_from_jsonl(path)

        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        if isinstance(data, dict):
            if all(isinstance(v, (int, float)) for v in data.values()):
                return {k: float(v) for k, v in data.items()}
            return _term_weights_from_texts(_extract_texts_from_doc(data))

        if isinstance(data, list):
            texts: List[str] = []
            for item in data:
                texts.extend(_extract_texts_from_doc(item))
            return _term_weights_from_texts(texts)
    except Exception:
        try:
            return _load_kb_from_jsonl(path)
        except Exception:
            return {}
    return {}


def merge_kb(base: Dict[str, float], extra: Dict[str, float]) -> Dict[str, float]:
    """Mescla extra em base; em conflito, extra prevalece."""
    out = dict(base)
    for k, v in extra.items():
        out[k] = v
    return out


def enrich_kb_from_chroma(
    query: str,
    chroma_path: str,
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
    k: int = 5,
    score_weight: float = 0.8,
) -> Dict[str, float]:
    """
    Busca em ChromaDB por query e retorna um dicionrio termo -> peso
    extrado dos trechos recuperados (palavras relevantes com score_weight).
    """
    try:
        from langchain_community.vectorstores import Chroma
        from langchain_community.embeddings import HuggingFaceEmbeddings
    except ImportError:
        return {}

    chroma_path = Path(chroma_path)
    if not chroma_path.exists() or not chroma_path.is_dir():
        return {}

    try:
        embeddings = HuggingFaceEmbeddings(model_name=embedding_model)
        vectorstore = Chroma(persist_directory=str(chroma_path), embedding_function=embeddings)
        docs = vectorstore.similarity_search(query, k=k)
    except Exception:
        return {}

    # Extrai termos dos textos e atribui peso
    term_scores: Dict[str, float] = {}
    for d in docs:
        text = d.page_content if hasattr(d, "page_content") else str(d)
        words = re.findall(r"[a-z]+", text.lower())
        for w in words:
            if len(w) > 2:
                term_scores[w] = term_scores.get(w, 0) + score_weight
    # Normaliza para [0, 1]
    if term_scores:
        m = max(term_scores.values())
        term_scores = {t: min(1.0, s / m) for t, s in term_scores.items()}
    return term_scores


def get_domain_knowledge_base(
    domain: str,
    config: Optional[Dict[str, Any]] = None,
    query_for_rag: Optional[str] = None,
) -> Dict[str, float]:
    """
    Retorna KB especializado para um domnio.
    - Usa KB especfico de domnio se configurado em domain_specific_kbs.
    - Se no existirem arquivos de domnio, usa o KB genrico em knowledge_base.path.
    - Enriquece com ChromaDB especfico do domnio (chroma_path/{domain}/).
    - Fallback: SEED_KNOWLEDGE_BASE.
    """
    PROJECT_ROOT = Path(__file__).resolve().parent
    try:
        from config_loader import load_config, PROJECT_ROOT as _root
        PROJECT_ROOT = _root
        if config is None:
            config = load_config()
    except Exception:
        pass
    if config is None:
        config = {}

    base: Dict[str, float] = {}

    domain_specific = config.get("knowledge_base", {}).get("domain_specific_kbs", {})
    if domain and isinstance(domain_specific, dict) and domain_specific.get(domain):
        domain_path = Path(domain_specific[domain])
        if not domain_path.is_absolute():
            domain_path = PROJECT_ROOT / domain_path
        if domain_path.exists():
            base = load_kb_from_file(domain_path)

    if not base:
        kb_path = config.get("knowledge_base", {}).get("path", "")
        if kb_path:
            path_obj = Path(kb_path) if Path(kb_path).is_absolute() else PROJECT_ROOT / kb_path
            if path_obj.exists():
                if path_obj.is_dir() and domain:
                    candidate = path_obj / f"kb_{domain}.json"
                    if candidate.exists():
                        base = load_kb_from_file(candidate)
                elif path_obj.is_file():
                    base = load_kb_from_file(path_obj)

    if not base:
        default_kb = config.get("knowledge_base", {}).get("default_kb", "")
        if default_kb:
            default_path = Path(default_kb) if Path(default_kb).is_absolute() else PROJECT_ROOT / default_kb
            if default_path.exists():
                base = load_kb_from_file(default_path)

    if not base:
        base = dict(SEED_KNOWLEDGE_BASE)

    # Chroma especfico do domnio
    chroma_base = config.get("knowledge_base", {}).get("chroma_path") or config.get("agent", {}).get("vector_db_path", "")
    if chroma_base:
        domain_chroma_path = Path(chroma_base) / domain
        if domain_chroma_path.exists() and domain_chroma_path.is_dir() and query_for_rag:
            extra = enrich_kb_from_chroma(
                query_for_rag,
                str(domain_chroma_path),
                config.get("agent", {}).get("embedding_model", "sentence-transformers/all-MiniLM-L6-v2"),
                k=5,
            )
            base = merge_kb(base, extra)

    return base


def get_knowledge_base(
    config: Optional[Dict[str, Any]] = None,
    query_for_rag: Optional[str] = None,
    domain: Optional[str] = None,
) -> Dict[str, float]:
    """
    Retorna KB geral ou domnio-especfico a partir da configurao.

    Se knowledge_base.path for um arquivo JSON, usa-o como fonte genrica.
    """
    config = config or {}
    kb_config = config.get("knowledge_base", {})
    kb_path = kb_config.get("path", "")
    default_kb = kb_config.get("default_kb", "")
    domain_specific = kb_config.get("domain_specific_kbs", {})

    project_root = Path(__file__).resolve().parent

    if isinstance(domain_specific, dict) and domain:
        domain_file = domain_specific.get(domain)
        if domain_file:
            domain_path = Path(domain_file) if Path(domain_file).is_absolute() else project_root / domain_file
            if domain_path.exists():
                return load_kb_from_file(domain_path)

    if kb_path:
        base_path = Path(kb_path) if Path(kb_path).is_absolute() else project_root / kb_path
        if base_path.exists():
            if base_path.is_file():
                return load_kb_from_file(base_path)
            if base_path.is_dir() and domain:
                candidate = base_path / f"kb_{domain}.json"
                if candidate.exists():
                    return load_kb_from_file(candidate)

    if default_kb:
        default_path = Path(default_kb) if Path(default_kb).is_absolute() else project_root / default_kb
        if default_path.exists():
            return load_kb_from_file(default_path)

    return dict(SEED_KNOWLEDGE_BASE)
# ========== l4_russell_equivalence.py ==========
"""
Base terica russelliana para a camada L4  Equivalncia e Correspondncia
============================================================================
Utiliza o arquivo data/russell.txt (Bertrand Russell, The Problems of Philosophy)
para fundamentar a sntese L4 no conceito de EQUIVALNCIA como correspondncia
entre crena/proposio e fato, e no apenas em agregao estatstica.

Conceitos extrados do Cap. XII (Truth and Falsehood):
  - Verdade = correspondncia entre crena e fato.
  - Fato = unidade complexa formada pelos objetos da crena na mesma ordem.
  - Crena verdadeira quando existe fato correspondente; falsa quando no existe.
  - Propriedade extrnseca: a verdade depende da relao da crena com algo externo.
"""

import os
import re
from dataclasses import dataclass, field
from typing import Dict, List, Tuple, Optional


#  Conceitos russellianos (extrados do texto) 
# Termos que indicam alta relevncia para equivalncia/correspondncia
CORRESPONDENCE_TERMS = [
    "correspondence", "correspond", "corresponding", "corresponds",
    "equivalence", "equivalent", "match", "accord", "agree", "fact",
    "belief", "true", "truth", "false", "falsehood", "beliefs", "facts",
    "object-terms", "object-relation", "complex unity", "constituents",
    "judgement", "judging", "sense-data", "physical object",
]
# Normalizados para matching em portugus/ingls
EQUIVALENCE_CONCEPTS_PT = [
    "correspondncia", "equivalncia", "crena", "fato", "verdade", "falsidade",
    "juzo", "objeto", "termos", "relao", "unidade", "complexo",
    "dado sensvel", "proposio", "conhecimento",
]


@dataclass
class RussellConceptBase:
    """
    Base de conceitos extrada de russell.txt para fundamentar a sntese L4.
    Permite ponderar proposies por alinhamento terico (correspondncia com fatos)
    e no apenas por estatstica.
    """
    # Trechos do texto sobre verdade/correspondncia (cap. XII e adjacentes)
    key_passages: List[str] = field(default_factory=list)
    # Termos do texto com peso conceitual (relevncia para equivalncia)
    term_weights: Dict[str, float] = field(default_factory=dict)
    # Princpio em forma de texto (para auditoria/interpretao)
    principle_summary: str = ""

    def concept_weight_for_terms(self, terms: List[str]) -> float:
        """
        Peso conceitual para um conjunto de termos: quanto mais os termos
        aparecem na base russelliana, mais a proposio  tratada como
        alinhada  teoria da equivalncia (correspondncia crenafato).
        """
        if not self.term_weights:
            return 1.0
        total = 0.0
        count = 0
        for t in terms:
            t_lower = t.lower().strip()
            if t_lower in self.term_weights:
                total += self.term_weights[t_lower]
                count += 1
        if count == 0:
            return 1.0
        return 1.0 + (total / count) * 0.5  # modulao suave


def _normalize_word(w: str) -> str:
    return re.sub(r"[^a-z0-9]", "", w.lower())


def load_russell_text(path: Optional[str] = None) -> str:
    """Carrega o contedo de data/russell.txt."""
    if path is None:
        base = os.path.dirname(os.path.abspath(__file__))
        path = os.path.join(base, "data", "russell.txt")
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_chapter_xii(content: str) -> str:
    """Extrai o captulo XII (Truth and Falsehood) e trechos adjacentes relevantes."""
    start = content.find("CHAPTER XII")
    if start == -1:
        start = content.find("TRUTH AND FALSEHOOD")
    if start == -1:
        return content[:15000]  # fallback: incio do livro
    end = content.find("CHAPTER XIV", start)
    if end == -1:
        end = content.find("CHAPTER XV", start)
    if end == -1:
        end = len(content)
    return content[start:end]


def extract_equivalence_passages(content: str) -> List[str]:
    """Extrai trechos que definem equivalncia/correspondncia."""
    chapter_xii = extract_chapter_xii(content)
    # Frases que contm os conceitos centrais
    sentences = re.split(r"[.!?]\s+", chapter_xii)
    key = []
    for s in sentences:
        s_lower = s.lower()
        if any(
            x in s_lower
            for x in (
                "correspondence",
                "correspond",
                "belief",
                "fact",
                "true",
                "false",
                "complex unity",
                "object-terms",
                "object-relation",
            )
        ):
            key.append(s.strip())
    return key[:50]  # limite razovel


def build_term_weights_from_russell(content: str) -> Dict[str, float]:
    """
    Constri pesos por termo a partir do texto de Russell: termos que aparecem
    em contextos de verdade/correspondncia recebem peso maior.
    """
    chapter = extract_chapter_xii(content)
    words = re.findall(r"[a-z]+", chapter.lower())
    # Frequncia no captulo de verdade
    freq: Dict[str, int] = {}
    for w in words:
        w = _normalize_word(w)
        if len(w) > 2:
            freq[w] = freq.get(w, 0) + 1
    # Normalizar para [0.2, 1.0] por relevncia conceitual
    concept_set = set(
        _normalize_word(t) for t in CORRESPONDENCE_TERMS + EQUIVALENCE_CONCEPTS_PT
    )
    max_f = max(freq.values()) if freq else 1
    term_weights: Dict[str, float] = {}
    for w, c in freq.items():
        if w in concept_set:
            term_weights[w] = 0.5 + 0.5 * (c / max_f)
        else:
            term_weights[w] = 0.2 + 0.3 * (c / max_f)
    return term_weights


def build_russell_concept_base(path: Optional[str] = None) -> RussellConceptBase:
    """
    Treina/constroi a base de conceitos russellianos a partir de russell.txt.
    Usado pela L4 para sntese fundamentada em equivalncia (correspondncia).
    """
    content = load_russell_text(path)
    passages = extract_equivalence_passages(content)
    term_weights = build_term_weights_from_russell(content)
    summary = (
        "Truth consists in correspondence between belief and fact. "
        "A belief is true when there is a corresponding fact (complex unity of the objects of the belief). "
        "Truth and falsehood are extrinsic properties: they depend on the relation of the belief to outside things."
    )
    return RussellConceptBase(
        key_passages=passages,
        term_weights=term_weights,
        principle_summary=summary,
    )


def score_proposition_by_concepts(
    proposition: str,
    knowledge_base: Dict[str, float],
    concept_base: RussellConceptBase,
) -> float:
    """
    Score conceitual da proposio: grau em que ela se alinha  teoria da
    equivalncia (correspondncia com fatos/BD), no apenas estatstica.

    - Termos da proposio que esto no KB com alta evidncia indicam
      melhor "correspondncia" com o mundo (fatos).
    - Termos que aparecem na base russelliana aumentam o peso terico.
    """
    words = re.findall(r"[a-z]+", proposition.lower())
    terms = [_normalize_word(w) for w in words if len(w) > 2]

    # 1) Alinhamento com fatos (KB): termos da proposio presentes no BD
    kb_match = 0.0
    n = 0
    for t in terms:
        for kb_term, ev in knowledge_base.items():
            if _normalize_word(kb_term) == t or t in _normalize_word(kb_term):
                kb_match += ev
                n += 1
                break
    fact_alignment = (kb_match / n) if n > 0 else 0.5  # neutro se nenhum termo no KB

    # 2) Peso conceitual russelliano (termos da teoria)
    concept_weight = concept_base.concept_weight_for_terms(terms)

    # Combinao: correspondncia com fatos (BD) + alinhamento terico
    return (0.7 * fact_alignment + 0.3 * concept_weight)


def save_concept_base(base: RussellConceptBase, path: str) -> None:
    """Salva a base de conceitos para uso posterior da L4."""
    import json
    data = {
        "principle_summary": base.principle_summary,
        "key_passages": base.key_passages[:20],
        "term_weights": base.term_weights,
    }
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_concept_base(path: str) -> RussellConceptBase:
    """Carrega base de conceitos previamente construda."""
    import json
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return RussellConceptBase(
        principle_summary=data.get("principle_summary", ""),
        key_passages=data.get("key_passages", []),
        term_weights=data.get("term_weights", {}),
    )
# ========== l4_chain_verification.py ==========
"""
Chain of Verification (CoVe) agent para a camada L4.
====================================================
Implementa o workflow Factor + Revise como etapa adicional de verificao
da sntese L4 antes da resposta final ser entregue.
"""

import os
import re
from typing import Any, Dict, List, Optional, Tuple


def generate_with_ollama(context: str, model: str = "doninha8:latest", ollama_host: str = "http://localhost:11434", temperature: float = 0.2) -> str:
    """Gera resposta usando Ollama com modelo base."""
    try:
        import ollama
        
        if ollama_host:
            os.environ["OLLAMA_HOST"] = ollama_host
        
        response = ollama.chat(
            model=model,
            messages=[{"role": "user", "content": context}],
            stream=False,
            options={
                "temperature": temperature,
                "num_ctx": 8192,
            },
        )
        if isinstance(response, dict):
            return response.get("message", {}).get("content", "").strip()
        return str(response).strip() if response else ""
    except Exception:
        return ""


def generate_with_custom_lm(context: str, model_path: str, max_new_tokens: int = 150, temperature: float = 0.7) -> str:
    try:
        from custom_lm_model import EpistemicLanguageModel, LMConfig, generate_text, load_lm
        from custom_tokenizer import CustomSPTokenizer, SPConfig
        import torch
        tokenizer = CustomSPTokenizer(SPConfig())
        tokenizer.load()
        vocab_size = tokenizer.vocab_size()
        model = load_lm(model_path, vocab_size)
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        out = generate_text(model, tokenizer, context, max_new_tokens=max_new_tokens, temperature=temperature, device=device)
        return out or ""
    except Exception:
        return ""


class ChainOfVerificationAgent:
    """Agente de Chain of Verification para a camada L4."""

    def __init__(self, config: Optional[Dict[str, Any]] = None) -> None:
        self.config = config or {}
        self.provider = self.config.get("provider", "template")
        self.ollama_model = self.config.get("ollama_model", "doninha8:latest")
        self.ollama_host = self.config.get("ollama_host", "http://localhost:11434")
        self.custom_lm_path = self.config.get("custom_lm_path", "")

    def verify(
        self,
        prompt: str,
        baseline_response: str,
        context_summary: str,
    ) -> Tuple[str, List[str]]:
        """Executa o workflow CoVe e retorna resposta revisada + log."""
        if self.provider == "ollama":
            output = self._verify_with_ollama(prompt, baseline_response, context_summary)
        elif self.provider == "custom_lm" and self.custom_lm_path:
            output = self._verify_with_custom_lm(prompt, baseline_response, context_summary)
        else:
            output = self._template_verify(prompt, baseline_response, context_summary)

        revised, log = self._parse_verification_output(output, baseline_response)
        return revised, log

    def _verify_with_ollama(self, prompt: str, baseline_response: str, context_summary: str) -> str:
        verification_prompt = self._build_agent_prompt(prompt, baseline_response, context_summary)
        return generate_with_ollama(verification_prompt, model=self.ollama_model, ollama_host=self.ollama_host)

    def _verify_with_custom_lm(self, prompt: str, baseline_response: str, context_summary: str) -> str:
        verification_prompt = self._build_agent_prompt(prompt, baseline_response, context_summary)
        return generate_with_custom_lm(verification_prompt, self.custom_lm_path)

    def _template_verify(self, prompt: str, baseline_response: str, context_summary: str) -> str:
        claims = self._extract_claims(baseline_response)
        questions = self._build_verification_questions(claims)
        verifications = [f"{idx+1}. {q}  Incerto; verificao externa necessria." for idx, q in enumerate(questions)]
        revised = baseline_response.strip()
        if verifications:
            revised += "\n\nNota: esta resposta foi revisada com base em verificao interna limitada; algumas afirmaes permanecem pendentes de confirmao externa."
        sections = [
            "Baseline Response:",
            baseline_response.strip(),
            "",
            "Verification Questions:",
            *questions,
            "",
            "Independent Verification Results:",
            *verifications,
            "",
            "Cross-Check & Revise:",
            "Nenhuma inconsistncia formal identificada no contedo disponvel localmente.",
            "",
            "Revised Response:",
            revised,
        ]
        return "\n".join(sections)

    def _build_agent_prompt(self, prompt: str, baseline_response: str, context_summary: str) -> str:
        lines = [
            "Voc  um engenheiro de prompts especialista em tcnicas avanadas de confiabilidade.",
            "A partir de agora, use o mtodo Chain of Verification (CoVe) - variante Factor + Revise para analisar e revisar a resposta.",
            "Responda usando sempre o fluxo: 1. Baseline Response, 2. Factoring, 3. Independent Verification, 4. Cross-Check & Revise.",
            "Seja rigoroso, conservador, e declare limitaes quando necessrio.",
            "",
            f"Pergunta original: {prompt}",
            "",
            "Contexto resumido de L4: ",
            context_summary or "Sem contexto adicional disponvel.",
            "",
            "Resposta inicial (Baseline Response):",
            baseline_response.strip(),
            "",
            "Tarefa:",
            "1. Gere de 6 a 12 perguntas de verificao independentes a partir das principais afirmaes da resposta inicial.",
            "2. Responda cada pergunta de forma independente, marcando como Confirmado, Refutado, Parcialmente correto ou Incerto.",
            "3. Compare a resposta inicial com os resultados e reescreva a resposta final incorporando apenas o que foi verificado.",
            "4. Entregue a estrutura completa com as sees claramente demarcadas e finalize com a resposta revisada.",
            "",
            "Formato de sada exigido:",
            "Baseline Response:",
            "<texto>",
            "",
            "Verification Questions:",
            "1. <pergunta>",
            "...",
            "",
            "Independent Verification Results:",
            "1. <marcao>  <resposta>",
            "...",
            "",
            "Cross-Check & Revise:",
            "<anlise>",
            "",
            "Revised Response:",
            "<texto revisado>",
        ]
        return "\n".join(lines)

    def _parse_verification_output(self, output: str, baseline_response: str) -> Tuple[str, List[str]]:
        if not output:
            return baseline_response, ["Nenhuma sada de verificao gerada."]

        revised = baseline_response
        log: List[str] = []
        if "Revised Response:" in output:
            parts = output.split("Revised Response:")
            revised = parts[-1].strip()
            log = [line.strip() for line in output.splitlines() if line.strip()]
        else:
            log = [line.strip() for line in output.splitlines() if line.strip()]
        return revised, log

    def _extract_claims(self, text: str) -> List[str]:
        sentences = [s.strip() for s in re.split(r"(?<=[.!?])\\s+", text) if s.strip()]
        claims = []
        for sentence in sentences:
            if len(claims) >= 12:
                break
            if len(sentence.split()) >= 5:
                claims.append(sentence)
        return claims[:12] if claims else sentences[:min(6, len(sentences))]

    def _build_verification_questions(self, claims: List[str]) -> List[str]:
        questions: List[str] = []
        for claim in claims[:12]:
            question = f"A afirmao a seguir est correta e fundamentada? {claim}"
            questions.append(question)
        if len(questions) < 6:
            questions.extend([
                "A estrutura lgica da resposta est consistente com a informao disponvel?",
                "H alguma suposio implcita que precisa ser explicitada ou verificada?",
            ])
        return questions[:12]
# ========== corpus_utils.py ==========
"""
Utilitrios de corpus
=====================

Funes para carregar texto de:
  - arquivos Markdown/TXT (ex.: README)
  - artigo completo em DOCX ("Uma verdadeira Epistemologia para a Inteligncia Artificial")
"""

from typing import List
import os

from docx import Document


def read_text_file(path: str, encoding: str = "utf-8") -> str:
    if not os.path.exists(path):
        raise FileNotFoundError(f"Arquivo de texto no encontrado: {path}")
    with open(path, "r", encoding=encoding) as f:
        return f.read()


def read_docx_file(path: str) -> str:
    if not os.path.exists(path):
        raise FileNotFoundError(f"Arquivo DOCX no encontrado: {path}")
    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


def load_main_corpus() -> List[str]:
    """
    Carrega o corpus principal deste projeto, incluindo materiais do projeto e bases de dados suplementares.

    Retorna uma lista de textos (documentos).
    """
    base_dir = os.path.dirname(__file__) or "."
    readme_path = os.path.join(base_dir, "README.md")
    article_path = os.path.join(
        base_dir,
        "Uma verdadeira Epistemologia para a Inteligncia Artificial.docx",
    )

    extra_paths = [
        os.path.join(base_dir, "data", "stanford_encyclopedia", "sep_texts_only.txt"),
        os.path.join(base_dir, "philosophy-corpus", "train_philosophy.txt"),
        os.path.join(base_dir, "philosophy-corpus", "train.txt"),
    ]

    texts: List[str] = []
    if os.path.exists(readme_path):
        texts.append(read_text_file(readme_path))
    if os.path.exists(article_path):
        texts.append(read_docx_file(article_path))

    for path in extra_paths:
        if os.path.exists(path):
            texts.append(read_text_file(path))

    if not texts:
        raise FileNotFoundError(
            "Nenhum corpus encontrado. Certifique-se de que README.md, o artigo DOCX ou os arquivos da base de dados esto disponveis."
        )
    return texts
# ========== l1_concept_table.py ==========
"""
CAMADA L1  Tbua de Conceitos (Aristteles: Categorias)
=========================================================
Mapeia cada termo do prompt a relaes semnticas fixas:
  - Sinonmia   : mesma denotao
  - Antonmia   : oposio semntica direta
  - Hiponmia   : relao especfico  geral
  - Homonmia   : mesma forma, sentidos distintos
  - Paronmia   : semelhana formal, sentidos distintos

As relaes so BINRIAS nesta camada  elimina a necessidade de
defuzzificao posterior na camada L3.
"""

from dataclasses import dataclass, field
from typing import Dict, List, Optional
import re
import json
import os
from knowledge_base import get_domain_knowledge_base


@dataclass
class ConceptNode:
    """Um conceito na tbua, com todas as suas relaes."""
    term: str
    definition: str = ""
    synonyms:   List[str] = field(default_factory=list)
    antonyms:   List[str] = field(default_factory=list)
    hyponyms:   List[str] = field(default_factory=list)   # mais especficos
    hypernyms:  List[str] = field(default_factory=list)   # mais gerais
    homonyms:   Dict[str, str] = field(default_factory=dict)  # sentido  definio
    paronyms:   List[str] = field(default_factory=list)
    domain:     str = "geral"
    application_context: str = ""
    canonical_source: str = ""
    canonical_context: Dict[str, str] = field(default_factory=dict)  # Verificao de atribuio cannica


class ConceptTable:
    """
    Tbua de conceitos fixos.  Em produo seria alimentada por um
    dicionrio / ontologia formal (WordNet-PT, OpenWordNet-PT, etc.).
    Aqui usamos um conjunto seminal suficiente para demonstrar todas
    as camadas do modelo.
    """

    def __init__(self) -> None:
        self._table: Dict[str, ConceptNode] = {}
        # Tbua seminal em portugus
        self._build_seed_table()
        # Banco de conceitos em ingls aprendido de dicionrio externo (se existir)
        self._load_external_concepts()

    # ------------------------------------------------------------------ #
    # API pblica                                                          #
    # ------------------------------------------------------------------ #

    def get(self, term: str) -> Optional[ConceptNode]:
        return self._table.get(self._normalize(term))

    def extract_concepts(self, text: str, llm_context: Optional[str] = None, domain: str = "geral", config: Optional[Dict] = None) -> List[ConceptNode]:
        """Extrai e retorna os ns de todos os termos encontrados no texto."""
        tokens = re.findall(r"[a-zA-Z]+", text)
        seen, result = set(), []
        for tok in tokens:
            key = self._normalize(tok)
            if key not in seen:
                node = self._table.get(key)
                if node:
                    seen.add(key)
                    result.append(self._clone_node(node))

        if result:
            self._enrich_concepts_with_application_context(result, text, llm_context, domain, config)
            combined_text = f"{llm_context.strip()} {text}" if llm_context else text
            result = [
                node for node in result
                if LogicLMSymbolicSolver.is_context_compatible(node, combined_text)
            ]
        return result

    def add(self, node: ConceptNode) -> None:
        self._table[self._normalize(node.term)] = node

    def relation_type(self, term_a: str, term_b: str) -> str:
        """Retorna o tipo de relao semntica entre dois termos."""
        a = self._normalize(term_a)
        b = self._normalize(term_b)
        node_a = self._table.get(a)
        if not node_a:
            return "desconhecida"
        if b in [self._normalize(s) for s in node_a.synonyms]:
            return "sinonmia"
        if b in [self._normalize(s) for s in node_a.antonyms]:
            return "antonmia"
        if b in [self._normalize(s) for s in node_a.hyponyms]:
            return "hiponmia"
        if b in [self._normalize(s) for s in node_a.hypernyms]:
            return "hiperonmia"
        if b in [self._normalize(s) for s in node_a.paronyms]:
            return "paronmia"
        if b in [self._normalize(k) for k in node_a.homonyms]:
            return "homonmia"
        return "sem_relao_direta"

    # ------------------------------------------------------------------ #
    # Construo da tbua seminal                                          #
    # ------------------------------------------------------------------ #

    def _clone_node(self, node: ConceptNode) -> ConceptNode:
        return ConceptNode(
            term=node.term,
            definition=node.definition,
            synonyms=list(node.synonyms),
            antonyms=list(node.antonyms),
            hyponyms=list(node.hyponyms),
            hypernyms=list(node.hypernyms),
            homonyms=dict(node.homonyms),
            paronyms=list(node.paronyms),
            domain=node.domain,
            application_context="",
            canonical_source=node.canonical_source,
            canonical_context=dict(node.canonical_context),
        )

    def _enrich_concepts_with_application_context(
        self,
        concepts: List[ConceptNode],
        prompt: str,
        llm_context: Optional[str] = None,
        domain: str = "geral",
        config: Optional[Dict] = None,
    ) -> None:
        """Aplica o solver simblico Logic-LM para adicionar contexto de uso aos conceitos."""
        LogicLMSymbolicSolver.enrich(concepts, prompt, llm_context, domain, config)

    def _build_seed_table(self) -> None:
        entries = [
            ConceptNode(
                term="quente",
                definition="Que possui temperatura alta.",
                synonyms=["aquecido", "clido", "morno", "tpido"],
                antonyms=["frio", "gelado", "fresco"],
                hypernyms=["temperatura"],
                hyponyms=["escaldante", "ardente"],
                domain="fsico",
                canonical_source="Newton - Philosophiae Naturalis Principia Mathematica - Livro I",
            ),
            ConceptNode(
                term="frio",
                definition="Que possui temperatura baixa.",
                synonyms=["gelado", "fresco", "frgido"],
                antonyms=["quente", "aquecido", "clido"],
                hypernyms=["temperatura"],
                hyponyms=["congelado", "glacial"],
                domain="fsico",
                canonical_source="Newton - Philosophiae Naturalis Principia Mathematica - Livro I",
            ),
            ConceptNode(
                term="morno",
                definition="Entre quente e frio; tpido.",
                synonyms=["tpido", "ameno"],
                antonyms=["escaldante", "glacial"],
                hypernyms=["temperatura", "quente", "frio"],
                hyponyms=[],
                domain="fsico",
                canonical_source="Galen - De Temperamentis - Seo 3",
            ),
            ConceptNode(
                term="temperatura",
                definition="Grandeza fsica que mede o grau de calor de um corpo.",
                synonyms=["calor", "grau"],
                antonyms=[],
                hypernyms=["grandeza_fsica"],
                hyponyms=["quente", "frio", "morno"],
                domain="fsico",
                canonical_source="Galileu - Discorsi e Dimostrazioni Matematiche - Seo 2",
            ),
            ConceptNode(
                term="gua",
                definition="Substncia H2O, geralmente em estado lquido.",
                synonyms=["H2O", "lquido"],
                antonyms=[],
                hypernyms=["substncia", "fluido"],
                hyponyms=["vapor", "gelo"],
                domain="fsico",
                canonical_source="Newton - Opticks - Definio 19",
            ),
            ConceptNode(
                term="verdadeiro",
                definition="Que est de acordo com os fatos ou a realidade.",
                synonyms=["correto", "real", "factual"],
                antonyms=["falso", "incorreto", "fictcio"],
                hypernyms=["valor_lgico"],
                domain="lgica",
                canonical_source="Aristteles - Metafsica - Livro Gamma",
                canonical_context={
                    "lgica_clssica": "Aristteles - Metafsica: valor de verdade binrio, NO lgica paraconsistente",
                    "epistemologia": "Plato - Teeteto: correspondncia com realidade, NO coerncia pura"
                }
            ),
            ConceptNode(
                term="falso",
                definition="Que no corresponde aos fatos ou  realidade.",
                synonyms=["incorreto", "errado", "fictcio"],
                antonyms=["verdadeiro", "correto", "real"],
                hypernyms=["valor_lgico"],
                domain="lgica",
                canonical_source="Aristteles - Metafsica - Livro Gamma",
                canonical_context={
                    "lgica_clssica": "Aristteles - Metafsica: negao do verdadeiro, NO dialtica hegeliana"
                }
            ),
            ConceptNode(
                term="banco",
                definition="Mvel para sentar; instituio financeira; repositrio de dados.",
                synonyms=[],
                antonyms=[],
                hypernyms=[],
                homonyms={
                    "assento": "mvel para sentar",
                    "financeiro": "instituio financeira",
                    "dados": "repositrio de dados",
                },
                domain="geral",
            ),
            ConceptNode(
                term="eminente",
                definition="Pessoa ilustre ou notvel.",
                synonyms=["ilustre", "notvel"],
                antonyms=[],
                paronyms=["iminente"],
                domain="geral",
            ),
            ConceptNode(
                term="iminente",
                definition="Que est prestes a acontecer.",
                synonyms=["prximo", "imediato"],
                antonyms=[],
                paronyms=["eminente"],
                domain="geral",
            ),
            ConceptNode(
                term="inteligncia",
                definition="Capacidade de compreender, raciocinar e resolver problemas.",
                synonyms=["cognio", "raciocnio", "entendimento"],
                antonyms=["ignorncia", "estupidez"],
                hypernyms=["capacidade_mental"],
                domain="cognitivo",
            ),
            ConceptNode(
                term="conhecimento",
                definition="Ato ou efeito de conhecer; saber, cincia, erudio.",
                synonyms=["saber", "cincia", "erudio"],
                antonyms=["ignorncia", "desconhecimento"],
                hypernyms=["epistemologia"],
                domain="filosfico",
                canonical_context={
                    "epistemologia": "Plato - Teeteto: justificao verdadeira, NO opinio infundada",
                    "kantiano": "Kant - Crtica da Razo Pura: a priori vs a posteriori, NO empirismo puro"
                }
            ),
            ConceptNode(
                term="verdade",
                definition="Conformidade entre o que se diz e o que .",
                synonyms=["veracidade", "factualidade", "realidade"],
                antonyms=["mentira", "falsidade", "iluso"],
                hypernyms=["epistemologia"],
                domain="filosfico",
                canonical_context={
                    "platnico": "Plato - Repblica: ideias eternas, NO relativismo",
                    "aristotlico": "Aristteles - Metafsica: correspondncia, NO coerncia"
                }
            ),
            ConceptNode(
                term="sntese regulativa",
                definition="Princpio que orienta o conhecimento sem constitu-lo.",
                synonyms=["regulativo", "orientador"],
                antonyms=[],
                hypernyms=["epistemologia", "kantismo"],
                domain="filosfico",
                canonical_source="Kant - Crtica da Razo Pura",
                canonical_context={
                    "kantismo": "Kant, CRP: princpio regulativo do conhecimento, NO Russell"
                }
            ),
        ]
        for node in entries:
            self.add(node)

    @staticmethod
    def _normalize(term: str) -> str:
        return term.strip().lower()

    # ------------------------------------------------------------------ #
    # Carregamento de conceitos externos (ex.: dicionrio em ingls)      #
    # ------------------------------------------------------------------ #

    def _load_external_concepts(self) -> None:
        """
        Carrega conceitos adicionais de um banco gerado a partir do
        dicionrio em ingls (arquivo JSON se existir).

        Formato esperado (lista de objetos):
          {
            "term": "abacus",
            "definition": "Frame with beads for calculating...",
            "synonyms": [],
            "antonyms": [],
            "hyponyms": [],
            "hypernyms": [],
            "domain": "geral"
          }
        """
        base_dir = os.path.dirname(__file__) or "."
        json_path = os.path.join(base_dir, "data", "concepts_en.json")
        if not os.path.exists(json_path):
            return
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                items = json.load(f)
        except Exception:
            return

        for item in items:
            term = item.get("term")
            if not term:
                continue
            node = ConceptNode(
                term=term,
                definition=item.get("definition", ""),
                synonyms=item.get("synonyms", []),
                antonyms=item.get("antonyms", []),
                hyponyms=item.get("hyponyms", []),
                hypernyms=item.get("hypernyms", []),
                homonyms=item.get("homonyms", {}),
                paronyms=item.get("paronyms", []),
                domain=item.get("domain", "geral"),
                application_context="",
                canonical_source=item.get("canonical_source", ""),
            )
            key = self._normalize(term)
            if key not in self._table:
                self._table[key] = node


class LogicLMSymbolicSolver:
    """Processador simblico inspirado em LLM-Symbolic Solver Logic-LM.

    Este mdulo faz uma pesquisa contextual nos parmetros de entrada da
    LLM base da IA Doninha e acrescenta  definio dos conceitos uma nota
    de aplicao prtica para o prompt atual.
    """

    CONTEXTUAL_KEYWORDS = {
        "fsico": ["temperatura", "calor", "energia", "massa", "volume"],
        "lgica": ["verdade", "falso", "proposio", "argumento", "inferencia", "inferncia"],
        "cognitivo": ["raciocnio", "inteligncia", "compreender", "resolver", "pensar"],
        "filosfico": ["verdade", "conhecimento", "epistemologia", "realidade", "tica"],
        "geral": ["aplicao", "uso", "contexto", "pergunta", "problema"],
    }

    @classmethod
    def enrich(
        cls,
        concepts: List[ConceptNode],
        prompt: str,
        llm_context: Optional[str] = None,
        domain: str = "geral",
        config: Optional[Dict] = None,
    ) -> List[ConceptNode]:
        text = prompt.strip()
        if llm_context:
            text = f"{llm_context.strip()} {text}"
        lower_text = text.lower()

        # Carrega KB especfico do domnio
        kb = get_domain_knowledge_base(domain, config, query_for_rag=text)

        for node in concepts:
            node.application_context = cls._infer_application_context(node, lower_text, concepts, kb)
        return concepts

    @classmethod
    def _infer_application_context(
        cls,
        node: ConceptNode,
        text: str,
        concepts: List[ConceptNode],
        kb: Dict[str, float],
    ) -> str:
        base_context = ""
        if node.term.lower() in text:
            if not cls.is_context_compatible(node, text):
                return ""
            relation = cls._infer_relation(node, text, concepts)
            if relation:
                base_context = relation
        else:
            base_context = cls._default_context(node)

        # Enriquece com termos relevantes do KB do domnio
        relevant_terms = [term for term, score in kb.items() if term.lower() in text.lower() and score > 0.5]
        if relevant_terms:
            kb_context = f" Contexto de conhecimento: {', '.join(relevant_terms[:3])}."
            base_context += kb_context

        return base_context

    @classmethod
    def is_context_compatible(
        cls,
        node: ConceptNode,
        text: str,
    ) -> bool:
        if not node.canonical_source:
            return True
        lower_text = text.lower()
        if node.term.lower() in lower_text:
            return True
        if node.domain:
            domain_keywords = cls.CONTEXTUAL_KEYWORDS.get(node.domain, [])
            if any(keyword in lower_text for keyword in domain_keywords):
                return True
        canonical_keywords = cls._extract_source_keywords(node.canonical_source)
        if any(keyword in lower_text for keyword in canonical_keywords):
            return True
        if node.application_context and any(
            part in lower_text for part in cls._tokenize(node.application_context)
        ):
            return True

        # Verificao de atribuio cannica
        if node.canonical_context:
            return cls._check_canonical_context_compatibility(node, text)
        return False

    @classmethod
    def _check_canonical_context_compatibility(
        cls,
        node: ConceptNode,
        text: str,
    ) -> bool:
        """Verifica se o contexto cannico do conceito  compatvel com o texto atual."""
        lower_text = text.lower()
        for context_key, context_value in node.canonical_context.items():
            # Verifica se o contexto cannico contm indicaes de incompatibilidade
            if "NO" in context_value.upper():
                # Extrai termos proibidos (aps "NO")
                not_parts = context_value.upper().split("NO")[1:]
                for not_part in not_parts:
                    prohibited_terms = cls._extract_prohibited_terms(not_part.strip())
                    if any(term in lower_text for term in prohibited_terms):
                        # Incompatvel - gera alerta para L7
                        cls._generate_canonical_alert(node, context_key, context_value, text)
                        return False
            # Verifica se o contexto cannico requer termos especficos
            elif ":" in context_value:
                required_terms = cls._extract_required_terms(context_value)
                if any(term in lower_text for term in required_terms):
                    return True
        return True  # Compatvel por padro se no h restries especficas

    @classmethod
    def _extract_prohibited_terms(cls, not_part: str) -> List[str]:
        """Extrai termos proibidos de uma parte 'NO ...'."""
        # Remove pontuao e divide por vrgulas ou 'ou'
        terms = re.split(r'[,\s]+ou[\s]+|[,;]', not_part)
        return [term.strip().lower() for term in terms if term.strip()]

    @classmethod
    def _extract_required_terms(cls, context_value: str) -> List[str]:
        """Extrai termos requeridos do contexto cannico."""
        # Assume formato "Fonte: descrio, termos requeridos"
        parts = context_value.split(":")
        if len(parts) > 1:
            description = parts[1].strip()
            terms = re.findall(r"[a-zA-Z]+", description)
            return [term.lower() for term in terms if len(term) > 3]
        return []

    @classmethod
    def _generate_canonical_alert(
        cls,
        node: ConceptNode,
        context_key: str,
        context_value: str,
        text: str,
    ) -> None:
        """Gera um alerta de incompatibilidade cannica para ser passado ao L7."""
        # Armazena o alerta em uma varivel global ou estrutura compartilhada
        # Por simplicidade, vamos usar um dicionrio global para alertas
        if not hasattr(cls, '_canonical_alerts'):
            cls._canonical_alerts = []
        alert = {
            'concept': node.term,
            'canonical_context': f"{context_key}: {context_value}",
            'incompatible_usage': text[:100] + "..." if len(text) > 100 else text,
            'alert_type': 'canonical_incompatibility'
        }
        cls._canonical_alerts.append(alert)

    @classmethod
    def get_canonical_alerts(cls) -> List[Dict]:
        """Retorna e limpa os alertas cannicos gerados."""
        if not hasattr(cls, '_canonical_alerts'):
            cls._canonical_alerts = []
        alerts = cls._canonical_alerts[:]
        cls._canonical_alerts.clear()
        return alerts

    @classmethod
    def _extract_source_keywords(cls, source: str) -> List[str]:
        return [
            token for token in re.findall(r"[a-zA-Z]+", source.lower())
            if len(token) > 3
        ]

    @classmethod
    def _tokenize(cls, text: str) -> List[str]:
        return [token for token in re.findall(r"[a-zA-Z]+", text.lower()) if len(token) > 3]

    @classmethod
    def _infer_relation(
        cls,
        node: ConceptNode,
        text: str,
        concepts: List[ConceptNode],
    ) -> str:
        domain_keywords = cls.CONTEXTUAL_KEYWORDS.get(node.domain, [])
        for keyword in domain_keywords:
            if keyword in text:
                return cls._build_context_sentence(node, keyword)

        related = cls._related_concepts(node, concepts, text)
        if related:
            return cls._build_related_context(node, related)

        return ""

    @classmethod
    def _related_concepts(
        cls,
        node: ConceptNode,
        concepts: List[ConceptNode],
        text: str,
    ) -> List[str]:
        related = []
        for other in concepts:
            if other.term == node.term:
                continue
            if other.term.lower() in text:
                related.append(other.term)
        return related

    @classmethod
    def _build_context_sentence(cls, node: ConceptNode, keyword: str) -> str:
        return (
            f"No contexto da pergunta, '{node.term}'  aplicado como um conceito de {node.domain}"
            f" relacionado a '{keyword}', indicando como o prompt utiliza seu significado prtico."
        )

    @classmethod
    def _build_related_context(cls, node: ConceptNode, related: List[str]) -> str:
        related_terms = ", ".join(related[:3])
        return (
            f"Neste caso, '{node.term}' aparece em conjunto com {related_terms},"
            f" o que sugere seu papel prtico na anlise do prompt."
        )

    @classmethod
    def _default_context(cls, node: ConceptNode) -> str:
        return (
            f"No contexto atual, '{node.term}' representa {node.definition.lower()}"
            f" e serve como um conceito relevante para o problema expresso no prompt."
        )

    @classmethod
    def summarize_application_context(cls, concepts: List[ConceptNode]) -> str:
        parts = [node.application_context for node in concepts if node.application_context]
        return " ".join(parts)

    @staticmethod
    def _normalize_text(text: str) -> str:
        return " ".join(text.split()).strip()
# ========== l5_generation.py ==========
"""
Camada L5  Gerao de resposta em texto livre.
================================================
A partir da sntese L4 (e contexto L1L3), gera resposta natural via LLM local (Ollama)
ou fallback para o template da L4. Opcional: LM customizado (EpistemicLanguageModel).
"""

import os
from typing import Optional

from layer_titles import LAYER_TITLES

# Resultado da L4
try:
    from l4_synthesis import SynthesisResult
except Exception:
    SynthesisResult = None  # type: ignore


def build_context_for_generation(
    prompt: str,
    synthesis_result: "SynthesisResult",
    concepts_summary: str = "",
    top_judgments: str = "",
) -> str:
    """Monta o contexto (texto) a ser enviado ao LLM para gerar a resposta final."""
    lines = [
        "## Contexto epistemolgico (L1L4)",
        f"Pergunta do usurio: {prompt}",
        "",
        f"Resposta sintetizada (L4): {synthesis_result.response}",
        f"Valor de verdade: {synthesis_result.truth_value:.2f} | Estado: {synthesis_result.state} | Certeza: {synthesis_result.certainty:+.2f}",
        "",
        "Use as seguintes nomenclaturas de seo para referenciar as etapas do raciocnio:",
        f"L1: {LAYER_TITLES['l1']}",
        f"L2: {LAYER_TITLES['l2']}",
        f"L3: {LAYER_TITLES['l3']}",
        f"L4: {LAYER_TITLES['l4']}",
        f"L5: {LAYER_TITLES['l5']}",
        f"L6: {LAYER_TITLES['l6']}",
        "",
    ]
    if synthesis_result.supporting_evidence:
        lines.append("Evidncias de suporte:")
        for ev in synthesis_result.supporting_evidence[:5]:
            lines.append(f"  - {ev}")
        lines.append("")
    if concepts_summary:
        lines.append(f"{LAYER_TITLES['l1']}: ")
        lines.append(concepts_summary)
        lines.append("")
    if top_judgments:
        lines.append(f"{LAYER_TITLES['l2']}: ")
        lines.append(top_judgments)
        lines.append("")
    lines.append("## Instruo")
    lines.append("Com base no contexto acima, elabore uma resposta final clara e precisa em portugus, sem repetir literalmente o texto da sntese. Seja conciso e cite a confiana quando relevante.")
    return "\n".join(lines)


def generate_with_ollama(
    context: str,
    model: str = "doninha8:latest",
    ollama_host: str = "http://localhost:11434",
    temperature: float = 0.3,
) -> str:
    """Gera resposta usando Ollama com o modelo local Doninha8."""
    try:
        import os
        import ollama

        if ollama_host:
            os.environ["OLLAMA_HOST"] = ollama_host

        if hasattr(ollama, "chat"):
            response = ollama.chat(
                model=model,
                messages=[{"role": "user", "content": context}],
                stream=False,
                options={
                    "temperature": temperature,
                    "num_ctx": 8192,
                },
            )
            if isinstance(response, dict):
                return response.get("message", {}).get("content", "").strip()
            return str(response).strip() if response else ""

        if hasattr(ollama, "Client"):
            client = ollama.Client(host=ollama_host)
            response = client.generate(
                model=model,
                prompt=context,
                stream=False,
                options={
                    "temperature": temperature,
                    "num_ctx": 8192,
                },
            )
            if isinstance(response, dict):
                return response.get("response", "").strip()
            return str(response).strip() if response else ""

        return ""
    except Exception as e:
        import sys
        print(f"Erro ao gerar com Ollama: {e}", file=sys.stderr)
        return ""


def generate_with_custom_lm(
    context: str,
    model_path: str,
    max_new_tokens: int = 150,
    temperature: float = 0.7,
) -> str:
    """Gera resposta usando EpistemicLanguageModel (custom_lm_model)."""
    try:
        from custom_lm_model import EpistemicLanguageModel, LMConfig, generate_text, load_lm
        from custom_tokenizer import CustomSPTokenizer, SPConfig
        import torch
        tokenizer = CustomSPTokenizer(SPConfig())
        tokenizer.load()
        vocab_size = tokenizer.vocab_size()
        model = load_lm(model_path, vocab_size)
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        out = generate_text(model, tokenizer, context, max_new_tokens=max_new_tokens, temperature=temperature, device=device)
        return out or ""
    except Exception:
        return ""


def generate_response(
    prompt: str,
    synthesis_result: "SynthesisResult",
    provider: str = "ollama",
    concepts_summary: str = "",
    top_judgments: str = "",
    custom_lm_path: str = "",
    ollama_model: str = "doninha8:latest",
    ollama_host: str = "http://localhost:11434",
) -> str:
    """
    Gera a resposta final em texto livre (ou template).
    provider: "ollama" | "template" | "custom_lm"
    """
    context = build_context_for_generation(prompt, synthesis_result, concepts_summary, top_judgments)

    if provider == "ollama":
        text = generate_with_ollama(context, model=ollama_model, ollama_host=ollama_host, temperature=0.3)
        if text:
            return text.strip()

    if provider == "custom_lm" and custom_lm_path:
        text = generate_with_custom_lm(context, custom_lm_path)
        if text:
            return text.strip()

    # Fallback: resposta da L4 (template)
    return synthesis_result.response
# ========== rag_hybrid_context_injection.py ==========
"""
RAG HBRIDO COM CONTEXT INJECTION
===================================
Camada de Retrieval-Augmented Generation (RAG) que trabalha de forma conjunta
com as camadas L1 e L2, usando um protocolo hbrido de:
  1. Context Injection (stuffing direto)  injeta contexto pr-selecionado
  2. Retrieval Seletivo por Domnios  busca documentos relevantes dinamicamente

A soluo  HIBRIDA: injeo direta + retrieval seletivo baseado em domnios.
Integrao com KB especializado (knowledge_base.py) e ChromaDB.
"""

from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import json
import re
from enum import Enum

try:
    from langchain_community.vectorstores import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings
    HAS_CHROMA = True
except ImportError:
    HAS_CHROMA = False

try:
    from knowledge_base import get_domain_knowledge_base, load_kb_from_file, merge_kb
except ImportError:
    get_domain_knowledge_base = None
    load_kb_from_file = None
    merge_kb = None


# 
# Enums e Estruturas de Dados
# 

class RetrievalStrategy(Enum):
    """Estratgia de retrieval seletivo."""
    DIRECT_INJECTION = "direct_injection"          # Apenas contexto injetado
    SEMANTIC_RETRIEVAL = "semantic_retrieval"      # Busca semntica em ChromaDB
    HYBRID = "hybrid"                               # Injeo + Retrieval seletivo
    DOMAIN_AWARE = "domain_aware"                   # Retrieval baseado em domnio


@dataclass
class DomainContext:
    """Contexto especializado de um domnio."""
    domain_name: str
    description: str = ""
    keywords: List[str] = field(default_factory=list)
    kb_path: str = ""                               # Caminho para KB do domnio
    chroma_collection: str = ""                     # Nome da coleo no ChromaDB
    system_prompt: str = ""                         # System prompt especializado
    injection_weight: float = 0.8                   # Peso da injeo direta [0,1]
    retrieval_weight: float = 0.2                   # Peso do retrieval [0,1]
    max_injected_docs: int = 3                      # Mx de docs injetados
    max_retrieved_docs: int = 5                     # Mx de docs recuperados


@dataclass
class RetrievedDocument:
    """Um documento recuperado do knowledge base."""
    content: str
    source: str = ""
    domain: str = ""
    relevance_score: float = 1.0
    is_injected: bool = False                       # Se vem de injeo direta
    metadata: Dict[str, Any] = field(default_factory=dict)

    def truncate(self, max_length: int = 500) -> str:
        """Trunca o contedo para no poluir o contexto."""
        if len(self.content) > max_length:
            return self.content[:max_length].rstrip() + "..."
        return self.content


@dataclass
class RAGContext:
    """Contexto hbrido compilado para injeo no prompt."""
    query: str
    domain: str = "geral"
    retrieved_documents: List[RetrievedDocument] = field(default_factory=list)
    injected_knowledge: Dict[str, float] = field(default_factory=dict)
    compiled_context: str = ""
    strategy: RetrievalStrategy = RetrievalStrategy.HYBRID
    confidence_score: float = 0.0


# 
# Sistema de Domnios Pr-configurados
# 

DEFAULT_DOMAINS: Dict[str, DomainContext] = {
    "filosofia": DomainContext(
        domain_name="filosofia",
        description="Filosofia, epistemologia, lgica clssica",
        keywords=["conhecimento", "verdade", "ser", "essncia", "substncia", "silogismo"],
        kb_path="data/kb_filosofia.json",
        chroma_collection="filosofia_corpus",
        system_prompt="""Voc  um especialista rigoroso em filosofia com acesso a uma base de conhecimento 
especializada em epistemologia, lgica e metafsica. Responda sempre usando o contexto fornecido quando 
relevante. Seja preciso, cite fontes filosficas e mantenha o rigor conceitual.""",
        injection_weight=0.8,
        retrieval_weight=0.2,
    ),
    "lgica": DomainContext(
        domain_name="lgica",
        description="Lgica formal, lgica paraconsistente, teoria de modelos",
        keywords=["proposio", "predicado", "quantificador", "inferncia", "validade", "contradio"],
        kb_path="data/kb_logica.json",
        chroma_collection="logica_corpus",
        system_prompt="""Voc  um especialista em lgica formal e paraconsistncia. Responda sempre 
usando o contexto fornecido quando relevante. Mantenha a preciso tcnica, use notao apropriada e 
cite definies formais quando necessrio.""",
        injection_weight=0.75,
        retrieval_weight=0.25,
    ),
    "epistemologia": DomainContext(
        domain_name="epistemologia",
        description="Epistemologia, teoria do conhecimento, justificao epistmica",
        keywords=["justificao", "crena", "conhecimento", "evidncia", "confiabilismo"],
        kb_path="data/kb_epistemologia.json",
        chroma_collection="epistemologia_corpus",
        system_prompt="""Voc  um especialista rigoroso em epistemologia com acesso a uma base de 
conhecimento especializada. Responda sempre usando o contexto fornecido quando relevante. Cite teorias 
epistemolgicas estabelecidas e seja preciso na caracterizao de conceitos.""",
        injection_weight=0.8,
        retrieval_weight=0.2,
    ),
    "geral": DomainContext(
        domain_name="geral",
        description="Conhecimento geral e enciclopdico",
        keywords=[],
        kb_path="data/kb.json",
        chroma_collection="general_corpus",
        system_prompt="""Voc  um especialista rigoroso com acesso a uma base de conhecimento especializada.
Responda sempre usando o contexto fornecido quando relevante. Seja preciso e cite fontes quando possvel.""",
        injection_weight=0.7,
        retrieval_weight=0.3,
    ),
}


# 
# Motor RAG Hbrido com Context Injection
# 

class HybridRAGContextInjectionEngine:
    """
    Motor principal de RAG hbrido que combina:
    - Context Injection (injeo direta de KB/documentos pr-selecionados)
    - Semantic Retrieval (busca em ChromaDB por similaridade)
    - Domain-Aware Selection (seleo baseada em domnio)
    
    A estratgia HYBRID usa injeo como contexto de base + retrieval seletivo.
    """

    def __init__(
        self,
        config: Optional[Dict[str, Any]] = None,
        embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
        chroma_path: str = "chromadb",
        verbose: bool = True,
    ):
        self.config = config or {}
        self.embedding_model = embedding_model
        self.chroma_path = Path(chroma_path)
        self.verbose = verbose
        self.domains = dict(DEFAULT_DOMAINS)
        self.chroma_stores: Dict[str, Any] = {}  # Cache de lojas ChromaDB
        self._initialize_chroma()

    def _initialize_chroma(self) -> None:
        """Inicializa conexes com ChromaDB para cada domnio."""
        if not HAS_CHROMA:
            if self.verbose:
                print("[RAG] ChromaDB no disponvel, usando apenas injeo direta.")
            return

        try:
            embeddings = HuggingFaceEmbeddings(model_name=self.embedding_model)
            for domain_name in self.domains:
                chroma_dir = self.chroma_path / domain_name
                if chroma_dir.exists() and chroma_dir.is_dir():
                    try:
                        store = Chroma(
                            persist_directory=str(chroma_dir),
                            embedding_function=embeddings,
                            collection_name=self.domains[domain_name].chroma_collection,
                        )
                        self.chroma_stores[domain_name] = store
                        if self.verbose:
                            print(f"[RAG] ChromaDB carregado para domnio '{domain_name}'")
                    except Exception as e:
                        if self.verbose:
                            print(f"[RAG] Erro ao carregar ChromaDB para '{domain_name}': {e}")
        except Exception as e:
            if self.verbose:
                print(f"[RAG] Erro ao inicializar ChromaDB: {e}")

    def register_domain(self, domain: DomainContext) -> None:
        """Registra um novo domnio."""
        self.domains[domain.domain_name] = domain

    def detect_domain(self, query: str, concepts: Optional[List[str]] = None) -> Tuple[str, float]:
        """
        Detecta qual domnio  mais relevante para a query usando keywords matching.
        Retorna (domain_name, confidence_score).
        """
        query_lower = query.lower()
        scores = {}

        for domain_name, domain_ctx in self.domains.items():
            score = 0.0
            if domain_ctx.keywords:
                for kw in domain_ctx.keywords:
                    if kw.lower() in query_lower:
                        score += 1.0
            if concepts:
                for concept in concepts:
                    if concept.lower() in query_lower:
                        score += 0.5

            scores[domain_name] = score

        # Normaliza scores
        max_score = max(scores.values()) if scores else 0.0
        if max_score > 0:
            best_domain = max(scores, key=scores.get)
            confidence = scores[best_domain] / (max_score + 1)
        else:
            best_domain = "geral"
            confidence = 0.1

        return best_domain, confidence

    def get_injected_knowledge(
        self,
        domain: str,
        query: Optional[str] = None,
    ) -> Dict[str, float]:
        """
        Recupera conhecimento para injeo direta do KB do domnio.
        Usa get_domain_knowledge_base se disponvel.
        """
        if not get_domain_knowledge_base:
            return {}

        try:
            kb = get_domain_knowledge_base(
                domain=domain,
                config=self.config,
                query_for_rag=query,
            )
            return kb
        except Exception as e:
            if self.verbose:
                print(f"[RAG] Erro ao recuperar KB do domnio '{domain}': {e}")
            return {}

    def retrieve_documents(
        self,
        query: str,
        domain: str = "geral",
        k: int = 5,
        strategy: RetrievalStrategy = RetrievalStrategy.HYBRID,
    ) -> List[RetrievedDocument]:
        """
        Recupera documentos relevantes usando a estratgia especificada.
        
        Strategies:
        - DIRECT_INJECTION: Sem retrieval, apenas contexto injetado
        - SEMANTIC_RETRIEVAL: Apenas busca em ChromaDB
        - HYBRID: Injeo + Retrieval seletivo
        - DOMAIN_AWARE: Retrieval especfico do domnio
        """
        results: List[RetrievedDocument] = []

        if strategy == RetrievalStrategy.DIRECT_INJECTION:
            # Apenas contexto injetado, sem retrieval dinmico
            return results

        domain_ctx = self.domains.get(domain, self.domains["geral"])
        max_injected = domain_ctx.max_injected_docs
        max_retrieved = domain_ctx.max_retrieved_docs

        # 
        # Estratgia HYBRID: Injeo + Retrieval seletivo
        # 
        if strategy in (RetrievalStrategy.HYBRID, RetrievalStrategy.DOMAIN_AWARE):
            # Etapa 1: Contexto injetado (KB direto)
            injected_kb = self.get_injected_knowledge(domain, query)
            if injected_kb:
                # Seleciona top-k termos por relevncia
                sorted_terms = sorted(injected_kb.items(), key=lambda x: x[1], reverse=True)
                for i, (term, score) in enumerate(sorted_terms[:max_injected]):
                    results.append(
                        RetrievedDocument(
                            content=f"Termo: {term}",
                            source=f"KB-{domain}",
                            domain=domain,
                            relevance_score=float(score),
                            is_injected=True,
                            metadata={"type": "kb_term", "weight": score},
                        )
                    )

        # Etapa 2: Retrieval semntico (ChromaDB)
        if strategy in (RetrievalStrategy.SEMANTIC_RETRIEVAL, RetrievalStrategy.HYBRID):
            if domain in self.chroma_stores:
                try:
                    chroma = self.chroma_stores[domain]
                    docs = chroma.similarity_search(query, k=max_retrieved)
                    for doc in docs:
                        # Extrai score se disponvel
                        score = getattr(doc, "metadata", {}).get("score", 0.8)
                        results.append(
                            RetrievedDocument(
                                content=doc.page_content if hasattr(doc, "page_content") else str(doc),
                                source=f"ChromaDB-{domain}",
                                domain=domain,
                                relevance_score=float(score),
                                is_injected=False,
                                metadata=getattr(doc, "metadata", {}),
                            )
                        )
                except Exception as e:
                    if self.verbose:
                        print(f"[RAG] Erro ao recuperar de ChromaDB-{domain}: {e}")

        # Ordena por relevncia
        results.sort(key=lambda x: x.relevance_score, reverse=True)
        return results[:k]

    def compile_context(
        self,
        query: str,
        retrieved_docs: List[RetrievedDocument],
        injected_kb: Optional[Dict[str, float]] = None,
        domain: str = "geral",
        include_system_prompt: bool = True,
    ) -> RAGContext:
        """
        Compila o contexto final para injeo no prompt.
        Combina documentos recuperados, KB injetado e system prompt.
        """
        domain_ctx = self.domains.get(domain, self.domains["geral"])
        lines = []

        # 
        # Parte 1: System Prompt especializado
        # 
        if include_system_prompt and domain_ctx.system_prompt:
            lines.append("## Instrues do Sistema")
            lines.append(domain_ctx.system_prompt)
            lines.append("")

        # 
        # Parte 2: Documentos Injetados (Context Injection)
        # 
        injected_docs = [d for d in retrieved_docs if d.is_injected]
        if injected_docs:
            lines.append("## Contexto Base Injetado (Domnio)")
            for doc in injected_docs:
                lines.append(f"- **{doc.source}** [{doc.relevance_score:.2f}]: {doc.truncate()}")
            lines.append("")

        # 
        # Parte 3: Documentos Recuperados (Semantic Retrieval)
        # 
        retrieved_only = [d for d in retrieved_docs if not d.is_injected]
        if retrieved_only:
            lines.append("## Contexto Recuperado (ChromaDB)")
            for doc in retrieved_only:
                lines.append(f"- **{doc.source}**: {doc.truncate()}")
            lines.append("")

        # 
        # Parte 4: Knowledge Base Terms (se fornecido)
        # 
        if injected_kb:
            lines.append("## Termos-Chave do Knowledge Base")
            sorted_terms = sorted(injected_kb.items(), key=lambda x: x[1], reverse=True)[:10]
            for term, score in sorted_terms:
                lines.append(f"- {term}: {score:.2f}")
            lines.append("")

        # 
        # Parte 5: Instruo de Resposta
        # 
        lines.append("## Pergunta do Usurio")
        lines.append(f"{query}")
        lines.append("")
        lines.append("---")
        lines.append("Baseando-se no contexto injetado e recuperado acima, elabore uma resposta rigorosa.")
        lines.append("")

        compiled = "\n".join(lines)

        # Calcula confidence score
        conf = 0.0
        if injected_docs:
            conf += domain_ctx.injection_weight * (sum(d.relevance_score for d in injected_docs) / len(injected_docs))
        if retrieved_only:
            conf += domain_ctx.retrieval_weight * (sum(d.relevance_score for d in retrieved_only) / len(retrieved_only))
        conf = min(1.0, conf)

        return RAGContext(
            query=query,
            domain=domain,
            retrieved_documents=retrieved_docs,
            injected_knowledge=injected_kb or {},
            compiled_context=compiled,
            strategy=RetrievalStrategy.HYBRID,
            confidence_score=conf,
        )

    def process(
        self,
        query: str,
        concepts: Optional[List[str]] = None,
        strategy: RetrievalStrategy = RetrievalStrategy.HYBRID,
        k: int = 8,
        auto_detect_domain: bool = True,
    ) -> RAGContext:
        """
        Pipeline completo de RAG hbrido.
        
        Etapas:
        1. Detecta domnio (se auto_detect_domain=True)
        2. Recupera documentos (injeo + retrieval)
        3. Compila contexto final
        4. Retorna RAGContext pronto para injeo
        """
        # Detecta domnio
        if auto_detect_domain:
            domain, conf = self.detect_domain(query, concepts)
            if self.verbose:
                print(f"[RAG] Domnio detectado: {domain} (confiana: {conf:.2f})")
        else:
            domain = "geral"

        # Recupera conhecimento injetado
        injected_kb = self.get_injected_knowledge(domain, query)

        # Recupera documentos
        retrieved_docs = self.retrieve_documents(
            query=query,
            domain=domain,
            k=k,
            strategy=strategy,
        )

        # Compila contexto
        rag_context = self.compile_context(
            query=query,
            retrieved_docs=retrieved_docs,
            injected_kb=injected_kb,
            domain=domain,
            include_system_prompt=True,
        )

        return rag_context

    def format_for_l1_l2(self, rag_context: RAGContext) -> Dict[str, Any]:
        """
        Formata o contexto RAG para consumo pelas camadas L1 (Conceitos) e L2 (Juzos).
        Retorna um dicionrio com:
        - domain: domnio detectado
        - injected_context: string do contexto injetado
        - kb_terms: dicionrio termo -> score
        - system_prompt: system prompt especializado
        - documents: lista de documentos
        """
        domain_ctx = self.domains.get(rag_context.domain, self.domains["geral"])

        return {
            "domain": rag_context.domain,
            "injected_context": rag_context.compiled_context,
            "kb_terms": rag_context.injected_knowledge,
            "system_prompt": domain_ctx.system_prompt,
            "documents": [
                {
                    "content": doc.truncate(1000),
                    "source": doc.source,
                    "relevance": doc.relevance_score,
                    "is_injected": doc.is_injected,
                }
                for doc in rag_context.retrieved_documents
            ],
            "confidence": rag_context.confidence_score,
        }


# 
# Funes Auxiliares de Alto Nvel
# 

def create_hybrid_rag_engine(
    config: Optional[Dict[str, Any]] = None,
    chroma_path: str = "chromadb",
) -> HybridRAGContextInjectionEngine:
    """Factory para criar uma instncia do motor RAG."""
    return HybridRAGContextInjectionEngine(config=config, chroma_path=chroma_path)


def process_query_with_rag(
    query: str,
    concepts: Optional[List[str]] = None,
    domain: Optional[str] = None,
    auto_detect: bool = True,
    config: Optional[Dict[str, Any]] = None,
) -> RAGContext:
    """
    Funo de convenincia para processar uma query com RAG hbrido.
    
    Exemplo:
        rag_ctx = process_query_with_rag("O que  conhecimento?", domain="epistemologia")
        print(rag_ctx.compiled_context)
    """
    engine = create_hybrid_rag_engine(config=config)
    return engine.process(
        query=query,
        concepts=concepts,
        auto_detect_domain=auto_detect,
        strategy=RetrievalStrategy.HYBRID,
    )
# ========== l2_kantian_judgments.py ==========
"""
CAMADA L2  Tbua de Juzos Kantianos
======================================
Antes de qualquer clculo estatstico o prompt  destrinchado nas
doze categorias da Tbua dos Juzos (Kritik der reinen Vernunft, 9).

Dimenses:
  Quantidade   Universal | Particular | Singular
  Qualidade    Afirmativo | Negativo | Infinito
  Relao      Categrico | Hipottico | Disjuntivo
  Modalidade   Problemtico | Assertrico | Apodtico

Cada hiptese gerada recebe um peso de prioridade; o Juzo Singular
Afirmativo Assertrico tem prioridade mxima ( a resposta-alvo).
"""

from dataclasses import dataclass, field
from typing import List, Tuple, Optional
from l1_concept_table import ConceptNode, ConceptTable
import re

try:
    from transformers import pipeline
except ImportError:
    pipeline = None


# 
# Estruturas de dados
# 

@dataclass
class EpistemicClassification:
    """Classificao epistemolgica sem restrio T+I+F=1."""
    truth: float = 0.0           # T  [0,1]  grau de verdade
    indeterminacy: float = 0.0   # I  [0,1]  grau de indeterminao
    falsity: float = 0.0         # F  [0,1]  grau de falsidade
    classification: str = "indeterminado"  # paraconsistncia | incompletude | vagueza | assertiva_confiante | indeterminado

    def __post_init__(self):
        self.classification = self._classify()

    def _classify(self) -> str:
        """Aplica regras epistemolgicas para classificar."""
        if self.truth + self.falsity > 1.0:
            return "paraconsistncia"
        if self.truth + self.indeterminacy + self.falsity < 1.0:
            return "incompletude"
        if self.indeterminacy > 0.6:
            return "vagueza"
        if self.truth > 0.7 and self.indeterminacy < 0.2 and self.falsity < 0.2:
            return "assertiva_confiante"
        return "indeterminado"

    def __str__(self) -> str:
        return (
            f"T={self.truth:.2f} I={self.indeterminacy:.2f} F={self.falsity:.2f} "
            f"[{self.classification}]"
        )


@dataclass
class KantianJudgment:
    """Uma proposio refinada segundo a tbua dos juzos."""
    quantidade:  str   # Universal | Particular | Singular
    qualidade:   str   # Afirmativo | Negativo | Infinito
    relacao:     str   # Categrico | Hipottico | Disjuntivo
    modalidade:  str   # Problemtico | Assertrico | Apodtico
    proposicao:  str   # texto da hiptese
    prioridade:  float = 0.0  # 0.0  1.0  (1.0 = resposta-alvo)
    epistemic_classification: EpistemicClassification = field(default_factory=EpistemicClassification)

    def __str__(self) -> str:
        return (
            f"[{self.quantidade}/{self.qualidade}/"
            f"{self.relacao}/{self.modalidade}] "
            f"(pri={self.prioridade:.2f}) {self.epistemic_classification} {self.proposicao}"
        )


@dataclass
class SyntaxProfile:
    """
    Perfil sinttico mnimo extrado do enunciado segundo a gramtica
    (aproximao heurstica baseada em listas inspiradas em grammar.txt).
    """
    quantifier_subject: Optional[str] = None   # "all", "some", "this", etc.
    quantifier_predicate: Optional[str] = None
    has_negation: bool = False
    has_infinite_like: bool = False           # construes do tipo "not-X"
    is_conditional: bool = False              # presena de "if", "then"
    is_disjunctive: bool = False              # presena de "or"
    modality_markers: Tuple[str, ...] = ()    # "can", "must", "might", etc.


# 
# Regras de prioridade entre modalidades (herana da "parte fraca")
# 
MODALIDADE_PESO = {
    "Apodtico":    1.0,
    "Assertrico":  0.7,
    "Problemtico": 0.4,
}
QUANTIDADE_PESO = {
    "Singular":   1.0,
    "Particular": 0.6,
    "Universal":  0.3,
}
QUALIDADE_PESO = {
    "Afirmativo": 1.0,
    "Infinito":   0.6,
    "Negativo":   0.4,
}
RELACAO_PESO = {
    "Categrico":  1.0,
    "Hipottico":  0.7,
    "Disjuntivo":  0.5,
}


def _priority(j: KantianJudgment) -> float:
    return (
        MODALIDADE_PESO[j.modalidade]
        * QUANTIDADE_PESO[j.quantidade]
        * QUALIDADE_PESO[j.qualidade]
        * RELACAO_PESO[j.relacao]
    )


# 
# Motor de gerao de juzos
# 

class BERTAssertionClassifier:
    """Classificador baseado em BERT para juzos assertricos.

    Processa proposies e retorna (T, I, F) sem restrio T+I+F=1,
    capturando paraconsistncia, incompletude e vagueza.
    """

    DOMAIN_CANDIDATES = {
        "fsico": [
            "empiricamente verificado",
            "teoricamente plausvel",
            "logicamente contraditrio",
            "empiricamente indeterminado",
        ],
        "lgica": [
            "logicamente contraditrio",
            "teoricamente plausvel",
            "empiricamente indeterminado",
            "empiricamente verificado",
        ],
        "cognitivo": [
            "empiricamente verificado",
            "teoricamente plausvel",
            "empiricamente indeterminado",
            "logicamente contraditrio",
        ],
        "filosfico": [
            "teoricamente plausvel",
            "empiricamente indeterminado",
            "logicamente contraditrio",
            "empiricamente verificado",
        ],
        "geral": [
            "teoricamente plausvel",
            "empiricamente verificado",
            "logicamente contraditrio",
            "empiricamente indeterminado",
        ],
    }
    GENERIC_CANDIDATES = [
        "empiricamente verificado",
        "teoricamente plausvel",
        "logicamente contraditrio",
        "empiricamente indeterminado",
    ]

    def __init__(self):
        self.classifier = None
        if pipeline is not None:
            try:
                self.classifier = pipeline(
                    "zero-shot-classification",
                    model="bert-base-multilingual-uncased",
                )
            except Exception:
                pass

    def classify(self, proposition: str, domain: str = "geral") -> EpistemicClassification:
        """Classifica uma proposio em (T, I, F) usando candidatos por domnio."""
        if self.classifier is None:
            return self._heuristic_classify(proposition)

        candidates = self.DOMAIN_CANDIDATES.get(domain, self.GENERIC_CANDIDATES)
        try:
            result = self.classifier(proposition, candidates, multi_class=True)
            scores = {label: score for label, score in zip(result["labels"], result["scores"])}
            return EpistemicClassification(
                truth=scores.get("empiricamente verificado", 0.0)
                + scores.get("teoricamente plausvel", 0.0) * 0.75,
                indeterminacy=scores.get("empiricamente indeterminado", 0.0),
                falsity=scores.get("logicamente contraditrio", 0.0),
            )
        except Exception:
            return self._heuristic_classify(proposition)

    def _heuristic_classify(self, proposition: str) -> EpistemicClassification:
        """Classificao heurstica quando BERT no est disponvel."""
        text = proposition.lower()
        t, i, f = 0.5, 0.3, 0.2

        if "verdadeiro" in text or "" in text or "sempre" in text:
            t = 0.8
            i = 0.1
            f = 0.1
        elif "falso" in text or "nunca" in text or "no " in text:
            t = 0.1
            i = 0.1
            f = 0.8
        elif "pode" in text or "talvez" in text or "possvel" in text:
            t = 0.4
            i = 0.5
            f = 0.3
        elif "contraditrio" in text or "e" in text and "ou" in text:
            t = 0.6
            i = 0.3
            f = 0.7
        elif "indeterminado" in text or "indefinido" in text:
            t = 0.3
            i = 0.7
            f = 0.3
        elif "incompleto" in text or "insuficiente" in text:
            t = 0.2
            i = 0.6
            f = 0.2

        return EpistemicClassification(truth=round(t, 3), indeterminacy=round(i, 3), falsity=round(f, 3))


class KantianJudgmentEngine:
    """
    Recebe um prompt e a lista de ConceptNodes extrados por L1 e devolve
    as 12 hipteses estruturadas segundo a tbua kantiana.
    
    Para juizos assertricos, aplica classificao BERT com (T, I, F).
    """

    def __init__(self, concept_table: ConceptTable) -> None:
        self.ct = concept_table
        self.bert_classifier = BERTAssertionClassifier()

    # ------------------------------------------------------------------ #
    # API pblica                                                          #
    # ------------------------------------------------------------------ #

    def refine(self, prompt: str, concepts: List[ConceptNode]) -> List[KantianJudgment]:
        """
        Gera as hipteses kantianas para o prompt e as ordena por
        prioridade descendente.
        """
        subject, predicates = self._parse_prompt(prompt, concepts)
        syntax = self._analyze_syntax(prompt)
        judgments: List[KantianJudgment] = []

        domain = self._infer_domain(concepts)
        for pred in predicates:
            antonym = self._antonym_of(pred, concepts)
            hypernym = self._hypernym_of(pred, concepts)

            #  Juzo principal guiado pela gramtica 
            qt = self._infer_quantity(syntax)
            ql = self._infer_quality(syntax)
            rel = self._infer_relation(syntax)
            mod = self._infer_modality(syntax)

            base_prop = f"{subject}  {pred}"
            if syntax.has_negation and antonym:
                base_prop = f"{subject} no  {antonym}"

            j = self._make(qt, ql, rel, mod, base_prop)
            if j.modalidade == "Assertrico":
                j.epistemic_classification = self.bert_classifier.classify(base_prop, domain=domain)
            judgments.append(j)

            #  Variaes cannicas (mantidas, mas ancoradas em L1) 
            judgments.append(self._make(
                "Universal", "Afirmativo", "Categrico", "Apodtico",
                f"Todo(a) {subject} com propriedade extrema  {pred}",
            ))
            judgments.append(self._make(
                "Particular", "Afirmativo", "Hipottico", "Problemtico",
                f"Algum(a) {subject} pode ser {pred}",
            ))
            j1 = self._make(
                "Singular", "Afirmativo", "Categrico", "Assertrico",
                f"Este(a) {subject} especfico  {pred}",
            )
            j1.epistemic_classification = self.bert_classifier.classify(j1.proposicao, domain=domain)
            judgments.append(j1)

            prop2 = (f"Este(a) {subject} no  {antonym}" if antonym else
                     f"Este(a) {subject} no possui a propriedade oposta a {pred}")
            j2 = self._make("Singular", "Negativo", "Categrico", "Assertrico", prop2)
            j2.epistemic_classification = self.bert_classifier.classify(j2.proposicao, domain=domain)
            judgments.append(j2)
            j3 = self._make(
                "Singular", "Infinito", "Categrico", "Assertrico",
                f"Este(a) {subject}  no-{antonym}" if antonym else
                f"Este(a) {subject}  indeterminado em relao a {pred}",
            )
            j3.epistemic_classification = self.bert_classifier.classify(j3.proposicao, domain=domain)
            judgments.append(j3)

            judgments.append(self._make(
                "Universal", "Afirmativo", "Hipottico", "Apodtico",
                f"Se {subject} possui condio X, ento  {pred}",
            ))
            j4 = self._make(
                "Universal", "Afirmativo", "Disjuntivo", "Assertrico",
                f"{subject}  {pred} OU {antonym} OU intermedirio"
                if antonym else f"{subject}  {pred} ou outra propriedade",
            )
            j4.epistemic_classification = self.bert_classifier.classify(j4.proposicao, domain=domain)
            judgments.append(j4)

            judgments.append(self._make(
                "Singular", "Afirmativo", "Categrico", "Problemtico",
                f"Este(a) {subject} pode ser {pred}?",
            ))
            j5 = self._make(
                "Singular", "Afirmativo", "Hipottico", "Assertrico",
                f"Este(a) {subject}  {pred} em razo das condies observadas",
            )
            j5.epistemic_classification = self.bert_classifier.classify(j5.proposicao, domain=domain)
            judgments.append(j5)
            judgments.append(self._make(
                "Universal", "Afirmativo", "Categrico", "Apodtico",
                f"{subject} deve ser {pred} quando condies necessrias presentes",
            ))

            #  HIPTESES COM INTERMEDIRIOS (hiperonmia) 
            if hypernym:
                j6 = self._make(
                    "Singular", "Afirmativo", "Categrico", "Assertrico",
                    f"Este(a) {subject} pertence  categoria {hypernym}",
                )
                j6.epistemic_classification = self.bert_classifier.classify(j6.proposicao, domain=domain)
                judgments.append(j6)
            if antonym:
                j7 = self._make(
                    "Singular", "Negativo", "Disjuntivo", "Assertrico",
                    f"Este(a) {subject} no  {pred} nem {antonym}: "
                    f"admite valor intermedirio",
                )
                j7.epistemic_classification = self.bert_classifier.classify(j7.proposicao, domain=domain)
                judgments.append(j7)

        # Calcula prioridades e ordena
        for j in judgments:
            j.prioridade = _priority(j)
        judgments.sort(key=lambda j: j.prioridade, reverse=True)
        return judgments

    def _infer_domain(self, concepts: List[ConceptNode]) -> str:
        """Inferncia simples de domnio majoritrio a partir dos conceitos extrados."""
        if not concepts:
            return "geral"
        domain_counts = {}
        for concept in concepts:
            domain = concept.domain.lower().strip() if concept.domain else "geral"
            domain_counts[domain] = domain_counts.get(domain, 0) + 1
        return max(domain_counts, key=domain_counts.get)

    # ------------------------------------------------------------------ #
    # Helpers                                                              #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _make(qt, ql, rel, mod, prop) -> KantianJudgment:
        j = KantianJudgment(
            quantidade=qt, qualidade=ql, relacao=rel,
            modalidade=mod, proposicao=prop,
        )
        j.prioridade = _priority(j)
        return j

    @staticmethod
    def _parse_prompt(prompt: str, concepts: List[ConceptNode]) -> Tuple[str, List[str]]:
        """
        Extrai sujeito e predicados candidatos do prompt de forma simples.
        Em produo seria substitudo por um parser sinttico.
        """
        tokens = re.findall(r"[a-zA-Z]+", prompt.lower())
        known = {c.term.lower() for c in concepts}
        subject = tokens[0] if tokens else "entidade"
        predicates = [t for t in tokens[1:] if t in known] or ["indeterminado"]
        return subject, predicates

    # ------------------------------------------------------------------ #
    # Anlise sinttica inspirada em grammar.txt                         #
    # ------------------------------------------------------------------ #

    def _analyze_syntax(self, prompt: str) -> SyntaxProfile:
        """
        Extrai um perfil sinttico mnimo usando listas de palavras
        alinhadas aos captulos de determiners, modals, negatives e
        conjunctions da grammar COBUILD.
        """
        text = prompt.lower()
        tokens = re.findall(r"[a-z]+", text)

        quant_all = {"all", "every", "each"}
        quant_some = {"some", "many", "several", "few", "a few"}
        quant_singular = {"this", "that", "these", "those", "a", "an", "one"}

        neg_markers = {"not", "no", "never", "none", "nothing", "nowhere"}
        infinite_patterns = {"not-", "non-"}

        cond_markers = {"if", "provided", "unless", "whenever", "as long as"}
        disj_markers = {"or", "either"}

        modal_poss = {"can", "could", "may", "might"}
        modal_necess = {"must", "have to", "need to", "should", "ought"}

        has_neg = any(tok in neg_markers for tok in tokens)
        has_inf = any(pat in text for pat in infinite_patterns)
        is_cond = any(tok in cond_markers for tok in tokens)
        is_disj = any(tok in disj_markers for tok in tokens)

        mods: list[str] = []
        for tok in tokens:
            if tok in modal_poss or tok in modal_necess:
                mods.append(tok)

        q_subj: Optional[str] = None
        q_pred: Optional[str] = None

        if tokens:
            first = tokens[0]
            if first in quant_all:
                q_subj = "all"
            elif first in quant_some:
                q_subj = "some"
            elif first in quant_singular:
                q_subj = "this"

        return SyntaxProfile(
            quantifier_subject=q_subj,
            quantifier_predicate=q_pred,
            has_negation=has_neg,
            has_infinite_like=has_inf,
            is_conditional=is_cond,
            is_disjunctive=is_disj,
            modality_markers=tuple(mods),
        )

    def _infer_quantity(self, syntax: SyntaxProfile) -> str:
        if syntax.quantifier_subject == "all":
            return "Universal"
        if syntax.quantifier_subject == "some":
            return "Particular"
        if syntax.quantifier_subject == "this":
            return "Singular"
        return "Singular"

    def _infer_quality(self, syntax: SyntaxProfile) -> str:
        if syntax.has_infinite_like:
            return "Infinito"
        if syntax.has_negation:
            return "Negativo"
        return "Afirmativo"

    def _infer_relation(self, syntax: SyntaxProfile) -> str:
        if syntax.is_conditional:
            return "Hipottico"
        if syntax.is_disjunctive:
            return "Disjuntivo"
        return "Categrico"

    def _infer_modality(self, syntax: SyntaxProfile) -> str:
        markers = {m for m in syntax.modality_markers}
        if any(m in {"must", "have", "need", "should", "ought"} for m in markers):
            return "Apodtico"
        if any(m in {"can", "could", "may", "might"} for m in markers):
            return "Problemtico"
        return "Assertrico"

    def _antonym_of(self, term: str, concepts: List[ConceptNode]) -> str:
        node = self.ct.get(term)
        if node and node.antonyms:
            return node.antonyms[0]
        return ""

    def _hypernym_of(self, term: str, concepts: List[ConceptNode]) -> str:
        node = self.ct.get(term)
        if node and node.hypernyms:
            return node.hypernyms[0]
        return ""
# ========== l3_paraconsistent.py ==========
"""
CAMADA L3  Avaliao Paraconsistente
======================================
Implementa a Lgica Anotada de Evidncias (LAE / PAL2v) de da Costa & Abe.

Cada proposio recebe um par de anotaes:
       [0,1]   grau de evidncia FAVORVEL
       [0,1]   grau de evidncia CONTRRIA

Estados resultantes:
  
    Verdadeiro     :  alto,  baixo                
    Falso          :  baixo,  alto                
    Inconsistente  :  alto,  alto  (contradio)  
    Indeterminado  :  baixo,  baixo               
    Intermedirio  : valores mdios  (morno, etc.)  
  

Princpio central:
    "Contradio Local + Consistncia Global  Trivializao"

A exploso  GENTIL: uma contradio local (quente e frio) no
trivializa o sistema  produz o estado "Intermedirio" (morno).
"""

from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional
import math
import re

import torch

try:
    from paraconsistent_rules import (
        state_from_rules,
        state_12_to_simple,
        load_rules_from_fuzzy_file,
        ParaconsistentRules,
    )
except Exception:
    state_from_rules = None  # type: ignore
    state_12_to_simple = None  # type: ignore
    load_rules_from_fuzzy_file = None  # type: ignore
    ParaconsistentRules = None  # type: ignore

try:
    # Import opcional do modelo neural; o sistema continua funcional sem ele.
    from neural_truth_model import TruthScoringModel, load_tokenizer, neural_annotations
except Exception:  # pragma: no cover - fallback em ambientes sem transformers
    TruthScoringModel = None  # type: ignore
    load_tokenizer = None  # type: ignore
    neural_annotations = None  # type: ignore


# 
# Constantes de limiar
# 
THRESHOLD_TRUE         = 0.7   #   este valor e   (1 - este)  Verdadeiro
THRESHOLD_FALSE        = 0.3   #   este e   (1 - este)  Falso
THRESHOLD_INCONSISTENT = 0.6   # ambos acima  Inconsistente (contradio local)
THRESHOLD_INDETERMINATE= 0.4   # ambos abaixo  Indeterminado

PROPOSITION_TYPE_A = "A"  # Universal Afirmativa
PROPOSITION_TYPE_E = "E"  # Universal Negativa
PROPOSITION_TYPE_I = "I"  # Particular Afirmativa
PROPOSITION_TYPE_O = "O"  # Particular Negativa

PROPOSITION_TYPE_LABELS = {
    PROPOSITION_TYPE_A: "Universal Afirmativa",
    PROPOSITION_TYPE_E: "Universal Negativa",
    PROPOSITION_TYPE_I: "Particular Afirmativa",
    PROPOSITION_TYPE_O: "Particular Negativa",
}


def infer_proposition_type(text: str) -> Optional[str]:
    """Heurstica simples para inferir o tipo A / E / I / O a partir do texto."""
    normalized = text.lower().strip()
    if not normalized:
        return None

    # Detecta padres de proposies particulares negativas antes das afirmativas.
    if re.search(r"\b(algum no|alguns no|alguma no|algumas no|nem todos|pelo menos um no)\b", normalized):
        return PROPOSITION_TYPE_O
    if re.search(r"\b(nenhum|nenhuma|nunca|jamais|sem nenhum|sem nenhuma|no existe|no h)\b", normalized):
        return PROPOSITION_TYPE_E
    if re.search(r"\b(algum|alguma|alguns|algumas|pelo menos um|h um|h algum|existem|existe)\b", normalized):
        return PROPOSITION_TYPE_I
    if re.search(r"\b(todo|todos|toda|todas|cada|sempre|qualquer)\b", normalized):
        return PROPOSITION_TYPE_A

    return None


def type_label(proposition_type: Optional[str]) -> str:
    return PROPOSITION_TYPE_LABELS.get(proposition_type, "Desconhecido")


@dataclass
class ParaconsistentValue:
    """
    Valor-verdade paraconsistente para uma proposio.
    Baseado na Lgica Anotada de Evidncias (LAE).
    """
    proposition: str
    mu: float          # evidncia favorvel   [0,1]
    lam: float         # evidncia contrria   [0,1]
    proposition_type: Optional[str] = None

    @property
    def proposition_kind(self) -> Optional[str]:
        """Retorna o tipo de proposio A/E/I/O, inferido do texto se necessrio."""
        return self.proposition_type or infer_proposition_type(self.proposition)

    @property
    def proposition_type_label(self) -> str:
        return type_label(self.proposition_kind)

    #  Graus derivados  #
    @property
    def certainty(self) -> float:
        """Grau de certeza: Gc =       [1, 1]"""
        return self.mu - self.lam

    @property
    def contradiction(self) -> float:
        """Grau de contradio: Gct =  +   1    [1, 1]"""
        return self.mu + self.lam - 1.0

    @property
    def state(self) -> str:
        """Estado lgico qualitativo. Usa regras do Fuzzy.txt se disponveis."""
        if state_from_rules is not None and state_12_to_simple is not None:
            state_12 = state_from_rules(self.mu, self.lam)
            return state_12_to_simple(state_12)
        # Fallback: limiares fixos
        if self.mu >= THRESHOLD_TRUE and self.lam <= (1 - THRESHOLD_TRUE):
            return "Verdadeiro"
        if self.mu <= THRESHOLD_FALSE and self.lam >= (1 - THRESHOLD_FALSE):
            return "Falso"
        if self.mu >= THRESHOLD_INCONSISTENT and self.lam >= THRESHOLD_INCONSISTENT:
            return "Inconsistente_local"   # exploso GENTIL  no trivializa
        if self.mu <= THRESHOLD_INDETERMINATE and self.lam <= THRESHOLD_INDETERMINATE:
            return "Indeterminado"
        return "Intermedirio"             # ex: morno entre quente e frio

    @property
    def state_12(self) -> Optional[str]:
        """Estado lgico de 12 valores (reticulado) conforme Fuzzy.txt, se regras carregadas."""
        if state_from_rules is not None:
            return state_from_rules(self.mu, self.lam)
        return None

    @property
    def truth_value(self) -> float:
        """Valor-verdade escalar normalizado para sada final."""
        return round((self.mu + (1 - self.lam)) / 2.0, 4)

    def __str__(self) -> str:
        type_label_text = self.proposition_type_label
        return (
            f"  ={self.mu:.3f}  ={self.lam:.3f}  "
            f"Gc={self.certainty:+.3f}  Gct={self.contradiction:+.3f}  "
            f"v={self.truth_value:.3f}  [{self.state}]\n"
            f"  Tipo={type_label_text}  \"{self.proposition}\""
        )


class ManyValuedRouter:
    """Roteador fuzzy para distinguir contradio lgica real de incerteza estatstica."""

    REAL_CONTRADICTION = "Contradio_real"
    STATISTICAL_UNCERTAINTY = "Incerteza_estatstica"
    AMBIGUOUS = "Ambguo"
    UNCLASSIFIED = "No_classificado"

    @staticmethod
    def _pair_strength(left: ParaconsistentValue, right: ParaconsistentValue) -> float:
        """Grau fuzzy de suporte conjunto entre duas proposies."""
        return round(min(left.mu, right.mu), 4)

    @classmethod
    def route_pair(
        cls,
        left: ParaconsistentValue,
        right: ParaconsistentValue,
    ) -> Tuple[str, float, str]:
        """Classifica o par de proposies como contradio real, incerteza ou ambguo."""
        left_type = left.proposition_kind
        right_type = right.proposition_kind
        label = f"{left_type or '?'} vs {right_type or '?'}"

        if {left_type, right_type} == {PROPOSITION_TYPE_A, PROPOSITION_TYPE_I}:
            strength = cls._pair_strength(left, right)
            return cls.REAL_CONTRADICTION, strength, f"Contraposio A/I detectada ({label})"

        if {left_type, right_type} == {PROPOSITION_TYPE_E, PROPOSITION_TYPE_I}:
            strength = cls._pair_strength(left, right)
            return cls.STATISTICAL_UNCERTAINTY, strength, f"Contraposio E/I detectada ({label})"

        if left_type is None or right_type is None:
            return cls.AMBIGUOUS, 0.0, "Tipo de proposio no identificado"

        return cls.UNCLASSIFIED, 0.0, f"Par {label} no corresponde a A/I nem E/I"

    @classmethod
    def route_pairwise(
        cls,
        values: List[ParaconsistentValue],
    ) -> List[Tuple[ParaconsistentValue, ParaconsistentValue, str, float, str]]:
        """Avalia todos os pares de proposies para identificao de rota lgica."""
        routes: List[Tuple[ParaconsistentValue, ParaconsistentValue, str, float, str]] = []
        n = len(values)
        for i in range(n):
            for j in range(i + 1, n):
                route, confidence, explanation = cls.route_pair(values[i], values[j])
                if route != cls.UNCLASSIFIED:
                    routes.append((values[i], values[j], route, confidence, explanation))
        return routes


# 
# Motor paraconsistente
# 

class ParaconsistentEngine:
    """
    Avalia as hipteses kantiana (L2) e atribui valores-verdade
    paraconsistentes a cada uma.

    Pode operar em dois modos:
      - Modo heurstico (padro): usa apenas o banco de conhecimento.
      - Modo neural: se um TruthScoringModel for fornecido, usa o modelo
        para calcular (, ) compatveis com a Lgica Anotada.
    """

    def __init__(
        self,
        neural_model: Optional["TruthScoringModel"] = None,
        neural_tokenizer=None,
        device: Optional[torch.device] = None,
    ) -> None:
        self.neural_model = neural_model
        self.neural_tokenizer = neural_tokenizer
        self.device = device or torch.device("cuda" if torch.cuda.is_available() else "cpu")

    def evaluate(
        self,
        propositions: List[Tuple[str, float]],  # (texto, peso_de_prioridade_L2)
        knowledge_base: Dict[str, float],        # termo  grau de evidncia no BD
    ) -> List[ParaconsistentValue]:
        """
        Para cada proposio:
           = evidncia favorvel extrada do banco de dados
           = evidncia contrria = 1  f(compatibilidade)
        """
        results: List[ParaconsistentValue] = []
        for prop_text, l2_priority in propositions:
            mu, lam = self._compute_annotations(prop_text, l2_priority, knowledge_base)
            pv = ParaconsistentValue(
                proposition=prop_text,
                mu=mu,
                lam=lam,
                proposition_type=infer_proposition_type(prop_text),
            )
            results.append(pv)

        # Ordena por valor-verdade descendente
        results.sort(key=lambda pv: pv.truth_value, reverse=True)
        return results

    def route_contradictions(
        self,
        values: List[ParaconsistentValue],
    ) -> List[Tuple[ParaconsistentValue, ParaconsistentValue, str, float, str]]:
        """Retorna rotas de pares de proposies classificadas pelo roteador fuzzy."""
        return ManyValuedRouter.route_pairwise(values)

    # ------------------------------------------------------------------ #
    # Anotao  /                                                        #
    # ------------------------------------------------------------------ #

    def _compute_annotations(
        self,
        text: str,
        l2_priority: float,
        kb: Dict[str, float],
    ) -> Tuple[float, float]:
        """
        Calcula (, ) para uma proposio.

        Se um modelo neural estiver disponvel, usa-o para obter anotaes
        compatveis com L3. Caso contrrio, volta para a heurstica original
        baseada apenas no banco de conhecimento e em contradies locais.
        """
        if self.neural_model is not None and self.neural_tokenizer is not None and neural_annotations is not None:
            mu, lam, _, _ = neural_annotations(self.neural_model.to(self.device), self.neural_tokenizer, text)
            # Pequena modulao pela prioridade de L2 para manter a integrao
            mu = min(1.0, mu * (0.5 + 0.5 * l2_priority))
            lam = max(0.0, lam * (1.0 - 0.3 * l2_priority))
            return round(mu, 4), round(lam, 4)

        import re
        tokens = set(re.findall(r"[a-z]+", text.lower()))

        kb_scores = [kb.get(t, 0.0) for t in tokens if kb.get(t, 0.0) > 0]
        mu_kb = sum(kb_scores) / len(kb_scores) if kb_scores else 0.3

        contradiction_detected = self._has_antonym_pair(tokens, kb)
        lam_base = 0.8 if contradiction_detected else (1.0 - mu_kb)

        mu = min(1.0, mu_kb * (0.5 + 0.5 * l2_priority))
        lam = max(0.0, lam_base * (1.0 - 0.3 * l2_priority))

        return round(mu, 4), round(lam, 4)

    ANTONYM_PAIRS = [
        ("quente", "frio"), ("quente", "gelado"),
        ("verdadeiro", "falso"), ("real", "fictcio"),
        ("afirmativo", "negativo"), ("possvel", "impossvel"),
    ]

    def _has_antonym_pair(self, tokens: set, kb: Dict[str, float]) -> bool:
        for a, b in self.ANTONYM_PAIRS:
            if a in tokens and b in tokens:
                return True
        return False

    # ------------------------------------------------------------------ #
    # Consistncia global: verifica se sistema trivializou                 #
    # ------------------------------------------------------------------ #

    @staticmethod
    def check_global_consistency(values: List[ParaconsistentValue]) -> bool:
        """
        Retorna True se o sistema  globalmente consistente
        (nenhuma trivializao  todos os estados vlidos).
        Uma trivializao ocorre se TODAS as proposies so
        'Inconsistente_local' sem nenhum 'Verdadeiro' ou 'Intermedirio'.
        """
        states = {pv.state for pv in values}
        if states == {"Inconsistente_local"}:
            return False   # trivializao global
        return True        # exploso gentil  sistema consistente
# ========== metrics.py ==========
"""
Mtricas de avaliao do pipeline.
==================================
Coerncia L3, similaridade semntica, BLEU/ROUGE (quando disponvel).
"""

import re
from typing import Dict, List, Optional, Any

# Resultado da L4
try:
    from l4_synthesis import SynthesisResult
except Exception:
    SynthesisResult = None  # type: ignore


def coherence_l3(truth_value: float, state: str, contradiction: float) -> Dict[str, float]:
    """
    Mtricas de coerncia com a camada L3.
    - truth_value alto e contradio baixa = bom.
    - state "Falso" ou "Indeterminado" com truth_value baixo = esperado coerente.
    """
    # Score de coerncia: valor alto  bom quando no h trivializao
    contradiction_penalty = abs(contradiction)  # contradio extrema penaliza
    coherence = max(0.0, 1.0 - contradiction_penalty) * (0.5 + 0.5 * truth_value)
    return {
        "coherence_score": round(coherence, 4),
        "truth_value": truth_value,
        "contradiction_abs": abs(contradiction),
    }


def tokenize_pt(text: str) -> List[str]:
    """Tokenizao simples para BLEU/ROUGE: palavras em minsculo."""
    return re.findall(r"[a-z]+", text.lower())


def bleu_sentence(reference: str, hypothesis: str, max_n: int = 2) -> float:
    """
    BLEU simplificado por frase (n-gram precision at max_n).
    Retorna valor em [0, 1].
    """
    ref_tok = tokenize_pt(reference)
    hyp_tok = tokenize_pt(hypothesis)
    if not hyp_tok:
        return 0.0
    if not ref_tok:
        return 0.0
    p_n = []
    for n in range(1, max_n + 1):
        ref_ngrams = [tuple(ref_tok[i : i + n]) for i in range(len(ref_tok) - n + 1)]
        hyp_ngrams = [tuple(hyp_tok[i : i + n]) for i in range(len(hyp_tok) - n + 1)]
        if not hyp_ngrams:
            continue
        matches = sum(1 for g in hyp_ngrams if g in ref_ngrams)
        p_n.append(matches / len(hyp_ngrams))
    if not p_n:
        return 0.0
    # Mdia geomtrica das precisions
    prod = 1.0
    for p in p_n:
        prod *= p
    return prod ** (1.0 / len(p_n))


def rouge_l_sentence(reference: str, hypothesis: str) -> float:
    """
    ROUGE-L simplificado (LCS de palavras).
    Retorna F1 em [0, 1].
    """
    ref_tok = tokenize_pt(reference)
    hyp_tok = tokenize_pt(hypothesis)
    if not ref_tok or not hyp_tok:
        return 0.0
    # LCS por palavras
    m, n = len(ref_tok), len(hyp_tok)
    dp = [[0] * (n + 1) for _ in range(m + 1)]
    for i in range(1, m + 1):
        for j in range(1, n + 1):
            if ref_tok[i - 1] == hyp_tok[j - 1]:
                dp[i][j] = dp[i - 1][j - 1] + 1
            else:
                dp[i][j] = max(dp[i - 1][j], dp[i][j - 1])
    lcs = dp[m][n]
    prec = lcs / n if n else 0
    rec = lcs / m if m else 0
    if prec + rec == 0:
        return 0.0
    return 2 * prec * rec / (prec + rec)


def semantic_similarity(reference: str, hypothesis: str) -> float:
    """
    Similaridade por embeddings (se sentence-transformers disponvel).
    Caso contrrio, retorna -1.0 para indicar indisponvel.
    """
    try:
        from sentence_transformers import SentenceTransformer
        model = SentenceTransformer("sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
        ref_emb = model.encode(reference)
        hyp_emb = model.encode(hypothesis)
        from numpy import dot
        from numpy.linalg import norm
        return float(dot(ref_emb, hyp_emb) / (norm(ref_emb) * norm(hyp_emb) + 1e-9))
    except Exception:
        return -1.0


def evaluate_response(
    synthesis_result: "SynthesisResult",
    reference_answer: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Agrega mtricas: coerncia L3 + BLEU/ROUGE (e opcionalmente similaridade)
    quando h resposta de referncia.
    """
    out: Dict[str, Any] = {}
    out["coherence"] = coherence_l3(
        synthesis_result.truth_value,
        synthesis_result.state,
        synthesis_result.contradiction,
    )
    if reference_answer:
        out["bleu"] = round(bleu_sentence(reference_answer, synthesis_result.response), 4)
        out["rouge_l"] = round(rouge_l_sentence(reference_answer, synthesis_result.response), 4)
        sim = semantic_similarity(reference_answer, synthesis_result.response)
        if sim >= 0:
            out["semantic_similarity"] = round(sim, 4)
    return out
# ========== l4_synthesis.py ==========
"""
CAMADA L4  Sntese por Equivalncia Russelliana
==================================================
A verdade cognoscvel por uma IA  sempre uma verdade de EQUIVALNCIA:
o grau de correspondncia entre a proposio refinada (sada de L2/L3)
e os dados do mundo real presentes no banco de dados de treinamento.

Base terica (data/russell.txt): Russell  verdade = correspondncia
entre crena e fato; sntese fundamentada em conceitos, no s estatstica.

Mapeamento Kantiano  IA:
  Intuio Sensvel (emprica)  equivalncia proposio  BD
  Intuio Pura (a priori)      estrutura da rede neural / KB
  Sntese                       clculo de equivalncia mediado
                                 por valores-verdade paraconsistentes

O resultado NO  uma predio de prxima palavra.
 o grau de equivalncia entre o conjunto de juzos e o BD.
"""

from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple
from l3_paraconsistent import ParaconsistentValue
from l4_chain_verification import ChainOfVerificationAgent
import math

try:
    from l4_russell_equivalence import (
        RussellConceptBase,
        build_russell_concept_base,
        score_proposition_by_concepts,
        load_concept_base,
    )
except Exception:
    RussellConceptBase = None  # type: ignore
    build_russell_concept_base = None  # type: ignore
    score_proposition_by_concepts = None  # type: ignore
    load_concept_base = None  # type: ignore


# 
# Estrutura do resultado final
# 

@dataclass
class SynthesisResult:
    """Resultado da sntese russelliana  resposta do sistema."""
    response:          str
    truth_value:       float    # paraconsistente  [0,1]
    certainty:         float    # Gc =      [1,1]
    contradiction:     float    # Gct =  +   1
    state:             str      # Verdadeiro | Falso | Intermedirio | ...
    supporting_evidence: List[str] = field(default_factory=list)
    falsified_hypotheses: List[str] = field(default_factory=list)
    verification_log: List[str] = field(default_factory=list)
    confidence_label:  str = ""

    def __post_init__(self):
        if not self.confidence_label:
            self.confidence_label = self._label()

    def _label(self) -> str:
        v = self.truth_value
        if v >= 0.85:  return "Alta Confiana"
        if v >= 0.65:  return "Confiana Moderada"
        if v >= 0.45:  return "Incerto / Intermedirio"
        if v >= 0.25:  return "Baixa Confiana"
        return "Indeterminado"

    def __str__(self) -> str:
        lines = [
            "" * 60,
            f"  RESPOSTA : {self.response}",
            f"  Estado   : {self.state}  ({self.confidence_label})",
            f"  v-verdade: {self.truth_value:.4f}  |  "
            f"Certeza: {self.certainty:+.4f}  |  "
            f"Contradio: {self.contradiction:+.4f}",
        ]
        if self.supporting_evidence:
            lines.append("  Evidncias de suporte:")
            for ev in self.supporting_evidence[:3]:
                lines.append(f"     {ev}")
        if self.falsified_hypotheses:
            lines.append("  Hipteses falsificadas:")
            for fh in self.falsified_hypotheses[:2]:
                lines.append(f"     {fh}")
        if self.verification_log:
            lines.append("  Chain of Verification:")
            for entry in self.verification_log[:4]:
                lines.append(f"    - {entry}")
        lines.append("" * 60)
        return "\n".join(lines)


# 
# Motor de sntese
# 

class RussellianSynthesisEngine:
    """
    Combina os valores-verdade paraconsistentes (L3) com o banco de
    conhecimento para produzir a sntese final (resposta).

    Sntese fundamentada em conceitos (Russell, russell.txt):
        equivalncia = correspondncia entre crena/proposio e fato (BD).
    O peso de cada proposio incorpora:
      - prioridade L2 (juzo kantiano)
      - certeza paraconsistente (Gc)
      - score conceitual de equivalncia (correspondncia com fatos/KB),
        no apenas agregao estatstica.
    """

    def __init__(
        self,
        knowledge_base: Dict[str, float],
        russell_concept_base: Optional["RussellConceptBase"] = None,
        use_concept_based_weights: bool = True,
        verification_config: Optional[Dict[str, Any]] = None,
    ) -> None:
        """
        knowledge_base: dicionrio termo  grau de evidncia [0,1]
        russell_concept_base: base terica extrada de russell.txt (equivalncia/correspondncia).
        use_concept_based_weights: se True, usa score conceitual na ponderao (recomendado).
        """
        self.kb = knowledge_base
        self.russell_base = russell_concept_base
        self.use_concept_weights = use_concept_based_weights and (russell_concept_base is not None)
        self.verifier = ChainOfVerificationAgent(verification_config)

    def synthesize(
        self,
        pv_list:     List[ParaconsistentValue],
        l2_priorities: Dict[str, float],     # proposicao[:40]  prioridade L2
        prompt:      str,
        kb: Optional[Dict[str, float]] = None,
    ) -> SynthesisResult:
        """
        Produz a SynthesisResult final integrando todas as camadas.
        """
        if not pv_list:
            return SynthesisResult(
                response="Sem hipteses vlidas para sntese.",
                truth_value=0.0, certainty=0.0,
                contradiction=0.0, state="Indeterminado",
            )

        #  Seleciona a hiptese com maior valor-verdade  #
        best = pv_list[0]
        supporting = [pv.proposition for pv in pv_list[1:4] if pv.state != "Falso"]
        falsified  = [pv.proposition for pv in pv_list if pv.state == "Falso"]

        #  Sntese ponderada: L2 + certeza + equivalncia (Russell)  #
        total_w, total_v = 0.0, 0.0
        for pv in pv_list:
            key = pv.proposition[:40]
            l2_w = l2_priorities.get(key, 0.5)
            # Peso base: prioridade kantiana e certeza paraconsistente
            weight = l2_w * (1.0 + max(pv.certainty, 0.0))
            # Peso conceitual: correspondncia proposio  fato (BD), conforme russell.txt
            if self.use_concept_weights and score_proposition_by_concepts is not None and self.russell_base is not None:
                concept_score = score_proposition_by_concepts(pv.proposition, self.kb, self.russell_base)
                weight *= concept_score
            total_v += pv.truth_value * weight
            total_w += weight

        v_final = total_v / total_w if total_w > 0 else best.truth_value

        #  Gera texto de resposta a partir da hiptese best + BD  #
        response = self._generate_response(best, prompt, kb)

        verified_response, verification_log = self.verifier.verify(
            prompt=prompt,
            baseline_response=response,
            context_summary=f"Hiptese principal: {best.proposition} | Estado L3: {best.state} | Certeza: {best.certainty:.2f}",
        )

        return SynthesisResult(
            response=verified_response,
            truth_value=round(v_final, 4),
            certainty=round(best.certainty, 4),
            contradiction=round(best.contradiction, 4),
            state=best.state,
            supporting_evidence=supporting,
            falsified_hypotheses=falsified,
            verification_log=verification_log,
        )

    # ------------------------------------------------------------------ #
    # Gerao de resposta textual                                          #
    # ------------------------------------------------------------------ #

    def _generate_response(self, best_pv: ParaconsistentValue, prompt: str, kb: Optional[Dict[str, float]] = None) -> str:
        """
        Gera resposta a partir da proposio com maior valor-verdade.
        Em produo seria substitudo pelo decoder do LLM com as
        hipteses kantianas como contexto hard-constrained.
        """
        # Extrai conceitos KB com alta evidncia
        kb = kb or self.kb
        top_kb = sorted(kb.items(), key=lambda x: x[1], reverse=True)[:3]
        kb_context = ", ".join(f"{k}({v:.2f})" for k, v in top_kb)

        state = best_pv.state
        v = best_pv.truth_value

        if state == "Verdadeiro":
            prefix = f"Com alta confiana (v={v:.2f}):"
        elif state == "Intermedirio":
            prefix = f"Com valor intermedirio (v={v:.2f}), sem trivializao:"
        elif state == "Inconsistente_local":
            prefix = f"Contradio local detectada (v={v:.2f}), exploso gentil:"
        elif state == "Falso":
            prefix = f"Evidncia insuficiente (v={v:.2f}):"
        else:
            prefix = f"Indeterminado (v={v:.2f}):"

        source_note = (
            "Fontes bibliogrficas/canonicais consultadas: Bertrand Russell, The Problems of Philosophy (cap. XII) + base local do projeto."
            if self.russell_base is not None
            else "Aviso de auditoria: nenhuma fonte bibliogrfica verificvel foi detectada na base local; a sntese L4 foi produzida com base semntica interna e sem citao externa confirmada."
        )

        return f"{prefix} {best_pv.proposition}  [KB: {kb_context}]\n\n[AUDIT L4] {source_note}"

    # ------------------------------------------------------------------ #
    # Verificao do limite fundamental (Crtica da IA Pura)              #
    # ------------------------------------------------------------------ #

    @staticmethod
    def check_fundamental_limits(query: str) -> Optional[str]:
        """
        Detecta perguntas que violam os limites fundamentais da IA
        (seo 10 do modelo): conscincia, imaginao, AGI, etc.
        Retorna aviso ou None.
        """
        limit_keywords = {
            "conscincia":    "IA no possui conscincia  atributo biolgico emergente.",
            "sentimento":     "IA no possui estados afetivos  limitada ao algoritmo.",
            "imaginao":     "Imaginao  liberdade humana (Sartre)  no computvel.",
            "agi":            "AGI  oximoro terico: algoritmo no supera seu criador.",
            "livre arbtrio": "Livre-arbtrio  problema no computvel.",
            "ser humano":     "IA  uma funo limite  mundo real exige mediao humana.",
        }
        q_lower = query.lower()
        for keyword, warning in limit_keywords.items():
            if keyword in q_lower:
                return f" Limite fundamental: {warning}"
        return None
# ========== l1_l2_rag_integration.py ==========
"""
INTEGRAO L1-L2 COM RAG HBRIDO
=================================
Estende as camadas L1 (Conceitos) e L2 (Juzos) para trabalhar com:
- RAG Hbrido (Context Injection + Retrieval Seletivo)
- Domain-Aware Knowledge Base
- Injeo de contexto nas tabelas de conceitos e juzos

Workflow:
1. RAG processa query e detecta domnio
2. Contexto injetado enriquece L1 (ConceptTable)
3. L2 refina juzos usando KB especializado do domnio
4. Sistema retorna tabelas enriquecidas
"""

from dataclasses import dataclass, field
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path

try:
    from l1_concept_table import ConceptTable, ConceptNode
except ImportError:
    ConceptTable = None  # type: ignore
    ConceptNode = None   # type: ignore

try:
    from l2_kantian_judgments import KantianJudgmentEngine, KantianJudgment, EpistemicClassification
except ImportError:
    KantianJudgmentEngine = None  # type: ignore
    KantianJudgment = None         # type: ignore
    EpistemicClassification = None # type: ignore

try:
    from rag_hybrid_context_injection import (
        HybridRAGContextInjectionEngine,
        RAGContext,
        RetrievalStrategy,
        DomainContext,
    )
except ImportError:
    HybridRAGContextInjectionEngine = None  # type: ignore
    RAGContext = None                        # type: ignore
    RetrievalStrategy = None                 # type: ignore
    DomainContext = None                     # type: ignore


@dataclass
class EnrichedL1Output:
    """Sada enriquecida da camada L1 com contexto RAG."""
    concepts: List[ConceptNode] = field(default_factory=list)
    domain: str = "geral"
    kb_terms: Dict[str, float] = field(default_factory=dict)
    injected_docs: int = 0
    domain_confidence: float = 0.0
    system_prompt: str = ""
    rag_context_summary: str = ""


@dataclass
class EnrichedL2Output:
    """Sada enriquecida da camada L2 com contexto RAG."""
    judgments: List[KantianJudgment] = field(default_factory=list)
    domain: str = "geral"
    top_judgment: Optional[KantianJudgment] = None
    domain_specialized_kb: Dict[str, float] = field(default_factory=dict)
    epistemic_evidence: Dict[str, float] = field(default_factory=dict)
    rag_impact_score: float = 0.0


class L1RAGEnricher:
    """
    Enriquece a Camada L1 (Tbua de Conceitos) com contexto RAG.
    
    Workflow:
    1. Recebe query + L1 ConceptTable
    2. Processa com RAG hbrido
    3. Enriquece conceitos com conhecimento injetado
    4. Retorna L1 expandido
    """

    def __init__(
        self,
        concept_table: Optional[Any] = None,
        rag_engine: Optional[HybridRAGContextInjectionEngine] = None,
        config: Optional[Dict[str, Any]] = None,
        verbose: bool = True,
    ):
        if ConceptTable is None:
            raise RuntimeError("l1_concept_table no pde ser importado")

        self.concept_table = concept_table or ConceptTable()
        self.rag_engine = rag_engine or HybridRAGContextInjectionEngine(config=config)
        self.config = config or {}
        self.verbose = verbose

    def extract_and_enrich(
        self,
        query: str,
        auto_detect_domain: bool = True,
    ) -> EnrichedL1Output:
        """
        Extrai conceitos do query e enriquece com contexto RAG.
        
        Etapas:
        1. RAG: Detecta domnio e recupera documentos
        2. Extrai conceitos padro (L1 original)
        3. Enriquece conceitos com KB injetado
        4. Retorna output enriquecido
        """
        # Etapa 1: Processa com RAG
        rag_context = self.rag_engine.process(
            query=query,
            auto_detect_domain=auto_detect_domain,
            strategy=RetrievalStrategy.HYBRID,
        )

        domain = rag_context.domain
        injected_kb = rag_context.injected_knowledge
        retrieved_docs = rag_context.retrieved_documents

        if self.verbose:
            print(f"\n[L1-RAG] Domnio: {domain}")
            print(f"[L1-RAG] Docs injetados: {sum(1 for d in retrieved_docs if d.is_injected)}")
            print(f"[L1-RAG] Docs recuperados: {sum(1 for d in retrieved_docs if not d.is_injected)}")

        # Etapa 2: Extrai conceitos originais (L1)
        concepts = self.concept_table.extract_concepts(query, domain=domain)

        # Etapa 3: Enriquece conceitos com KB injetado
        enriched_concepts = self._enrich_concepts_with_kb(
            concepts=concepts,
            injected_kb=injected_kb,
            domain=domain,
            rag_context=rag_context,
        )

        # Etapa 4: Compila output
        return EnrichedL1Output(
            concepts=enriched_concepts,
            domain=domain,
            kb_terms=injected_kb,
            injected_docs=sum(1 for d in retrieved_docs if d.is_injected),
            domain_confidence=rag_context.confidence_score,
            system_prompt=self.rag_engine.domains[domain].system_prompt,
            rag_context_summary=self._summarize_context(retrieved_docs),
        )

    def _enrich_concepts_with_kb(
        self,
        concepts: List[Any],
        injected_kb: Dict[str, float],
        domain: str,
        rag_context: RAGContext,
    ) -> List[Any]:
        """
        Enriquece cada conceito com informaes do KB injetado e documentos recuperados.
        Adiciona: domnio, contexto de aplicao, fonte cannica.
        """
        if not concepts:
            return concepts

        enriched = []
        for concept in concepts:
            if not isinstance(concept, ConceptNode):
                enriched.append(concept)
                continue

            # Clone do conceito para enriquecer
            enriched_concept = self.concept_table._clone_node(concept)

            # Atribui domnio
            enriched_concept.domain = domain

            # Busca evidncia no KB
            kb_score = injected_kb.get(concept.term.lower(), 0.0)
            if kb_score > 0:
                enriched_concept.definition += f"\n[KB-{domain}: {kb_score:.2f}]"

            # Busca fonte nos documentos
            for doc in rag_context.retrieved_documents:
                if concept.term.lower() in doc.content.lower():
                    enriched_concept.canonical_source = f"{doc.source}"
                    enriched_concept.application_context = doc.truncate(200)
                    break

            enriched.append(enriched_concept)

        return enriched

    def _summarize_context(self, docs: List[Any]) -> str:
        """Cria um resumo textual do contexto recuperado."""
        if not docs:
            return "Nenhum contexto recuperado."

        lines = []
        injected = sum(1 for d in docs if getattr(d, "is_injected", False))
        retrieved = len(docs) - injected

        lines.append(f"Contexto RAG: {injected} injetado(s), {retrieved} recuperado(s)")
        for i, doc in enumerate(docs[:3]):
            source = getattr(doc, "source", "desconhecido")
            score = getattr(doc, "relevance_score", 0.0)
            lines.append(f"  {i+1}. [{source}] score={score:.2f}")

        return "; ".join(lines)


class L2RAGEnricher:
    """
    Enriquece a Camada L2 (Juzos Kantianos) com contexto RAG.
    
    Workflow:
    1. Recebe query + L2 KantianJudgmentEngine
    2. Processa com RAG hbrido (domnio especializado)
    3. Refina juzos usando KB especializado
    4. Retorna L2 expandido com classificao epistemolgica aprimorada
    """

    def __init__(
        self,
        concept_table: Optional[Any] = None,
        judgment_engine: Optional[Any] = None,
        rag_engine: Optional[HybridRAGContextInjectionEngine] = None,
        config: Optional[Dict[str, Any]] = None,
        verbose: bool = True,
    ):
        if KantianJudgmentEngine is None:
            raise RuntimeError("l2_kantian_judgments no pde ser importado")

        self.concept_table = concept_table
        self.judgment_engine = judgment_engine or KantianJudgmentEngine(concept_table)
        self.rag_engine = rag_engine or HybridRAGContextInjectionEngine(config=config)
        self.config = config or {}
        self.verbose = verbose

    def analyze_and_enrich(
        self,
        query: str,
        concepts: Optional[List[Any]] = None,
        auto_detect_domain: bool = True,
    ) -> EnrichedL2Output:
        """
        Analisa query com L2 e enriquece juzos com contexto RAG.
        
        Etapas:
        1. RAG: Recupera contexto especializado do domnio
        2. L2: Gera juzos kantianos (original)
        3. Enriquece: Refina modalidades e evidncias usando KB
        4. Retorna L2 enriquecido
        """
        # Etapa 1: RAG especializado
        rag_context = self.rag_engine.process(
            query=query,
            concepts=[c.term if isinstance(c, ConceptNode) else str(c) for c in (concepts or [])] if concepts else None,
            auto_detect_domain=auto_detect_domain,
            strategy=RetrievalStrategy.HYBRID,
        )

        domain = rag_context.domain
        domain_kb = rag_context.injected_knowledge

        if self.verbose:
            print(f"\n[L2-RAG] Domnio: {domain}")
            print(f"[L2-RAG] KB Terms: {len(domain_kb)}")

        # Etapa 2: Anlise L2 padro
        judgments = self.judgment_engine.infer_from_prompt(query)

        # Etapa 3: Enriquecimento baseado em RAG
        enriched_judgments = self._enrich_judgments_with_kb(
            judgments=judgments,
            domain_kb=domain_kb,
            domain=domain,
            rag_context=rag_context,
        )

        # Calcula top judgment
        top_judgment = max(enriched_judgments, key=lambda j: j.prioridade) if enriched_judgments else None

        # Computa epistemic evidence
        epistemic_evidence = self._compute_epistemic_evidence(enriched_judgments, domain_kb)

        # Computa RAG impact
        rag_impact = self._compute_rag_impact(enriched_judgments, rag_context)

        return EnrichedL2Output(
            judgments=enriched_judgments,
            domain=domain,
            top_judgment=top_judgment,
            domain_specialized_kb=domain_kb,
            epistemic_evidence=epistemic_evidence,
            rag_impact_score=rag_impact,
        )

    def _enrich_judgments_with_kb(
        self,
        judgments: List[Any],
        domain_kb: Dict[str, float],
        domain: str,
        rag_context: RAGContext,
    ) -> List[Any]:
        """
        Enriquece juzos com informao do KB e documentos.
        - Atualiza prioridade baseado em KB relevncia
        - Refina classificao epistemolgica
        - Adiciona evidncia de suporte
        """
        if not judgments:
            return judgments

        enriched = []
        for judgment in judgments:
            if not isinstance(judgment, KantianJudgment):
                enriched.append(judgment)
                continue

            # Extrai termos da proposio
            prop_terms = [w.lower() for w in judgment.proposicao.split() if len(w) > 3]

            # Calcula boost de prioridade baseado em KB
            kb_boost = 0.0
            for term in prop_terms:
                if term in domain_kb:
                    kb_boost += domain_kb[term] * 0.1

            # Refina prioridade
            judgment.prioridade = min(1.0, judgment.prioridade + kb_boost)

            # Refina classificao epistemolgica
            judgment.epistemic_classification = self._refine_epistemic_classification(
                judgment.epistemic_classification,
                domain_kb,
                prop_terms,
            )

            enriched.append(judgment)

        return enriched

    def _refine_epistemic_classification(
        self,
        current_classification: Any,
        domain_kb: Dict[str, float],
        terms: List[str],
    ) -> Any:
        """
        Refina a classificao epistemolgica usando informao do KB.
        """
        if not EpistemicClassification:
            return current_classification

        # Calcula scores baseado em KB
        kb_truth_score = sum(domain_kb.get(t, 0.0) for t in terms) / max(len(terms), 1)
        kb_indeterminacy = 1.0 - kb_truth_score if kb_truth_score < 0.5 else 0.0

        # Cria nova classificao refinada
        refined = EpistemicClassification(
            truth=min(1.0, current_classification.truth + kb_truth_score * 0.2),
            indeterminacy=max(0.0, current_classification.indeterminacy - kb_indeterminacy * 0.1),
            falsity=max(0.0, current_classification.falsity - kb_truth_score * 0.1),
        )

        return refined

    def _compute_epistemic_evidence(
        self,
        judgments: List[Any],
        domain_kb: Dict[str, float],
    ) -> Dict[str, float]:
        """Computa evidncia epistemolgica para cada dimenso."""
        evidence = {
            "truth_evidence": 0.0,
            "indeterminacy_evidence": 0.0,
            "falsity_evidence": 0.0,
            "clarity_evidence": 0.0,
        }

        if not judgments or not domain_kb:
            return evidence

        kb_values = list(domain_kb.values())
        if not kb_values:
            return evidence

        avg_kb_value = sum(kb_values) / len(kb_values)

        evidence["truth_evidence"] = avg_kb_value
        evidence["indeterminacy_evidence"] = 1.0 - avg_kb_value
        evidence["clarity_evidence"] = 1.0 - abs(avg_kb_value - 0.5)

        return evidence

    def _compute_rag_impact(self, judgments: List[Any], rag_context: RAGContext) -> float:
        """Computa o impacto do RAG na qualidade dos juzos."""
        if not judgments:
            return 0.0

        # Baseado na confiana do contexto e qualidade dos juzos
        base_confidence = rag_context.confidence_score
        judgment_quality = sum(
            min(1.0, j.prioridade if hasattr(j, "prioridade") else 0.0)
            for j in judgments
        ) / len(judgments)

        return base_confidence * judgment_quality


class IntegratedL1L2RAGPipeline:
    """
    Pipeline completo integrado de L1 + L2 com RAG Hbrido.
    
    Workflow completo:
    1. Recebe query
    2. RAG Hbrido processa (domain detection + context retrieval)
    3. L1 extrai e enriquece conceitos
    4. L2 analisa e enriquece juzos
    5. Retorna output estruturado com ambas as camadas
    """

    def __init__(
        self,
        config: Optional[Dict[str, Any]] = None,
        verbose: bool = True,
    ):
        self.config = config or {}
        self.verbose = verbose

        # Inicializa componentes
        self.concept_table = ConceptTable() if ConceptTable else None
        self.rag_engine = HybridRAGContextInjectionEngine(config=config, verbose=verbose)
        self.l1_enricher = L1RAGEnricher(
            concept_table=self.concept_table,
            rag_engine=self.rag_engine,
            config=config,
            verbose=verbose,
        )
        self.l2_enricher = L2RAGEnricher(
            concept_table=self.concept_table,
            rag_engine=self.rag_engine,
            config=config,
            verbose=verbose,
        )

    def process(
        self,
        query: str,
        auto_detect_domain: bool = True,
    ) -> Dict[str, Any]:
        """
        Processa query atravs do pipeline completo L1-L2-RAG.
        
        Retorna:
        {
            'query': str,
            'domain': str,
            'l1_output': EnrichedL1Output,
            'l2_output': EnrichedL2Output,
            'compiled_context': str,  # Contexto injetado para LLM
            'system_prompt': str,
            'confidence': float,
        }
        """
        if self.verbose:
            print(f"\n{'='*70}")
            print(f"[L1-L2-RAG PIPELINE] Processando query: {query[:60]}...")
            print(f"{'='*70}")

        # L1: Extrao e enriquecimento de conceitos
        l1_output = self.l1_enricher.extract_and_enrich(
            query=query,
            auto_detect_domain=auto_detect_domain,
        )

        if self.verbose:
            print(f"\n[L1] Conceitos extrados: {len(l1_output.concepts)}")
            print(f"[L1] Domnio: {l1_output.domain}")

        # L2: Anlise e enriquecimento de juzos
        l2_output = self.l2_enricher.analyze_and_enrich(
            query=query,
            concepts=l1_output.concepts,
            auto_detect_domain=False,  # J detectado por L1-RAG
        )

        if self.verbose:
            print(f"\n[L2] Juzos gerados: {len(l2_output.judgments)}")
            if l2_output.top_judgment:
                print(f"[L2] Top judgment: {str(l2_output.top_judgment)[:100]}...")

        # Compila sada final
        return {
            "query": query,
            "domain": l1_output.domain,
            "l1_output": l1_output,
            "l2_output": l2_output,
            "compiled_context": self._compile_final_context(l1_output, l2_output),
            "system_prompt": l1_output.system_prompt,
            "confidence": max(l1_output.domain_confidence, l2_output.rag_impact_score),
            "rag_context_summary": l1_output.rag_context_summary,
        }

    def _compile_final_context(self, l1_output: EnrichedL1Output, l2_output: EnrichedL2Output) -> str:
        """Compila o contexto final para injeo em LLM."""
        lines = [
            "## Contexto Estruturado (L1-L2-RAG)",
            "",
            f"**Domnio Detectado**: {l1_output.domain}",
            f"**Confiana**: {max(l1_output.domain_confidence, l2_output.rag_impact_score):.2%}",
            "",
            "### Camada L1 (Conceitos)",
            f"Conceitos extrados: {len(l1_output.concepts)}",
        ]

        for concept in l1_output.concepts[:5]:
            concept_name = concept.term if hasattr(concept, "term") else str(concept)
            lines.append(f"  - {concept_name}")

        lines.extend([
            "",
            "### Camada L2 (Juzos Kantianos)",
            f"Juzos gerados: {len(l2_output.judgments)}",
        ])

        if l2_output.top_judgment:
            lines.append(f"  **Top Judgment**: {str(l2_output.top_judgment)[:150]}...")

        lines.extend([
            "",
            "### Knowledge Base (Injetado)",
            f"Termos-chave: {len(l2_output.domain_specialized_kb)}",
        ])

        for term, score in list(l2_output.domain_specialized_kb.items())[:5]:
            lines.append(f"  - {term}: {score:.2f}")

        lines.extend([
            "",
            "---",
            "Use o contexto acima para formular uma resposta rigorosa e bem fundamentada.",
            "",
        ])

        return "\n".join(lines)


# 
# Funes de Convenincia
# 

def create_l1_l2_rag_pipeline(
    config: Optional[Dict[str, Any]] = None,
) -> IntegratedL1L2RAGPipeline:
    """Factory para criar pipeline integrado."""
    return IntegratedL1L2RAGPipeline(config=config, verbose=True)


def process_with_l1_l2_rag(query: str, config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """
    Funo de convenincia para processar query com pipeline L1-L2-RAG.
    
    Exemplo:
        result = process_with_l1_l2_rag("O que  conhecimento?")
        print(result["compiled_context"])
    """
    pipeline = create_l1_l2_rag_pipeline(config=config)
    return pipeline.process(query)
# ========== syllogism_module.py ==========
"""
MDULO  Silogismo Cientfico Aristotlico + Paradoxo de Hempel + Popper
=========================================================================
Integrado entre L2 e L3 (etapa 4 e 5 do fluxo).

Filtra as hipteses kantianas pelas 8 regras do silogismo cientfico e
aplica o princpio da falseabilidade: toda concluso  tratada como
FALSA at que se encontre evidncia verdadeira equivalente.

Paradoxo de Hempel implementado como filtro negativo:
    Nem toda palavra posterior pode ser inferida da anterior.
    Objetos irrelevantes no validam uma teoria.
"""

from dataclasses import dataclass
from typing import List, Optional, Tuple
from l2_kantian_judgments import KantianJudgment


# 
# Estrutura de um silogismo
# 

@dataclass
class Syllogism:
    major:      str   # premissa maior (universal)
    minor:      str   # premissa menor (particular/singular)
    conclusion: str   # concluso derivada
    valid:      bool  = True
    violations: List[str] = None

    def __post_init__(self):
        if self.violations is None:
            self.violations = []

    def __str__(self) -> str:
        status = " VLIDO" if self.valid else f" INVLIDO ({'; '.join(self.violations)})"
        return (
            f"  Maior : {self.major}\n"
            f"  Menor : {self.minor}\n"
            f"  Concl.: {self.conclusion}\n"
            f"  Status: {status}"
        )


# 
# As 8 regras do silogismo cientfico
# 

class AristotelianSyllogismValidator:
    """
    Valida um silogismo segundo as 8 regras aristotlicas e retorna
    a lista de violaes (vazia se vlido).
    """

    def validate(self, major: str, minor: str, conclusion: str) -> List[str]:
        violations: List[str] = []
        m_neg = self._is_negative(major)
        n_neg = self._is_negative(minor)
        c_neg = self._is_negative(conclusion)
        m_part = self._is_particular(major)
        n_part = self._is_particular(minor)
        c_part = self._is_particular(conclusion)

        # R1  Apenas trs termos, cada um no mesmo sentido
        terms_m = self._extract_key_terms(major)
        terms_n = self._extract_key_terms(minor)
        terms_c = self._extract_key_terms(conclusion)
        all_terms = terms_m | terms_n | terms_c
        if len(all_terms) > 6:          # heurstica liberal
            violations.append("R1: mais de trs termos distintos detectados")

        # R2  Termo mdio no aparece na concluso
        middle = terms_m & terms_n - terms_c
        if not middle and terms_m & terms_n:
            violations.append("R2: termo mdio pode estar na concluso")

        # R3  Concluso no excede extenso das premissas
        if not c_part and (m_part or n_part):
            violations.append("R3: concluso mais extensa que as premissas")

        # R4  Termo mdio deve ser universal pelo menos uma vez
        if m_part and n_part:
            violations.append("R4: termo mdio nunca  universal")

        # R5  De duas negativas, nada se conclui
        if m_neg and n_neg:
            violations.append("R5: duas premissas negativas  concluso invlida")

        # R6  Duas afirmativas  concluso afirmativa
        if not m_neg and not n_neg and c_neg:
            violations.append("R6: premissas afirmativas exigem concluso afirmativa")

        # R7  De duas particulares, nada se conclui
        if m_part and n_part:
            violations.append("R7: duas premissas particulares  concluso invlida")

        # R8  "Parte Fraca": concluso segue a premissa mais fraca
        if (m_neg or n_neg) and not c_neg:
            violations.append("R8: premissa negativa exige concluso negativa")
        if (m_part or n_part) and not c_part and not c_neg:
            violations.append("R8: premissa particular exige concluso particular")

        return violations

    #  helpers  #

    @staticmethod
    def _is_negative(text: str) -> bool:
        neg_markers = {"no", "nunca", "nenhum", "jamais", "nem", "negativo"}
        return any(w in text.lower().split() for w in neg_markers)

    @staticmethod
    def _is_particular(text: str) -> bool:
        part_markers = {"algum", "alguma", "alguns", "algumas", "certo",
                        "parte", "pode", "possvel"}
        return any(w in text.lower().split() for w in part_markers)

    @staticmethod
    def _extract_key_terms(text: str) -> set:
        stop = {"", "so", "de", "do", "da", "em", "com", "por",
                "para", "este", "esta", "esse", "toda", "todo", "um", "uma"}
        import re
        tokens = re.findall(r"[a-zA-Z]+", text.lower())
        return {t for t in tokens if t not in stop and len(t) > 2}


# 
# Filtro de Hempel (anti-confirmao espria)
# 

class HempelFilter:
    """
    Paradoxo de Hempel: objetos irrelevantes no devem confirmar hipteses.
    Implementado como deteco de correlaes esprias entre termos do
    prompt e termos do banco de dados sem relao semntica real.
    """

    def __init__(self, relevance_threshold: float = 0.25) -> None:
        self.threshold = relevance_threshold

    def is_spurious(self, judgment: KantianJudgment, prompt_terms: set) -> bool:
        """
        Retorna True se a hiptese  provavelmente espria
        (confirmao por objeto irrelevante).
        """
        import re
        hyp_terms = set(re.findall(
            r"[a-zA-Z]+",
            judgment.proposicao.lower()
        ))
        overlap = len(hyp_terms & prompt_terms) / max(len(hyp_terms), 1)
        return overlap < self.threshold   # pouca sobreposio = provvel esprio


# 
# Princpio da Falseabilidade de Popper
# 

class PopperFalsifiability:
    """
    Toda concluso  tratada como FALSA at que se encontre evidncia
    verdadeira equivalente no banco de dados.

    Implementa o princpio do Cisne Negro: a proposio universal
    "todo cisne  branco"  falsa at que seja falsificada por um cisne preto.
    """

    def __init__(self, falsifiability_floor: float = 0.1) -> None:
        """
        falsifiability_floor : score mnimo de evidncia para aceitar
                               a hiptese como no-falsificada.
        """
        self.floor = falsifiability_floor

    def apply(
        self,
        hypotheses: List[Tuple[KantianJudgment, float]],   # (juzo, score_BD)
    ) -> List[Tuple[KantianJudgment, float, bool]]:
        """
        Retorna triplas (juzo, score, falsificada?).
        Hipteses universais afirmativas partem sempre de score 0
        (falsas at prova em contrrio).
        """
        result = []
        for j, score in hypotheses:
            # Proposies universais: presume falso at evidncia forte
            if j.quantidade == "Universal":
                adjusted = score if score >= self.floor else 0.0
                falsified = adjusted < self.floor
            # Singulares: usa o score direto
            else:
                adjusted = score
                falsified = False
            result.append((j, adjusted, falsified))
        return result


# 
# Pipeline integrado
# 

class ScientificSyllogismPipeline:
    """
    Integra: Silogismo Aristotlico + Filtro de Hempel + Falseabilidade.
    Chamado entre L2 e L3.
    """

    def __init__(self) -> None:
        self.validator = AristotelianSyllogismValidator()
        self.hempel    = HempelFilter()
        self.popper    = PopperFalsifiability()

    def run(
        self,
        judgments: List[KantianJudgment],
        prompt_terms: set,
        kb_scores: dict,          # termo_proposicao  score [0,1]
    ) -> List[Tuple[KantianJudgment, float]]:
        """
        Filtra e pontua as hipteses.
        Retorna lista ordenada de (juzo, score_final).
        """
        # 1. Remove hipteses esprias (Hempel)
        non_spurious = [
            j for j in judgments
            if not self.hempel.is_spurious(j, prompt_terms)
        ]

        # 2. Valida via silogismo (usa prioridade L2 como par maior/menor)
        scored: List[Tuple[KantianJudgment, float]] = []
        for j in non_spurious:
            # Constri silogismo sinttico para validao
            major = f"Universal: {j.proposicao}"
            minor = f"Singular: {j.proposicao}"
            conclusion = j.proposicao
            violations = self.validator.validate(major, minor, conclusion)
            penalty = len(violations) * 0.1
            base_score = kb_scores.get(j.proposicao[:30], j.prioridade)
            scored.append((j, max(0.0, base_score - penalty)))

        # 3. Aplica falseabilidade (Popper)
        with_falsifiability = self.popper.apply(scored)

        # 4. Remove falsificadas e reordena
        valid = [
            (j, score)
            for j, score, falsified in with_falsifiability
            if not falsified
        ]
        valid.sort(key=lambda x: x[1], reverse=True)
        return valid
# ========== l6_final_response.py ==========
"""
CAMADA L6  Resposta Final em Texto Fluido
==========================================
Transforma o output estruturado das camadas L1L5 em um texto contnuo,
claro e preciso, com tom profissional e acessvel.

Fluxo de processamento:
  Motor de Raciocnio  Output Estruturado  Sntese  Resposta Final
"""

from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional
from l4_synthesis import SynthesisResult
from layer_titles import LAYER_TITLES

try:
    from l5_generation import generate_with_custom_lm, generate_with_ollama
except Exception:
    generate_with_custom_lm = None  # type: ignore
    generate_with_ollama = None  # type: ignore


@dataclass
class EpistemicContext:
    """Contexto epistemolgico agregado para L6."""
    proposition_states: List[Dict[str, Any]] = field(default_factory=list)
    many_valued_routes: List[Dict[str, Any]] = field(default_factory=list)
    bert_classifications: List[Dict[str, Any]] = field(default_factory=list)
    application_context: str = ""


class FinalResponseEngine:
    """Gera a resposta final em texto fluido a partir da sntese das camadas anteriores."""

    def finalize_response(
        self,
        prompt: str,
        synthesis_result: SynthesisResult,
        epistemic_context: Optional[EpistemicContext] = None,
        generated_text: str = "",
        concepts_summary: str = "",
        top_judgments: str = "",
        agent_context: str = "",
    ) -> str:
        """Produz a resposta final nica e contnua seguindo regras de redao clara."""
        main_text = self._normalize_text(generated_text or synthesis_result.response or "")
        if not main_text:
            return "No h informao suficiente para formular uma resposta final."

        intro = self._build_intro(synthesis_result)
        conclusion = self._build_conclusion(synthesis_result)
        context_note = self._build_context_note(concepts_summary, top_judgments, agent_context, epistemic_context)

        if intro:
            if main_text.lower().startswith(intro.lower()):
                final = main_text
            else:
                final = f"{intro} {main_text}"
        else:
            final = main_text

        if conclusion:
            final = f"{final} {conclusion}"

        if context_note:
            final = f"{final} {context_note}"

        return final.strip()

    def _build_context_note(
        self,
        concepts_summary: str,
        top_judgments: str,
        agent_context: str,
        epistemic_context: Optional[EpistemicContext] = None,
    ) -> str:
        note_parts = []
        if concepts_summary:
            note_parts.append("o raciocnio integrou conceitos extrados e evidncias relevantes")
        if top_judgments:
            note_parts.append("os juzos kantianos foram usados para priorizar hipteses")
        if agent_context:
            note_parts.append("informaes de busca externas tambm foram consideradas")
        if epistemic_context is not None:
            if epistemic_context.application_context:
                note_parts.append("o contexto aplicacional do LogicLMSolver tambm foi considerado")
            if epistemic_context.many_valued_routes:
                note_parts.append("as rotas paraconsistentes do ManyValuedRouter foram analisadas")
            if epistemic_context.bert_classifications:
                note_parts.append("classificaes BERT (T/I/F) das principais hipteses influenciaram a formulao")
        if not note_parts:
            return ""
        return "Essa resposta reflete o processamento integrado das camadas anteriores, com ateno a evidncias e juzos relevantes."  

    # ------------------------------------------------------------------ #
    # Componentes textuais adaptativos                                    #
    # ------------------------------------------------------------------ #

    def _build_intro(self, synthesis_result: SynthesisResult) -> str:
        if synthesis_result.truth_value >= 0.85:
            return "Com base no motor de raciocnio L1L5, a melhor concluso indica"  
        if synthesis_result.truth_value >= 0.65:
            return "A partir da sntese das camadas L1L5, o cenrio mais slido sugere"  
        if synthesis_result.truth_value >= 0.45:
            return "Com certa cautela, a anlise das camadas L1L5 aponta"  
        return "A anlise das camadas L1L5 indica"  

    def _build_conclusion(self, synthesis_result: SynthesisResult) -> str:
        if synthesis_result.state in {"Indeterminado", "N"}:
            return (
                "Esta questo tem uma dimenso genuinamente indeterminada  no por falta de rigor, "
                "mas porque a evidncia emprica ainda no existe. Isso  diferente de 'no sabemos' "
                "  'no h dados para saber'."
            )
        if synthesis_result.state == "Inconsistente_local" or synthesis_result.contradiction > 0.25:
            return "Essa concluso  apresentada como a melhor interpretao disponvel, embora exista uma contradio local que recomenda prudncia."  
        if synthesis_result.truth_value < 0.65:
            return "Dado o grau de incerteza, vale considerar essa resposta como provisria at que evidncias adicionais sejam avaliadas."  
        return ""

    def _normalize_text(self, text: str) -> str:
        normalized_lines = []
        for line in text.splitlines():
            line = " ".join(line.split()).strip()
            if line:
                normalized_lines.append(line)
        return "\n\n".join(normalized_lines)

    def _ensure_single_paragraph(self, text: str) -> str:
        normalized_lines = []
        for line in text.splitlines():
            line = " ".join(line.split()).strip()
            if line:
                normalized_lines.append(line)
        return "\n\n".join(normalized_lines)

    def rewrite_response(
        self,
        prompt: str,
        synthesis_result: SynthesisResult,
        epistemic_context: Optional[EpistemicContext] = None,
        generated_text: str = "",
        concepts_summary: str = "",
        top_judgments: str = "",
        agent_context: str = "",
        provider: str = "template",
        custom_lm_path: str = "",
        ollama_model: str = "doninha8:latest",
        ollama_host: str = "http://localhost:11434",
    ) -> str:
        """Refina a resposta final com um segundo prompt no estilo de um agente escritor."""
        draft = self.finalize_response(
            prompt=prompt,
            synthesis_result=synthesis_result,
            epistemic_context=epistemic_context,
            generated_text=generated_text,
            concepts_summary=concepts_summary,
            top_judgments=top_judgments,
            agent_context=agent_context,
        )
        writer_prompt = self._build_writer_prompt(
            prompt,
            draft,
            synthesis_result,
            epistemic_context,
            concepts_summary,
            top_judgments,
            agent_context,
        )

        if provider == "ollama" and generate_with_ollama:
            text = generate_with_ollama(writer_prompt, model=ollama_model, ollama_host=ollama_host, temperature=0.3)
            if text:
                return text.strip()

        if provider == "custom_lm" and generate_with_custom_lm and custom_lm_path:
            text = generate_with_custom_lm(writer_prompt, custom_lm_path)
            if text:
                return text.strip()

        return self._polish_writer_text(draft)

    def _build_writer_prompt(
        self,
        prompt: str,
        draft: str,
        synthesis_result: SynthesisResult,
        epistemic_context: Optional[EpistemicContext],
        concepts_summary: str,
        top_judgments: str,
        agent_context: str,
    ) -> str:
        lines = [
            "Voc  um agente escritor tcnico e comunicador.",
            "Transforme o raciocnio completo gerado pelas camadas L1 a L6 em um texto fluido, natural, coeso e fcil de ler.",
            "Respeite as proposies e concluses encontradas entre L1 e L6 e mantenha o rigor lgico e tcnico.",
            "Comece direto pela resposta principal, depois explique o caminho se necessrio.",
            "Use linguagem clara, conversacional e precisa e mencione incertezas ou trade-offs de forma elegante quando existirem.",
            "No separe o texto em passos numerados ou listas.",
            "Ao se referir s etapas, use os ttulos de seo designados abaixo:",
            f"L1: {LAYER_TITLES['l1']}",
            f"L2: {LAYER_TITLES['l2']}",
            f"L3: {LAYER_TITLES['l3']}",
            f"L4: {LAYER_TITLES['l4']}",
            f"L5: {LAYER_TITLES['l5']}",
            f"L6: {LAYER_TITLES['l6']}",
            "",
            f"Pergunta do usurio: {prompt}",
            "",
            "Texto preliminar:",
            draft,
            "",
            "Contexto de sntese:",
            f"Resposta de sntese L4: {synthesis_result.response}",
            f"Valor de verdade: {synthesis_result.truth_value:.2f}",
            f"Estado: {synthesis_result.state}",
            f"Certeza: {synthesis_result.certainty:+.2f}",
            f"Contradio: {synthesis_result.contradiction:+.2f}",
        ]
        if concepts_summary:
            lines.extend(["", f"Conceitos L1: {concepts_summary}"])
        if top_judgments:
            lines.extend(["", f"Juzos L2: {top_judgments}"])
        if agent_context:
            lines.extend(["", "Contexto de busca externo:", agent_context])
        if epistemic_context is not None:
            lines.extend(["", "Detalhes epistemolgicos:", self._summarize_epistemic_context(epistemic_context)])
        lines.extend([
            "",
            "Raciocnio completo:",
            "O texto deve sintetizar a extrao de conceitos, os juzos kantianos, a avaliao paraconsistente, a sntese russelliana e a formulao final.",
        ])
        return "\n".join(lines)

    def _summarize_epistemic_context(self, epistemic_context: EpistemicContext) -> str:
        parts: List[str] = []
        if epistemic_context.application_context:
            parts.append(f"Contexto de aplicao: {epistemic_context.application_context}")
        if epistemic_context.proposition_states:
            top_props = epistemic_context.proposition_states[:3]
            summary = ", ".join(
                f"{item.get('state', 'Desconhecido')} ({item.get('truth_value', 'n/a')})" for item in top_props
            )
            parts.append(f"Proposies avaliadas: {summary}")
        if epistemic_context.many_valued_routes:
            parts.append("Rotas paraconsistentes avaliadas.")
        if epistemic_context.bert_classifications:
            parts.append("Classificaes BERT (T/I/F) foram usadas para ajustar prioridades epistemolgicas.")
        return " ".join(parts)

    def _polish_writer_text(self, draft: str) -> str:
        return draft.strip()
# ========== l7_final_text.py ==========
"""
CAMADA L7  Texto Final Definitivo (Automtico e Robusto)
==========================================================
Gera o texto final de alta qualidade a partir do raciocnio acumulado
nas camadas L1 a L6.

A camada L7 funciona como um prompt adicional de escrita: ela recebe os
sumrios das camadas anteriores e transforma o contedo em um nico
bloco contnuo, fluido e persuasivo.

Suporta mltiplos providers:
- ollama: para modelos rodando localmente
- custom_lm: para modelos customizados
- template: fallback que retorna texto sem LLM
"""

from typing import Optional, Dict, Any
import logging
from l4_synthesis import SynthesisResult
from layer_titles import LAYER_TITLES

logger = logging.getLogger(__name__)

try:
    from l5_generation import generate_with_custom_lm
except Exception:
    generate_with_custom_lm = None  # type: ignore

try:
    import ollama
except Exception:
    ollama = None  # type: ignore


class FinalTextEngine:
    """Gera o texto final definitivo a partir do raciocnio L1L6."""

    # Classificao de audincia
    AUDIENCE_PROFILES = {
        "leigo": {
            "description": "Pblico geral sem conhecimento tcnico especializado",
            "style": "Linguagem simples e acessvel, analogias concretas do dia a dia, evitar notao formal, foco em aplicaes prticas e concluses teis",
            "examples": ["o que ", "como funciona", "explicar", "simples"]
        },
        "tcnico": {
            "description": "Profissional da rea com conhecimento tcnico intermedirio",
            "style": "Usar terminologia especfica da rea, incluir referncias conceituais, evitar tabelas de estados complexas, manter rigor tcnico sem excesso de formalismo",
            "examples": ["anlise", "implementao", "mtodo", "tcnica", "profissional"]
        },
        "acadmico": {
            "description": "Pesquisador ou acadmico com formao avanada",
            "style": "Notao completa e formal, referncias bibliogrficas detalhadas, incluir modo debug/disponibilidade de estados internos, rigor acadmico completo",
            "examples": ["teoria", "formal", "demonstrao", "referncia", "acadmico", "pesquisa"]
        }
    }

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Inicializa a FinalTextEngine com configuraes opcionais.
        
        Args:
            config: Dicionrio com configuraes de L7, incluindo:
                - provider: 'ollama', 'custom_lm', ou 'template'
                - model: Nome do modelo (para ollama)
                - temperature: Temperatura de gerao (padro: 0.7)
                - max_tokens: Nmero mximo de tokens (padro: 4096)
                - custom_lm_path: Caminho do modelo customizado (para custom_lm)
        """
        self.config = config or {}
        self.l7_config = self.config.get("l7", {})
        
    def _build_l7_prompt(self, 
                        prompt: str,
                        l1_summary: str,
                        l2_summary: str,
                        l3_summary: str,
                        l4_response: str,
                        l5_text: str,
                        l6_text: str,
                        audience_profile: str = "tcnico",
                        full_synthesis: Optional[str] = None) -> str:
        """
        Constri automaticamente o prompt L7 para gerao de texto final.
        
        Este mtodo agrega todo o raciocnio das camadas L1-L6 e produz
        um prompt bem estruturado e adaptado ao perfil de audincia.
        
        Args:
            prompt: Pergunta/prompt original do usurio
            l1_summary: Resumo de conceitos extrados (L1)
            l2_summary: Resumo de juzos kantianos (L2)
            l3_summary: Resumo de anlise paraconsistente (L3)
            l4_response: Resposta da sntese russelliana (L4)
            l5_text: Texto gerado em L5 (se disponvel)
            l6_text: Texto refinado de L6
            audience_profile: Perfil de audincia ('leigo', 'tcnico', 'acadmico')
            full_synthesis: Sntese completa opcional
            
        Returns:
            String com prompt bem estruturado para gerao automtica
        """
        lines = []
        
        # SEO 1: Instruo Base
        lines.append("Voc  um excelente escritor tcnico e comunicador, especializado em sintetizar raciocnios complexos em textos claros, profundos e agradveis de ler.")
        lines.append("")
        lines.append("Sua funo  gerar o TEXTO FINAL DEFINITIVO a partir de todo o raciocnio desenvolvido nas camadas L1 a L6.")
        lines.append("")
        
        # SEO 2: Contexto do Prompt Original
        lines.append("" * 70)
        lines.append("PROMPT ORIGINAL DO USURIO:")
        lines.append("" * 70)
        lines.append(prompt)
        lines.append("")
        
        # SEO 3: Resumo das Camadas
        lines.append("" * 70)
        lines.append("RACIOCNIO ACUMULADO (CAMADAS L1L6):")
        lines.append("" * 70)
        lines.append(f"L1 - Conceitos Extrados: {l1_summary or 'No disponvel'}")
        lines.append(f"L2 - Juzos Kantianos: {l2_summary or 'No disponvel'}")
        lines.append(f"L3 - Anlise Paraconsistente: {l3_summary or 'No disponvel'}")
        lines.append(f"L4 - Sntese Russelliana: {l4_response or 'No disponvel'}")
        lines.append(f"L5 - Gerao de Resposta: {l5_text or 'No disponvel'}")
        lines.append(f"L6 - Refinamento Final: {l6_text or 'No disponvel'}")
        lines.append("")
        
        # SEO 4: Perfil de Audincia
        profile_data = self.AUDIENCE_PROFILES.get(audience_profile, self.AUDIENCE_PROFILES["tcnico"])
        lines.append("" * 70)
        lines.append(f"PERFIL DE AUDINCIA: {audience_profile.upper()}")
        lines.append("" * 70)
        lines.append(f"Descrio: {profile_data['description']}")
        lines.append(f"Estilo recomendado: {profile_data['style']}")
        lines.append("")
        
        # SEO 5: Diretivas de Formatao (OBRIGATRIAS)
        lines.append("" * 70)
        lines.append("DIRETIVAS DE FORMATAO (OBRIGATRIAS):")
        lines.append("" * 70)
        lines.append(" Formato: Texto fluido, com pargrafos quando necessrio, sem ttulos, subttulos, bullets ou numerao.")
        lines.append(" Abertura: Comece diretamente com a tese ou resposta principal (1-2 frases fortes e claras).")
        lines.append(" Estrutura: Desenvolvimento gradual das premissas, nuances e evoluo do pensamento.")
        lines.append(" Integrao: Harmonize todas as camadas de forma natural, mostrando a evoluo do raciocnio.")
        lines.append(" Tone: Profissional, confiante e acessvel. Explique termos tcnicos quando necessrios.")
        lines.append(" Variao: Use frases de tamanhos variados com transies naturais e sofisticadas.")
        lines.append(" nfase: Destaque ideias importantes via posicionamento e repetio sutil (no bvia).")
        lines.append(" Rigor: Inclua tenses, trade-offs e incertezas com elegncia e maturidade intelectual.")
        lines.append("")
        
        # SEO 6: Tarefa Final
        lines.append("" * 70)
        lines.append("TAREFA:")
        lines.append("" * 70)
        lines.append("Transforme todo esse raciocnio em uma DISSERTAO EXPOSITIVA FLUIDA, COESA E NATURAL.")
        lines.append("Escreva o texto final agora:")
        lines.append("" * 70)
        lines.append("")
        
        return "\n".join(lines)

    @classmethod
    def _classify_audience(cls, prompt: str, l1_summary: str, l2_summary: str, l3_summary: str) -> str:
        """
        Classifica o perfil da audincia baseado no prompt e contexto das camadas.
        Retorna: 'leigo', 'tcnico', ou 'acadmico'
        """
        # Combinar todo o contexto para anlise
        full_context = f"{prompt} {l1_summary} {l2_summary} {l3_summary}".lower()

        # Contar termos tcnicos e indicadores de nvel
        technical_indicators = {
            "leigo": 0,
            "tcnico": 0,
            "acadmico": 0
        }

        # Anlise de vocabulrio e termos
        for profile, data in cls.AUDIENCE_PROFILES.items():
            for keyword in data["examples"]:
                if keyword.lower() in full_context:
                    technical_indicators[profile] += 1

        # Anlise de extenso e complexidade
        prompt_length = len(prompt.split())
        has_formal_terms = any(term in full_context for term in [
            "formal", "demonstrao", "teorema", "axioma", "paradigma",
            "epistemologia", "ontologia", "metafsica", "transcendental"
        ])
        has_technical_jargon = any(term in full_context for term in [
            "lgica paraconsistente", "juzo kantiano", "sntese russelliana",
            "valor de verdade", "contradio", "cognio"
        ])

        # Regras de classificao
        if has_formal_terms or prompt_length > 50 or "referncia" in full_context:
            return "acadmico"
        elif has_technical_jargon or technical_indicators["tcnico"] > technical_indicators["leigo"]:
            return "tcnico"
        elif technical_indicators["leigo"] > 0 or prompt_length < 20:
            return "leigo"
        else:
            # Padro: analisar padro da pergunta
            question_patterns = {
                "acadmico": ["por que", "como se explica", "qual a teoria", "demonstre"],
                "tcnico": ["como implementar", "qual mtodo", "anlise de", "tcnica para"],
                "leigo": ["o que ", "para que serve", "como funciona", "exemplo"]
            }

            for profile, patterns in question_patterns.items():
                if any(pattern in prompt.lower() for pattern in patterns):
                    return profile

            return "tcnico"  # padro seguro

    def _enhance_with_writer_prompt(
        self,
        base_text: str,
        prompt: str,
        audience_profile: str,
        synthesis_result: Optional[SynthesisResult] = None,
        canonical_alerts: Optional[list] = None
    ) -> str:
        """
        Aprimora o texto base usando o prompt de redao (fallback sem LLM).
        Usada quando nenhum provider est disponvel.
        """
        # Aqui poderamos aplicar algumas transformaes/enhancements
        # que no requerem LLM, como formatting, reorganizao, etc.
        return self._build_writer_prompt(
            prompt=prompt,
            l1_summary="",
            l2_summary="",
            l3_summary="",
            l4_response="",
            l5_text="",
            l6_text=base_text,
            synthesis_result=synthesis_result,
            canonical_alerts=canonical_alerts,
            audience_profile=audience_profile
        )


    def finalize_text(
        self,
        prompt: str,
        l1_summary: str = "",
        l2_summary: str = "",
        l3_summary: str = "",
        l4_response: str = "",
        l5_text: str = "",
        l6_text: str = "",
        synthesis_result: Optional[SynthesisResult] = None,
        provider: str = "ollama",
        model: str = "doninha8:latest",
        custom_lm_path: str = "",
        canonical_alerts: Optional[list] = None,
        audience_profile: Optional[str] = None,
        **kwargs) -> str:
        """
        Gera o texto final definitivo de forma automtica e robusta.
        
        Suporta mltiplos providers:
        - ollama: Executa modelos locais via Ollama
        - custom_lm: Usa modelo LM customizado
        - template: Retorna o melhor resultado de L6 sem LLM (fallback)
        
        Args:
            prompt: Pergunta/prompt original do usurio
            l1_summary: Resumo de conceitos (L1)
            l2_summary: Resumo de juzos kantianos (L2)
            l3_summary: Resumo de anlise paraconsistente (L3)
            l4_response: Resposta da sntese (L4)
            l5_text: Texto gerado (L5)
            l6_text: Texto refinado (L6)
            synthesis_result: Resultado da sntese L4
            provider: 'ollama', 'custom_lm', ou 'template'
            model: Nome do modelo Ollama
            custom_lm_path: Caminho do modelo customizado
            canonical_alerts: Alertas de incompatibilidade semntica
            audience_profile: 'leigo', 'tcnico', ou 'acadmico'
            **kwargs: Argumentos adicionais (temperature, max_tokens, etc.)
            
        Returns:
            String com o texto final gerado
        """
        
        # 1. Classificar audincia se no foi fornecida
        if audience_profile is None:
            audience_profile = self._classify_audience(prompt, l1_summary, l2_summary, l3_summary)
        
        # 2. Construir prompt L7 automtico
        l7_prompt = self._build_l7_prompt(
            prompt=prompt,
            l1_summary=l1_summary,
            l2_summary=l2_summary,
            l3_summary=l3_summary,
            l4_response=l4_response,
            l5_text=l5_text,
            l6_text=l6_text,
            audience_profile=audience_profile,
            full_synthesis=synthesis_result.response if synthesis_result else None
        )
        
        # 3. Gerar texto usando o provider selecionado
        generated_text = None
        
        if provider == "ollama" and ollama:
            generated_text = self._generate_with_ollama(
                prompt=l7_prompt,
                model=self.l7_config.get("model", model),
                temperature=kwargs.get("temperature", 0.7),
                max_tokens=kwargs.get("max_tokens", 4096)
            )
        elif provider == "custom_lm" and generate_with_custom_lm:
            custom_path = custom_lm_path or self.l7_config.get("custom_lm_path", "")
            if custom_path:
                generated_text = self._generate_with_custom_lm(
                    prompt=l7_prompt,
                    model_path=custom_path
                )

        if provider == "template":
            return (l6_text or l5_text or l4_response or "").strip()

        if generated_text:
            return generated_text.strip()

        return (l6_text or l5_text or l4_response or "").strip()

    def _generate_with_ollama(self, prompt: str, model: str, temperature: float = 0.7, max_tokens: int = 4096) -> Optional[str]:
        """
        Gera texto usando Ollama (modelos locais).
        
        Args:
            prompt: Prompt para gerao
            model: Nome do modelo (e.g., 'llama2', 'neural-chat', 'mistral')
            temperature: Controla criatividade (0.0-1.0)
            max_tokens: Limite de tokens de sada
            
        Returns:
            Texto gerado ou None se falhar
        """
        try:
            if not ollama:
                logger.error("Ollama no est instalado")
                return None
            
            response = ollama.chat(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                stream=False,
                options={
                    "temperature": temperature,
                    "num_predict": max_tokens,
                    "num_ctx": 8192
                }
            )
            
            generated = response.get("message", {}).get("content", "").strip()
            if generated:
                logger.info(f"L7 (ollama/{model}): Texto gerado com sucesso ({len(generated)} chars)")
                return generated
            else:
                logger.warning("Ollama retornou resposta vazia")
                return None
                
        except Exception as e:
            logger.error(f"Erro ao usar Ollama: {e}")
            return None

    def _generate_with_custom_lm(self, prompt: str, model_path: str) -> Optional[str]:
        """
        Gera texto usando modelo LM customizado.
        
        Args:
            prompt: Prompt para gerao
            model_path: Caminho do modelo customizado
            
        Returns:
            Texto gerado ou None se falhar
        """
        try:
            if not generate_with_custom_lm:
                logger.error("generate_with_custom_lm no est disponvel")
                return None
            
            generated = generate_with_custom_lm(prompt, model_path)
            if generated:
                logger.info(f"L7 (custom_lm): Texto gerado com sucesso ({len(generated)} chars)")
                return generated
            else:
                logger.warning("Custom LM retornou resposta vazia")
                return None
                
        except Exception as e:
            logger.error(f"Erro ao usar Custom LM: {e}")
            return None



    def _build_writer_prompt(
        self,
        prompt: str,
        l1_summary: str,
        l2_summary: str,
        l3_summary: str,
        l4_response: str,
        l5_text: str,
        l6_text: str,
        synthesis_result: Optional[SynthesisResult] = None,
        canonical_alerts: Optional[list] = None,
        audience_profile: str = "tcnico",
    ) -> str:
        lines = [
            "Voc  um excelente escritor tcnico e comunicador, com capacidade de sintetizar raciocnios complexos em textos claros e persuasivos.",
            "Sua funo  gerar o texto final de alta qualidade a partir do raciocnio desenvolvido nas camadas L1 a L6.",
            "Tarefa: transforme todo o raciocnio acumulado nas camadas L1 a L6 em uma dissertao expositiva fluida, coesa e natural, usando pargrafos claros quando for apropriado.",
            "Pblico-alvo: leitor inteligente de nvel intermedirio (no  especialista no tema).",
            "Formato: texto fluido, com pargrafos quando necessrio, sem ttulos, subttulos, bullets ou qualquer marcao.",
            "Estrutura recomendada: comece diretamente com a tese ou resposta principal em 1-2 frases fortes e claras. Em seguida, desenvolva as premissas, nuances e evolues do pensamento.",
            "Integre harmoniosamente o contedo das camadas anteriores, mostrando a evoluo natural do raciocnio e destacando tenses, trade-offs e incertezas com elegncia.",
            "Linguagem: clara, conversacional e precisa. Use termos tcnicos quando necessrios, explicando-os na sequncia.",
            "Estilo: profissional, acessvel, rigoroso e fcil de ler.",
            "",
            f"PERFIL DA AUDINCIA CLASSIFICADO: {audience_profile.upper()}",
        ]

        # Adicionar instrues especficas do perfil
        profile_data = self.AUDIENCE_PROFILES.get(audience_profile, self.AUDIENCE_PROFILES["tcnico"])
        lines.append(f"Descrio do perfil: {profile_data['description']}")
        lines.append(f"Instrues de estilo especficas: {profile_data['style']}")
        lines.append("")
        lines.append(f"Pergunta do usurio: {prompt}")
        lines.append("")
        lines.append("Raciocnio acumulado L1L6:")
        lines.append(f"L1 - {LAYER_TITLES['l1']}: {l1_summary or 'No disponvel.'}")
        lines.append(f"L2 - {LAYER_TITLES['l2']}: {l2_summary or 'No disponvel.'}")
        lines.append(f"L3 - {LAYER_TITLES['l3']}: {l3_summary or 'No disponvel.'}")
        lines.append(f"L4 - {LAYER_TITLES['l4']}: {l4_response or 'No disponvel.'}")
        lines.append(f"L5 - {LAYER_TITLES['l5']}: {l5_text or 'No disponvel.'}")
        lines.append(f"L6 - {LAYER_TITLES['l6']}: {l6_text or 'No disponvel.'}")
        lines.append(f"L7 - {LAYER_TITLES['l7']}: texto final de sntese e redao.")
        lines.append("")

        # Adicionar informaes da sntese L4 se disponvel
        if synthesis_result:
            lines.append(f"Estado da sntese L4: {synthesis_result.state}")
            lines.append(f"Valor de verdade L4: {synthesis_result.truth_value:.2f}")
            lines.append(f"Certeza L4: {synthesis_result.certainty:+.2f}")
            lines.append(f"Contradio L4: {synthesis_result.contradiction:+.2f}")
            lines.append("")
        else:
            lines.append("Estado da sntese L4: No disponvel.")
            lines.append("")

        # Adicionar alertas de incompatibilidade cannica
        if canonical_alerts:
            lines.append("Alertas de incompatibilidade cannica:")
            for alert in canonical_alerts:
                lines.append(f"- Conceito '{alert['concept']}': {alert['canonical_context']}")
                lines.append(f"  Uso incompatvel detectado: {alert['incompatible_usage']}")
            lines.append("")
            lines.append("IMPORTANTE: Inclua ressalvas no texto final sobre estes usos incompatveis dos conceitos.")
            lines.append("")

        lines.append("Escreva o texto final agora.")

        return "\n".join(lines)

    def _normalize_text(self, text: str) -> str:
        return " ".join(text.split()).strip()

    def _ensure_single_paragraph(self, text: str) -> str:
        return " ".join(text.replace("\n", " ").split()).strip()
# ========== agente_sintese_final.py ==========
"""
Agente de sntese final L1L7.
==============================
Gera um texto final fluido e coeso a partir do raciocnio acumulado
nas camadas L1 a L6, usando a engine L7 j disponvel no projeto.

Uso rpido:
    python agente_sintese_final.py --prompt "Explique..."
"""


import argparse
import json
import os
from pathlib import Path
from typing import Any, Dict, Optional

from l7_final_text import FinalTextEngine


DEFAULT_PROVIDER = os.getenv("FINAL_TEXT_PROVIDER", "ollama")
DEFAULT_MODEL = os.getenv("FINAL_TEXT_MODEL", "doninha8:latest")


def synthesize_final_text(
    prompt: str,
    l1_summary: str = "",
    l2_summary: str = "",
    l3_summary: str = "",
    l4_response: str = "",
    l5_text: str = "",
    l6_text: str = "",
    provider: str = DEFAULT_PROVIDER,
    model: str = DEFAULT_MODEL,
    temperature: float = 0.7,
    max_tokens: int = 4096,
) -> str:
    """Gera o texto final de alta qualidade usando a camada L7."""
    engine = FinalTextEngine(config={"l7": {"model": model}})
    return engine.finalize_text(
        prompt=prompt,
        l1_summary=l1_summary,
        l2_summary=l2_summary,
        l3_summary=l3_summary,
        l4_response=l4_response,
        l5_text=l5_text,
        l6_text=l6_text,
        provider=provider,
        model=model,
        temperature=temperature,
        max_tokens=max_tokens,
    )


def load_json_file(path: str) -> Dict[str, Any]:
    """Carrega um arquivo JSON com o raciocnio acumulado."""
    file_path = Path(path)
    if not file_path.exists():
        return {}
    try:
        return json.loads(file_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Agente de sntese final L1L7")
    parser.add_argument("--prompt", required=True, help="Pergunta ou prompt original para sintetizar.")
    parser.add_argument("--l1", default="", help="Resumo da camada L1 em texto livre.")
    parser.add_argument("--l2", default="", help="Resumo da camada L2 em texto livre.")
    parser.add_argument("--l3", default="", help="Resumo da camada L3 em texto livre.")
    parser.add_argument("--l4", default="", help="Resposta da camada L4 em texto livre.")
    parser.add_argument("--l5", default="", help="Texto de gerao da camada L5.")
    parser.add_argument("--l6", default="", help="Texto refinado da camada L6.")
    parser.add_argument("--json", default="", help="Caminho para JSON com os campos l1..l6.")
    parser.add_argument("--provider", default=DEFAULT_PROVIDER, help="Provider para L7 (ollama, template, custom_lm).")
    parser.add_argument("--model", default=DEFAULT_MODEL, help="Modelo Ollama para a sntese final.")
    parser.add_argument("--temperature", type=float, default=0.7, help="Temperatura da gerao final.")
    parser.add_argument("--max-tokens", type=int, default=4096, help="Mximo de tokens para a sntese final.")
    return parser


def main(argv: Optional[list[str]] = None) -> None:
    args = build_parser().parse_args(argv)

    data: Dict[str, Any] = {}
    if args.json:
        data = load_json_file(args.json)

    final_text = synthesize_final_text(
        prompt=args.prompt,
        l1_summary=data.get("l1_summary", data.get("l1", args.l1)),
        l2_summary=data.get("l2_summary", data.get("l2", args.l2)),
        l3_summary=data.get("l3_summary", data.get("l3", args.l3)),
        l4_response=data.get("l4_response", data.get("l4", args.l4)),
        l5_text=data.get("l5_text", data.get("l5", args.l5)),
        l6_text=data.get("l6_text", data.get("l6", args.l6)),
        provider=args.provider,
        model=args.model,
        temperature=args.temperature,
        max_tokens=args.max_tokens,
    )

    print(final_text)


if __name__ == "__main__":
    main()
# ========== pipeline.py ==========
"""
PIPELINE PRINCIPAL  Modelo Hbrido de LLM
===========================================
Orquestra as 10 etapas do fluxo completo:

  1. Recepo do prompt
  2. Extrao de conceitos [L1]
  3. Refinamento por Juzos Kantianos [L2]
  4. Silogismo Cientfico + Hempel
  5. Falseabilidade de Popper
  6. Avaliao Paraconsistente [L3]
  7. Sntese por Equivalncia [L4]
  8. Gerao da Resposta [L5  opcional]
  9. Resposta Final em Texto Fluida [L6]
 10. Texto Final Definitivo [L7]

Usa config_loader, knowledge_base (KB escalvel + RAG opcional), l5_generation
e opcionalmente o agente de pesquisa para enriquecer contexto.
"""

import sys
import re
import time
import os
from pathlib import Path
from typing import Dict, List, Optional, Any

import torch

from neural_truth_model import TruthScoringModel, load_tokenizer
from l1_concept_table import ConceptTable, ConceptNode, LogicLMSymbolicSolver
from l2_kantian_judgments import KantianJudgmentEngine, KantianJudgment
from syllogism_module import ScientificSyllogismPipeline
from l3_paraconsistent import ParaconsistentEngine, ParaconsistentValue
from l4_synthesis import RussellianSynthesisEngine, SynthesisResult
from l6_final_response import EpistemicContext, FinalResponseEngine
from l7_final_text import FinalTextEngine

try:
    from l4_russell_equivalence import load_concept_base
except Exception:
    load_concept_base = None  # type: ignore

try:
    from config_loader import load_config, PROJECT_ROOT
except Exception:
    load_config = None  # type: ignore
    PROJECT_ROOT = Path(__file__).resolve().parent

try:
    from knowledge_base import get_knowledge_base, SEED_KNOWLEDGE_BASE
except Exception:
    get_knowledge_base = None  # type: ignore
    SEED_KNOWLEDGE_BASE = {}

try:
    from l5_generation import generate_response as l5_generate
except Exception:
    l5_generate = None  # type: ignore

try:
    from agente_busca_web import run_search_for_context
except Exception:
    run_search_for_context = None  # type: ignore

try:
    from agente_sintese_final import synthesize_final_text as synthesize_final_agent
except Exception:
    synthesize_final_agent = None  # type: ignore


def _get_kb(config: Optional[Dict[str, Any]], prompt: str, use_agent: bool) -> Dict[str, float]:
    if get_knowledge_base is None:
        return dict(SEED_KNOWLEDGE_BASE) if SEED_KNOWLEDGE_BASE else {}
    return get_knowledge_base(
        config=config,
        query_for_rag=prompt if use_agent else None,
    )


class HybridLLMPipeline:
    """
    Pipeline completo do Modelo Hbrido de LLM.
    Suporta config, KB escalvel, L5 (gerao), agente opcional e chat.
    """

    def __init__(
        self,
        knowledge_base: Optional[Dict[str, float]] = None,
        config: Optional[Dict[str, Any]] = None,
        verbose: bool = True,
    ) -> None:
        self._config = config or (load_config() if load_config else {})
        self.kb = knowledge_base or _get_kb(self._config, "", False)
        if not self.kb:
            self.kb = dict(SEED_KNOWLEDGE_BASE) if SEED_KNOWLEDGE_BASE else {}
        self.verbose = verbose

        self.L1 = ConceptTable()
        self.L2 = KantianJudgmentEngine(self.L1)
        self.SYL = ScientificSyllogismPipeline()

        # L3
        l3_cfg = self._config.get("l3", {})
        model_path = l3_cfg.get("model_path", "truth_scoring_model.pt")
        backbone_name = l3_cfg.get("backbone", "bert-base-multilingual-cased")
        if not Path(model_path).is_absolute():
            model_path = str(PROJECT_ROOT / model_path)
        neural_model = None
        neural_tokenizer = None
        if os.path.exists(model_path):
            try:
                device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
                neural_tokenizer = load_tokenizer(backbone_name)
                neural_model = TruthScoringModel(backbone_name=backbone_name)
                state = torch.load(model_path, map_location=device)
                neural_model.load_state_dict(state)
                neural_model.to(device)
                if self.verbose:
                    print(f"[L3] Modelo neural carregado de '{model_path}'")
                self.L3 = ParaconsistentEngine(neural_model=neural_model, neural_tokenizer=neural_tokenizer, device=device)
            except Exception as exc:
                if self.verbose:
                    print(f"[L3] Falha ao carregar modelo neural: {exc}")
                self.L3 = ParaconsistentEngine()
        else:
            self.L3 = ParaconsistentEngine()

        # L4
        russell_base = None
        rpath = self._config.get("l4", {}).get("russell_concepts_path", "l4_russell_concepts.json")
        if not Path(rpath).is_absolute():
            rpath = str(PROJECT_ROOT / rpath)
        if load_concept_base and os.path.exists(rpath):
            try:
                russell_base = load_concept_base(rpath)
                if self.verbose:
                    print("[L4] Base russelliana carregada.")
            except Exception:
                pass
        if russell_base is None and load_concept_base:
            try:
                from l4_russell_equivalence import build_russell_concept_base
                russell_base = build_russell_concept_base()
            except Exception:
                pass
        self.L4 = RussellianSynthesisEngine(
            self.kb,
            russell_concept_base=russell_base,
            use_concept_based_weights=(russell_base is not None),
            verification_config=self._config.get("l4_chain_verification", {}),
        )
        self.L6 = FinalResponseEngine()
        self.L7 = FinalTextEngine(config=self._config)  # Passa config para suportar mltiplos providers

    def _infer_domain(self, concepts: List[ConceptNode]) -> str:
        """Inferncia simples de domnio majoritrio a partir dos conceitos extrados."""
        if not concepts:
            return "geral"
        domain_counts = {}
        for concept in concepts:
            domain = concept.domain.lower().strip() if concept.domain else "geral"
            domain_counts[domain] = domain_counts.get(domain, 0) + 1
        return max(domain_counts, key=domain_counts.get)

    def _collect_canonical_sources(self, concepts: List[ConceptNode]) -> List[str]:
        sources = []
        seen = set()
        for concept in concepts:
            source = (concept.canonical_source or "").strip()
            if source and source not in seen:
                seen.add(source)
                sources.append(source)
        return sources

    def _summarize_judgments(self, judgments: List[KantianJudgment]) -> str:
        parts = []
        for idx, judgment in enumerate(judgments[:5], start=1):
            cls = getattr(judgment.epistemic_classification, "classification", "no_classificado")
            truth = getattr(judgment.epistemic_classification, "truth", 0.0)
            ind = getattr(judgment.epistemic_classification, "indeterminacy", 0.0)
            fals = getattr(judgment.epistemic_classification, "falsity", 0.0)
            parts.append(
                f"L2-{idx}: {judgment.proposicao[:120]} | pri={judgment.prioridade:.2f} | class={cls} | T/I/F={truth:.2f}/{ind:.2f}/{fals:.2f}"
            )
        return " ; ".join(parts) if parts else "nenhum juzo L2 disponvel"

    def _summarize_paraconsistent(self, pv_list: List[ParaconsistentValue]) -> str:
        parts = []
        for idx, pv in enumerate(pv_list[:5], start=1):
            parts.append(
                f"L3-{idx}: ={pv.mu:.3f} ={pv.lam:.3f} state={pv.state} truth={pv.truth_value:.3f} certainty={pv.certainty:+.3f} contradiction={pv.contradiction:+.3f}"
            )
        return " ; ".join(parts) if parts else "nenhuma avaliao L3 disponvel"

    def _build_citation_note(self, concepts: List[ConceptNode], agent_context: str) -> str:
        sources = self._collect_canonical_sources(concepts)
        if sources:
            return "Fontes bibliogrficas/canonicais disponveis para auditoria: " + "; ".join(sources[:6])
        if agent_context:
            return "Aviso de auditoria: a base local no exps fontes bibliogrficas verificveis; a correspondncia foi analisada com contexto RAG externo e sem citaes cannicas confirmadas."
        return "Aviso de auditoria: nenhuma fonte bibliogrfica verificvel foi detectada na base local; a resposta foi produzida com base interna e deve ser tratada como no corroborada por referncias externas."

    def _append_audit_block(self, text: str, label: str, details: str) -> str:
        block = f"\n\n[AUDIT {label}] {details}"
        return (text + block).strip()

    def process(
        self,
        prompt: str,
        chat_session: Optional[Any] = None,
        use_agent: Optional[bool] = None,
        skip_l5: bool = False,
        skip_l6: bool = False,
    ) -> SynthesisResult:
        """Executa o pipeline e retorna SynthesisResult (com response j gerada por L5 se ativo)."""
        t0 = time.perf_counter()
        use_agent = use_agent if use_agent is not None else self._config.get("agent", {}).get("use_agent", False)
        if chat_session and hasattr(chat_session, "get_context_for_prompt"):
            prompt_for_kb = chat_session.get_context_for_prompt(prompt, self._config.get("chat", {}).get("max_turns_in_context", 10))
        else:
            prompt_for_kb = prompt

        # KB pode ser enriquecido por RAG (Chroma) quando use_agent
        if use_agent and get_knowledge_base:
            self.kb = _get_kb(self._config, prompt_for_kb, True)
            if not self.kb:
                self.kb = dict(SEED_KNOWLEDGE_BASE) if SEED_KNOWLEDGE_BASE else {}

        self._log("\n" + "" * 60)
        self._log(f"  PROMPT: {prompt[:200]}{'...' if len(prompt) > 200 else ''}")
        self._log("" * 60)

        limit = RussellianSynthesisEngine.check_fundamental_limits(prompt)
        if limit:
            self._log(f"\n{limit}")

        self._log("\n[ETAPA 2] L1  Extrao de Conceitos")
        concepts: List[ConceptNode] = self.L1.extract_concepts(prompt, llm_context=prompt_for_kb, domain="geral", config=self._config)
        domain = self._infer_domain(concepts)
        if domain != "geral":
            # Re-extrai com domnio especfico para enriquecer com KB do domnio
            concepts = self.L1.extract_concepts(prompt, llm_context=prompt_for_kb, domain=domain, config=self._config)
        concepts_summary = ""
        if self.verbose and concepts:
            for c in concepts:
                syns = ", ".join(c.synonyms[:2]) or ""
                self._log(f"   {c.term:15s} | sinnimos: {syns}")
            concepts_summary = "; ".join(f"{c.term}({', '.join(c.synonyms[:2])})" for c in concepts[:8])

        self._log("\n[ETAPA 3] L2  Juzos Kantianos")
        judgments: List[KantianJudgment] = self.L2.refine(prompt, concepts)
        top_judgments = ""
        if judgments:
            top_judgments = "\n".join(j.proposicao for j, _ in list(zip(judgments, [None] * 6))[:6])

        self._log("\n[ETAPAS 4+5] Silogismo + Hempel + Popper")
        prompt_terms = set(re.findall(r"[a-zA-Z]+", prompt.lower()))
        kb_scores = {j.proposicao[:30]: self.kb.get(j.proposicao.split()[0], 0.3) for j in judgments}
        filtered = self.SYL.run(judgments, prompt_terms, kb_scores)
        self._log(f"  {len(judgments)} hipteses  {len(filtered)} aps filtros")

        self._log("\n[ETAPA 6] L3  Lgica Paraconsistente + Classificao Epistemolgica L2")
        props_with_priority = [(j.proposicao, score) for j, score in filtered]
        pv_list: List[ParaconsistentValue] = self.L3.evaluate(props_with_priority, self.kb)
        consistent = self.L3.check_global_consistency(pv_list)
        self._log(f"  Consistncia global: {'' if consistent else ''}")

        epistemic_context = EpistemicContext(
            proposition_states=[
                {
                    "proposition": pv.proposition,
                    "proposition_type": pv.proposition_kind or "Desconhecido",
                    "mu": pv.mu,
                    "lambda": pv.lam,
                    "certainty": pv.certainty,
                    "contradiction": pv.contradiction,
                    "truth_value": pv.truth_value,
                    "state": pv.state,
                }
                for pv in pv_list
            ],
            many_valued_routes=[
                {
                    "left": left.proposition,
                    "left_type": left.proposition_kind or "Desconhecido",
                    "right": right.proposition,
                    "right_type": right.proposition_kind or "Desconhecido",
                    "route": route,
                    "confidence": confidence,
                    "explanation": explanation,
                }
                for left, right, route, confidence, explanation in self.L3.route_contradictions(pv_list)
            ],
            bert_classifications=[
                {
                    "proposition": judgment.proposicao,
                    "priority": judgment.prioridade,
                    "truth": judgment.epistemic_classification.truth,
                    "indeterminacy": judgment.epistemic_classification.indeterminacy,
                    "falsity": judgment.epistemic_classification.falsity,
                    "classification": judgment.epistemic_classification.classification,
                }
                for judgment, _ in filtered[:8]
            ],
            application_context=LogicLMSymbolicSolver.summarize_application_context(concepts),
        )

        self._log("\n[ETAPA 7] L4  Sntese Russelliana")
        l2_priorities = {j.proposicao[:40]: j.prioridade for j, _ in filtered}
        result: SynthesisResult = self.L4.synthesize(pv_list, l2_priorities, prompt)
        l4_result = result
        l5_text = result.response

        # Contexto do agente (busca web/local) se ativo
        agent_context = ""
        if use_agent and run_search_for_context:
            try:
                agent_context = run_search_for_context(prompt)
                if agent_context and self.verbose:
                    self._log("\n[AGENTE] Contexto de busca obtido.")
            except Exception:
                pass

        # L4  nota de fontes e auditoria (agora com contexto RAG disponvel)
        l4_sources_note = self._build_citation_note(concepts, agent_context)
        l4_result.response = self._append_audit_block(
            l4_result.response,
            "L4",
            f"truth={l4_result.truth_value:.4f} certainty={l4_result.certainty:+.4f} contradiction={l4_result.contradiction:+.4f} state={l4_result.state} | {l4_sources_note}"
        )
        result.response = l4_result.response
        l5_text = result.response

        # L5  Gerao de resposta em texto livre
        gen_cfg = self._config.get("generation", {})
        final_cfg = self._config.get("finalization", {})
        l7_cfg = self._config.get("l7", {})

        base_provider = gen_cfg.get("provider", "ollama")
        if base_provider not in {"ollama", "custom_lm", "template"}:
            base_provider = "ollama"

        if final_cfg.get("provider") and final_cfg.get("provider") != base_provider:
            self._log(
                f"[PIPELINE] Ignorando provider de finalizao '{final_cfg.get('provider')}' para usar provider base '{base_provider}'."
            )
        if l7_cfg.get("provider") and l7_cfg.get("provider") != base_provider:
            self._log(
                f"[PIPELINE] Ignorando provider L7 '{l7_cfg.get('provider')}' para usar provider base '{base_provider}'."
            )

        provider = base_provider
        if not skip_l5 and l5_generate and provider != "template":
            final_response = l5_generate(
                prompt,
                result,
                provider=provider,
                concepts_summary=concepts_summary,
                top_judgments=top_judgments,
                custom_lm_path=gen_cfg.get("custom_lm_path", ""),
                ollama_model=gen_cfg.get("ollama_model", "doninha8:latest"),
                ollama_host=gen_cfg.get("ollama_host", "http://localhost:11434"),
            )
            if agent_context and final_response:
                final_response = final_response + "\n\n[Contexto da busca]\n" + agent_context[:800]
            elif agent_context:
                final_response = result.response + "\n\n[Contexto da busca]\n" + agent_context[:800]
            else:
                final_response = final_response or result.response
            result = SynthesisResult(
                response=self._append_audit_block(
                    final_response,
                    "L5",
                    f"provider={provider} model={gen_cfg.get('ollama_model', 'doninha8:latest')} | audit=L1-L5: {self._summarize_judgments(judgments)}"
                ),
                truth_value=result.truth_value,
                certainty=result.certainty,
                contradiction=result.contradiction,
                state=result.state,
                supporting_evidence=result.supporting_evidence,
                falsified_hypotheses=result.falsified_hypotheses,
                confidence_label=result.confidence_label,
            )
        elif agent_context and result.response:
            result = SynthesisResult(
                response=self._append_audit_block(
                    result.response + "\n\n[Contexto da busca]\n" + agent_context[:800],
                    "L5",
                    f"provider={provider} model={gen_cfg.get('ollama_model', 'doninha8:latest')} | audit=L1-L5: {self._summarize_judgments(judgments)}"
                ),
                truth_value=result.truth_value,
                certainty=result.certainty,
                contradiction=result.contradiction,
                state=result.state,
                supporting_evidence=result.supporting_evidence,
                falsified_hypotheses=result.falsified_hypotheses,
                confidence_label=result.confidence_label,
            )

        l5_text = result.response

        if not skip_l6:
            final_text = self.L6.finalize_response(
                prompt=prompt,
                synthesis_result=result,
                epistemic_context=epistemic_context,
                generated_text=result.response,
                concepts_summary=concepts_summary,
                top_judgments=top_judgments,
                agent_context=agent_context,
            )
            final_text = self.L6.rewrite_response(
                prompt=prompt,
                synthesis_result=result,
                epistemic_context=epistemic_context,
                generated_text=final_text,
                concepts_summary=concepts_summary,
                top_judgments=top_judgments,
                agent_context=agent_context,
                provider=base_provider,
                custom_lm_path=final_cfg.get("custom_lm_path", gen_cfg.get("custom_lm_path", "")),
                ollama_model=final_cfg.get("ollama_model", gen_cfg.get("ollama_model", "doninha8:latest")),
                ollama_host=final_cfg.get("ollama_host", gen_cfg.get("ollama_host", "http://localhost:11434")),
            )
            result = SynthesisResult(
                response=self._append_audit_block(
                    final_text,
                    "L6",
                    f"provider={base_provider} model={final_cfg.get('ollama_model', gen_cfg.get('ollama_model', 'doninha8:latest'))} | epistemic={self._summarize_paraconsistent(pv_list)}"
                ),
                truth_value=result.truth_value,
                certainty=result.certainty,
                contradiction=result.contradiction,
                state=result.state,
                supporting_evidence=result.supporting_evidence,
                falsified_hypotheses=result.falsified_hypotheses,
                confidence_label=result.confidence_label,
            )

        l3_summary = ""
        if epistemic_context is not None and epistemic_context.proposition_states:
            top_states = epistemic_context.proposition_states[:3]
            l3_summary = "; ".join(
                f"{item.get('proposition', 'desconhecida')}  {item.get('state', 'n/a')} ({item.get('truth_value', 0):.2f})"
                for item in top_states
            )
            if epistemic_context.many_valued_routes:
                l3_summary += f"; rotas paraconsistentes: {len(epistemic_context.many_valued_routes)}"

        l7_cfg = self._config.get("l7", {})
        # Coletar alertas de incompatibilidade cannica gerados durante L1
        canonical_alerts = LogicLMSymbolicSolver.get_canonical_alerts() if LogicLMSymbolicSolver else []
        
        # === L7  Texto Final Definitivo (Automtico e Integrado) ===
        # Usa o agente de sntese final quando disponvel, com fallback para a engine L7 original.
        if synthesize_final_agent is not None:
            final_text_l7 = synthesize_final_agent(
                prompt=prompt,
                l1_summary=concepts_summary,
                l2_summary=top_judgments,
                l3_summary=l3_summary,
                l4_response=l4_result.response,
                l5_text=l5_text,
                l6_text=result.response,
                provider=base_provider,
                model=l7_cfg.get("model", gen_cfg.get("ollama_model", "doninha8:latest")),
                temperature=l7_cfg.get("temperature", 0.7),
                max_tokens=l7_cfg.get("max_tokens", 4096),
            )
        else:
            final_text_l7 = self.L7.finalize_text(
                prompt=prompt,
                l1_summary=concepts_summary,
                l2_summary=top_judgments,
                l3_summary=l3_summary,
                l4_response=l4_result.response,
                l5_text=l5_text,
                l6_text=result.response,
                synthesis_result=l4_result,
                provider=base_provider,
                model=l7_cfg.get("model", gen_cfg.get("ollama_model", "doninha8:latest")),  # Padro para ollama
                custom_lm_path=l7_cfg.get("custom_lm_path", gen_cfg.get("custom_lm_path", "")),
                canonical_alerts=canonical_alerts,
                temperature=l7_cfg.get("temperature", 0.7),
                max_tokens=l7_cfg.get("max_tokens", 4096),
            )
        result = SynthesisResult(
            response=self._append_audit_block(
                final_text_l7,
                "L7",
                f"provider={base_provider} model={l7_cfg.get('model', gen_cfg.get('ollama_model', 'doninha8:latest'))} | sources={'; '.join(self._collect_canonical_sources(concepts)[:6]) or 'nenhuma fonte local'} | L2={self._summarize_judgments(judgments)} | L3={self._summarize_paraconsistent(pv_list)}"
            ),
            truth_value=result.truth_value,
            certainty=result.certainty,
            contradiction=result.contradiction,
            state=result.state,
            supporting_evidence=result.supporting_evidence,
            falsified_hypotheses=result.falsified_hypotheses,
            confidence_label=result.confidence_label,
        )

        elapsed = (time.perf_counter() - t0) * 1000
        self._log(f"\n[ETAPA 10] L7  Texto Final Definitivo  ({elapsed:.1f} ms)\n")
        self._log(str(result))
        return result

    def _log(self, msg: str) -> None:
        if self.verbose:
            print(msg)

    def repl(self) -> None:
        print("\n" + "" * 60)
        print("  MODELO HBRIDO DE LLM  Fonseca")
        print("  Digite 'sair' para encerrar")
        print("" * 60)
        while True:
            try:
                prompt = input("\nPrompt  ").strip()
            except (EOFError, KeyboardInterrupt):
                break
            if not prompt:
                continue
            if prompt.lower() in {"sair", "exit", "quit"}:
                break
            self.process(prompt)


def main() -> None:
    import argparse
    parser = argparse.ArgumentParser(description="Modelo Hbrido de LLM  Pipeline L1L7")
    parser.add_argument("--prompt", "-p", type=str, help="Pergunta nica (imprime s a resposta)")
    parser.add_argument("--repl", action="store_true", help="Modo interativo")
    parser.add_argument("--demo", action="store_true", help="Rodar demonstrao com prompts fixos")
    parser.add_argument("--config", type=str, help="Caminho para config.yaml")
    args, _ = parser.parse_known_args()

    config = load_config(Path(args.config)) if load_config and args.config else (load_config() if load_config else {})
    pipeline = HybridLLMPipeline(config=config, verbose=not args.prompt)

    if args.prompt:
        r = pipeline.process(args.prompt)
        print(r.response)
        return
    if args.repl:
        pipeline.repl()
        return
    if args.demo:
        for p in ["A gua a 35 graus est quente ou fria?", "O que  a verdade?"]:
            pipeline.process(p)
            print()
        return
    # Default: demo + repl se --repl no argv antigo
    if "--repl" in sys.argv:
        pipeline.repl()
        return
    for p in ["A gua a 35 graus est quente ou fria?", "O que  a verdade?"]:
        pipeline.process(p)
        print()


if __name__ == "__main__":
    main()
# ========== api.py ==========
"""
API REST do Modelo Hbrido de LLM.
==================================
FastAPI expondo /process, /chat e /agent. Usa config, pipeline e chat_session.
"""

import os
import sys
from pathlib import Path
from typing import Any, Dict, Optional
from uuid import uuid4

# Raiz do projeto
ROOT = Path(__file__).resolve().parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

try:
    from fastapi import FastAPI, HTTPException
    from pydantic import BaseModel
except ImportError:
    FastAPI = None  # type: ignore
    HTTPException = None  # type: ignore
    BaseModel = object  # type: ignore


# -----------------------------------------------------------------------------
# Modelos de request/response
# -----------------------------------------------------------------------------
class ProcessRequest(BaseModel):
    prompt: str
    session_id: Optional[str] = None
    use_agent: Optional[bool] = None
    skip_l5: bool = False


class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None


class ProcessResponse(BaseModel):
    response: str
    truth_value: float
    state: str
    certainty: float
    contradiction: float
    confidence_label: str
    session_id: Optional[str] = None


# -----------------------------------------------------------------------------
# Estado global (sesses de chat, pipeline, config)
# -----------------------------------------------------------------------------
def _load_app_state():
    from config_loader import load_config
    from pipeline import HybridLLMPipeline
    from chat_session import ChatSession
    config = load_config()
    pipeline = HybridLLMPipeline(config=config, verbose=False)
    sessions: Dict[str, ChatSession] = {}
    max_turns = config.get("chat", {}).get("max_turns_in_context", 10)
    return config, pipeline, sessions, max_turns


if FastAPI is None:
    app = None
else:
    app = FastAPI(title="Modelo Hbrido de LLM", version="1.0")
    _config, _pipeline, _sessions, _max_turns = _load_app_state()

    @app.get("/health")
    def health():
        return {"status": "ok", "model": "hybrid_llm"}

    @app.post("/process", response_model=ProcessResponse)
    def process(req: ProcessRequest):
        session_id = req.session_id or str(uuid4())
        session = _sessions.get(session_id)
        if session:
            session.add_user(req.prompt)
        try:
            result = _pipeline.process(
                req.prompt,
                chat_session=session,
                use_agent=req.use_agent,
                skip_l5=req.skip_l5,
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
        if session:
            session.add_assistant(result.response)
        return ProcessResponse(
            response=result.response,
            truth_value=result.truth_value,
            state=result.state,
            certainty=result.certainty,
            contradiction=result.contradiction,
            confidence_label=result.confidence_label,
            session_id=session_id,
        )

    @app.post("/chat", response_model=ProcessResponse)
    def chat(req: ChatRequest):
        session_id = req.session_id or str(uuid4())
        if session_id not in _sessions:
            from chat_session import ChatSession
            _sessions[session_id] = ChatSession(max_turns=_max_turns)
        session = _sessions[session_id]
        session.add_user(req.message)
        try:
            result = _pipeline.process(req.message, chat_session=session, use_agent=None, skip_l5=False)
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
        session.add_assistant(result.response)
        return ProcessResponse(
            response=result.response,
            truth_value=result.truth_value,
            state=result.state,
            certainty=result.certainty,
            contradiction=result.contradiction,
            confidence_label=result.confidence_label,
            session_id=session_id,
        )

    class AgentRequest(BaseModel):
        query: str

    @app.post("/agent")
    def agent_search(req: AgentRequest):
        """Chama apenas o agente de pesquisa (busca local + internet)."""
        try:
            from agente_busca_web import run_search_for_context
            text = run_search_for_context(req.query)
            return {"answer": text, "query": req.query}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))


def run_api():
    if app is None:
        print("Instale fastapi e uvicorn: pip install fastapi uvicorn", file=sys.stderr)
        sys.exit(1)
    import uvicorn
    from config_loader import load_config
    cfg = load_config()
    api_cfg = cfg.get("api", {})
    host = api_cfg.get("host", "0.0.0.0")
    port = int(api_cfg.get("port", 8000))
    uvicorn.run("api:app", host=host, port=port, reload=False)


if __name__ == "__main__":
    run_api()